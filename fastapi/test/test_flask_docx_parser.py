#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Flask DOCX 파서 테스트 스크립트
applib/app.py의 DOCX to JSON 변환 기능을 테스트합니다.
"""

import os
import sys
import json
import tempfile
import requests
from pathlib import Path

# applib 경로 추가
sys.path.append('/home/wizice/regulation/fastapi/applib')
sys.path.append('/home/wizice/regulation/fastapi')

# 색상 출력용 ANSI 코드
GREEN = '\033[92m'
RED = '\033[91m'
YELLOW = '\033[93m'
BLUE = '\033[94m'
RESET = '\033[0m'

def print_status(status, message):
    """상태 메시지 출력"""
    if status == 'success':
        print(f"{GREEN}✓{RESET} {message}")
    elif status == 'error':
        print(f"{RED}✗{RESET} {message}")
    elif status == 'warning':
        print(f"{YELLOW}⚠{RESET} {message}")
    elif status == 'info':
        print(f"{BLUE}ℹ{RESET} {message}")
    else:
        print(f"  {message}")

def create_sample_docx():
    """테스트용 DOCX 파일 생성"""
    try:
        import docx

        print_status('info', 'DOCX 샘플 파일 생성 중...')

        # 새 문서 생성
        doc = docx.Document()

        # 제목 추가
        doc.add_heading('정신건강의학과 입원환자 관리 규정', 0)

        # 메타데이터 테이블 추가
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'

        # 메타데이터 입력
        metadata = [
            ('제정일', '2024.01.01.'),
            ('최종개정일', '2024.03.15.'),
            ('최종검토일', '2024.03.20.'),
            ('담당부서', '정신건강의학과'),
            ('관련기준', '4주기 정신의료기관평가기준: CMS.1.1\nJCI Standard 7th Edition: CMS.1')
        ]

        for i, (key, value) in enumerate(metadata):
            table.rows[i].cells[0].text = key
            table.rows[i].cells[1].text = value

        doc.add_paragraph()  # 빈 줄

        # 본문 추가
        doc.add_heading('제1장 총칙', 1)

        doc.add_heading('제1조 (목적)', 2)
        doc.add_paragraph('이 규정은 정신건강의학과 입원환자의 체계적인 관리를 통해 '
                         '환자 안전과 의료 서비스의 질을 향상시키는 것을 목적으로 한다.')

        doc.add_heading('제2조 (적용범위)', 2)
        doc.add_paragraph('이 규정은 정신건강의학과에 입원하는 모든 환자에게 적용된다.')

        doc.add_heading('제2장 입원관리', 1)

        doc.add_heading('제3조 (입원절차)', 2)
        doc.add_paragraph('① 환자의 입원은 정신건강의학과 전문의의 진단과 판단에 따른다.')
        doc.add_paragraph('② 입원 시 다음 각 호의 절차를 따른다.')
        doc.add_paragraph('  1. 환자 또는 보호자의 동의서 작성')
        doc.add_paragraph('  2. 초기 평가 및 진단')
        doc.add_paragraph('  3. 치료계획 수립')

        doc.add_heading('제4조 (퇴원절차)', 2)
        doc.add_paragraph('담당 전문의의 판단에 따라 환자의 상태가 호전되어 '
                         '일상생활이 가능하다고 판단될 때 퇴원할 수 있다.')

        doc.add_paragraph()  # 빈 줄
        doc.add_heading('부칙', 1)
        doc.add_paragraph('제1조 (시행일) 이 규정은 2024년 1월 1일부터 시행한다.')

        # 임시 파일로 저장
        temp_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        doc.save(temp_file.name)

        print_status('success', f'샘플 DOCX 파일 생성 완료: {temp_file.name}')
        return temp_file.name

    except ImportError:
        print_status('error', 'python-docx 라이브러리가 설치되지 않았습니다.')
        print_status('info', 'pip install python-docx 명령으로 설치해주세요.')
        return None
    except Exception as e:
        print_status('error', f'DOCX 파일 생성 실패: {str(e)}')
        return None

def test_direct_parsing(docx_path):
    """직접 파싱 테스트 (Flask 서버 없이)"""
    print("\n" + "="*60)
    print("직접 파싱 테스트 (Flask 서버 없이)")
    print("="*60)

    try:
        # 필요한 모듈 임포트
        import docx
        from utils.docx_parser import extract_metadata
        from utils.sequential_numbers import extract_numbers_from_docx, convert_to_sections_format

        print_status('info', 'DOCX 파일 파싱 중...')

        # 문서 열기
        doc = docx.Document(docx_path)

        # 메타데이터 추출
        metadata = extract_metadata(doc)
        print_status('success', '메타데이터 추출 완료')
        print(f"\n메타데이터:")
        for key, value in metadata.items():
            if value:
                print(f"  - {key}: {value}")

        # 순차적 번호 추출 및 섹션 변환
        extract_results = extract_numbers_from_docx(docx_path)
        sections = convert_to_sections_format(extract_results)

        print_status('success', f'조문 추출 완료: {len(sections)}개')

        # 처음 3개 조문 출력
        print(f"\n처음 3개 조문:")
        for i, section in enumerate(sections[:3], 1):
            print(f"  [{i}] 레벨: {section.get('레벨', '')}, "
                  f"번호: {section.get('번호', '')}, "
                  f"내용: {section.get('내용', '')[:50]}...")

        # 문서 정보 구조화
        doc_title = doc.paragraphs[0].text.strip() if doc.paragraphs else "제목 없음"

        document_info = {
            "규정명": doc_title,
            "내규종류": "규정",
            "제정일": metadata.get("제정일", ""),
            "최종개정일": metadata.get("최종개정일", ""),
            "최종검토일": metadata.get("최종검토일", ""),
            "담당부서": metadata.get("담당부서", ""),
            "조문갯수": len(sections)
        }

        # 최종 JSON 구조
        document_structure = {
            "문서정보": document_info,
            "조문내용": sections
        }

        # JSON 파일로 저장
        json_path = docx_path.replace('.docx', '.json')
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(document_structure, f, ensure_ascii=False, indent=2)

        print_status('success', f'JSON 파일 저장 완료: {json_path}')

        # 파일 크기 확인
        json_size = os.path.getsize(json_path)
        print_status('info', f'JSON 파일 크기: {json_size:,} bytes')

        return True

    except ImportError as e:
        print_status('error', f'필요한 모듈을 임포트할 수 없습니다: {e}')
        return False
    except Exception as e:
        print_status('error', f'파싱 실패: {str(e)}')
        import traceback
        traceback.print_exc()
        return False

def test_flask_server(docx_path):
    """Flask 서버를 통한 테스트"""
    print("\n" + "="*60)
    print("Flask 서버 테스트")
    print("="*60)

    base_url = "http://localhost:5002"

    try:
        # 서버 상태 확인
        print_status('info', 'Flask 서버 연결 테스트 중...')
        response = requests.get(f"{base_url}/", timeout=2)

        if response.status_code == 200:
            print_status('success', 'Flask 서버 연결 성공')
        else:
            print_status('warning', f'서버 응답 코드: {response.status_code}')

        # 파일 업로드 테스트
        print_status('info', 'DOCX 파일 업로드 중...')

        with open(docx_path, 'rb') as f:
            files = {'file': (os.path.basename(docx_path), f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
            response = requests.post(f"{base_url}/upload", files=files)

        if response.status_code == 200:
            data = response.json()

            if data.get('success'):
                print_status('success', '파일 업로드 및 파싱 성공')

                print(f"\n파싱 결과:")
                print(f"  - 원본 파일: {data.get('filename', '')}")
                print(f"  - JSON 파일: {data.get('json_filename', '')}")

                doc_info = data.get('document_info', {})
                print(f"\n문서 정보:")
                print(f"  - 규정명: {doc_info.get('규정명', '')}")
                print(f"  - 담당부서: {doc_info.get('담당부서', '')}")
                print(f"  - 조문개수: {doc_info.get('조문갯수', 0)}")

                sections = data.get('sections', [])
                print(f"\n조문 수: {len(sections)}개")

                # 처음 2개 조문 출력
                if sections:
                    print("\n처음 2개 조문:")
                    for i, section in enumerate(sections[:2], 1):
                        print(f"  [{i}] {json.dumps(section, ensure_ascii=False, indent=4)}")

                return True
            else:
                print_status('error', f"파싱 실패: {data.get('error', '알 수 없는 오류')}")
                return False
        else:
            print_status('error', f'서버 오류: {response.status_code}')
            print(response.text)
            return False

    except requests.exceptions.ConnectionError:
        print_status('warning', 'Flask 서버에 연결할 수 없습니다.')
        print_status('info', '서버 시작 명령: cd /home/wizice/regulation/fastapi/applib && python app.py')
        return False
    except Exception as e:
        print_status('error', f'테스트 실패: {str(e)}')
        return False

def main():
    """메인 테스트 함수"""
    print("\n" + "="*60)
    print("Flask DOCX to JSON 파서 테스트")
    print("="*60)

    # 1. 샘플 DOCX 파일 생성
    docx_path = create_sample_docx()
    if not docx_path:
        print_status('error', 'DOCX 파일 생성에 실패했습니다.')
        return

    # 2. 직접 파싱 테스트
    direct_result = test_direct_parsing(docx_path)

    # 3. Flask 서버 테스트 (선택사항)
    server_result = test_flask_server(docx_path)

    # 4. 결과 요약
    print("\n" + "="*60)
    print("테스트 결과 요약")
    print("="*60)

    if direct_result:
        print_status('success', '직접 파싱 테스트: 성공')
    else:
        print_status('error', '직접 파싱 테스트: 실패')

    if server_result:
        print_status('success', 'Flask 서버 테스트: 성공')
    else:
        print_status('warning', 'Flask 서버 테스트: 실패 (서버가 실행 중이 아닐 수 있음)')

    # 임시 파일 정리
    print_status('info', '임시 파일 정리 중...')
    try:
        os.unlink(docx_path)
        json_path = docx_path.replace('.docx', '.json')
        if os.path.exists(json_path):
            os.unlink(json_path)
        print_status('success', '임시 파일 정리 완료')
    except:
        pass

    print("\n테스트 완료!")

if __name__ == '__main__':
    main()