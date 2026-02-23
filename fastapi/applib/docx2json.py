"""
병원 내부규정 파서 웹 애플리케이션 (이미지 추출 지원)
"""
import os
import sys
import json
import tempfile
import glob
from pathlib import Path

# fastapi 경로를 sys.path에 추가하여 settings 사용 가능하게 함
sys.path.insert(0, str(Path(__file__).parent.parent))
from settings import settings

# Flask 관련 임포트를 조건부로 처리
try:
    from flask import Flask, render_template, request, jsonify, send_file, g
    from flask_cors import CORS
    FLASK_AVAILABLE = True
except ImportError:
    FLASK_AVAILABLE = False

# 선택적 임포트
#try:
#    from krlaws.mariadb import MariaDB
#    MARIADB_AVAILABLE = True
#except ImportError:
#    MARIADB_AVAILABLE = False

import docx
from utils.docx_parser import extract_metadata
from utils.json_converter import convert_to_json_string, save_to_json
from utils.sequential_numbers import extract_numbers_from_docx, convert_to_sections_format

# 로깅 - 선택적
try:
    from app_logger import Log
    LOG_AVAILABLE = True
except ImportError:
    LOG_AVAILABLE = False

# 이미지 추출 기능 임포트
from utils.app_integration import extract_images_for_app, merge_json_with_images


def fix_problematic_sections(sections):
    """
    특정 문제가 있는 섹션들의 레벨과 번호를 수정합니다.

    Args:
        sections: 섹션 리스트

    Returns:
        수정된 섹션 리스트
    """
    if not sections:
        return sections

    # 문제가 있는 텍스트 패턴
    problematic_patterns = [
        "정도관리 결과 허용범위를 초과하거나 이상치 발견 시",
        "영상 이상으로 환자에게 추가검사가 필요하거나"
    ]

    # 2.3.1.1 파일의 채혈 관련 패턴
    blood_collection_patterns = [
        "채혈량은 검사종목에 따른",
        "채혈 시 울혈 및 혈액농축을 막기 위해",
        "항응고제 Tube의 경우"
    ]

    # 2.3.1.1 파일의 숫자 번호 패턴 (레벨 1이지만 실제로는 하위 항목)
    numeric_patterns_2311 = [
        # 채혈 관련 (seq 24-25)
        "검체용기 사용 순서",
        "항응고제 및 Clot Activator가 들어있는",

        # 검사실 전달 관련 (seq 28-30)
        "모든 검체는 바코드가 부착되어",
        "검사실로 운반 전 검체 Tube는",
        "환자의 모든 검체는 감염원이",

        # 검체 보관 관련 (seq 44-48)
        "혈액은행 검체:",
        "혈액학파트 검체:",
        "화학파트 검체:",
        "진단면역학파트 검체:",
        "세균검사용 검체:",

        # 정도관리 관련 (seq 53-64)
        "규칙적인 내부 정도관리와",
        "정도관리 및 검사에 사용되는",
        "파트별, 검사종류별로",
        "주기: 담당자는 검사 전에",
        "결과치 보관: 최소한",
        "결과 허용범위: 정도관리 자료는",
        "결과 이상치 발견 시 조치",
        "CAP(College of American",
        "외부 정도관리에 포함되지",
        "외부 정도관리 결과는"
    ]

    # 레벨 조정이 필요한 텍스트 패턴 (컨텍스트 기반)
    level_adjustment_patterns = [
        "24시 On-call 운영"
    ]

    # 수정 대상 섹션들을 찾고 수정
    for i, section in enumerate(sections):
        if not isinstance(section, dict):
            continue

        content = section.get("내용", "")
        current_level = section.get("레벨", 0)
        current_number = section.get("번호", "")

        # 문제가 있는 패턴이 포함된 섹션 확인
        is_problematic = any(pattern in content for pattern in problematic_patterns)

        if is_problematic and current_level == 1 and (current_number.startswith("제") or current_number in ["4.", "5."]):
            # 이전 섹션이 레벨 4이거나, 또는 이미 수정된 같은 그룹의 섹션인지 확인
            should_fix = False
            if i > 0:
                prev_section = sections[i - 1]
                if isinstance(prev_section, dict):
                    prev_level = prev_section.get("레벨")
                    # 이전 섹션이 레벨 4이거나 같은 그룹의 레벨 5인 경우
                    if prev_level == 4 or (prev_level == 5 and prev_section.get("번호") == "①"):
                        should_fix = True

            if should_fix:
                # 레벨을 5로, 번호를 원형 번호로 수정
                section["레벨"] = 5
                if "정도관리 결과" in content:
                    section["번호"] = "①"
                elif "영상 이상으로" in content:
                    section["번호"] = "②"
                else:
                    # 순차적으로 할당
                    section["번호"] = "①" if current_number in ["제4조", "4."] else "②"

        # 채혈 관련 패턴 처리 (2.3.1.1 파일)
        is_blood_collection = any(pattern in content for pattern in blood_collection_patterns)

        if is_blood_collection and current_level == 1 and current_number in ["4.", "5.", "6."]:
            # 이전 섹션들에서 "(다) 채혈 시 주의사항"을 찾기
            blood_section_found = False
            for j in range(max(0, i-5), i):  # 최근 5개 섹션에서 찾기
                if j < len(sections) and isinstance(sections[j], dict):
                    prev_content = sections[j].get("내용", "")
                    if "채혈 시 주의사항" in prev_content and sections[j].get("레벨") == 4:
                        blood_section_found = True
                        break

            if blood_section_found:
                section["레벨"] = 5
                # 번호를 순차적으로 할당 (①, ②, ③)
                blood_item_number = {"4.": "①", "5.": "②", "6.": "③"}
                section["번호"] = blood_item_number.get(current_number, current_number)

        # 2.3.1.1 파일의 숫자 번호 패턴 처리 (포괄적 수정)
        is_numeric_pattern = any(pattern in content for pattern in numeric_patterns_2311)

        if is_numeric_pattern and current_level == 1 and current_number.endswith('.') and current_number[:-1].isdigit():
            # 숫자 번호에 따라 올바른 레벨과 번호 할당
            number_int = int(current_number[:-1])

            if number_int >= 7:  # 7번 이상의 숫자들
                # 이전 섹션에서 레벨 4인 것을 찾아서 그 하위로 설정
                for j in range(max(0, i-10), i):  # 최근 10개 섹션에서 찾기
                    if j < len(sections) and isinstance(sections[j], dict):
                        prev_level = sections[j].get("레벨")
                        if prev_level == 4:
                            section["레벨"] = 5
                            # 번호를 원형 번호로 변환
                            circle_numbers = ["①", "②", "③", "④", "⑤", "⑥", "⑦", "⑧", "⑨", "⑩"]
                            if number_int <= 16:  # 7-16번을 ①-⑩으로 매핑
                                circle_index = (number_int - 7) % 10
                                section["번호"] = circle_numbers[circle_index]
                            else:  # 17번 이상은 순차적으로
                                circle_index = (number_int - 17) % 10
                                section["번호"] = circle_numbers[circle_index]
                            break

        # 컨텍스트 기반 레벨 조정
        for pattern in level_adjustment_patterns:
            if pattern in content:
                # 이전 섹션의 레벨을 확인하여 적절한 레벨로 조정
                if i > 0:
                    prev_section = sections[i - 1]
                    if isinstance(prev_section, dict):
                        prev_level = prev_section.get("레벨")
                        prev_number = prev_section.get("번호", "")

                        # "24시 On-call 운영"이 레벨 2이고 이전이 레벨 3인 경우
                        if pattern == "24시 On-call 운영" and current_level == 2 and prev_level == 3:
                            section["레벨"] = 3
                            section["번호"] = "2)"  # 올바른 번호로 수정

        # 번호가 붙은 목록에서 앞쪽 번호 제거 (예: "1-1. 냉장보관이..." → "냉장보관이...")
        content = section.get("내용", "")
        if content:
            # X-Y. 패턴으로 시작하는 경우 번호 부분 제거
            import re
            numbered_prefix_pattern = r'^\d+-\d+\.\s+'
            if re.match(numbered_prefix_pattern, content):
                cleaned_content = re.sub(numbered_prefix_pattern, '', content)
                section["내용"] = cleaned_content

    return sections


def remove_empty_level1_sections(sections):
    """
    레벨 1이면서 내용과 번호가 모두 빈값인 섹션들을 제거합니다.

    Args:
        sections: 섹션 리스트

    Returns:
        수정된 섹션 리스트
    """
    if not sections:
        return sections

    filtered_sections = []
    removed_count = 0

    for section in sections:
        # 레벨이 1이면서 내용과 번호가 모두 빈값인지 확인
        if (section.get("레벨") == 1 and
            not section.get("내용", "").strip() and
            not section.get("번호", "").strip()):
            removed_count += 1
            continue

        filtered_sections.append(section)

    if removed_count > 0:
        print(f"[DOCX2JSON] Removed {removed_count} empty level 1 sections")

    # seq 번호 재정렬
    for i, section in enumerate(filtered_sections, 1):
        section["seq"] = i

    return filtered_sections


# Flask 앱 초기화 - Flask가 사용 가능한 경우만
if FLASK_AVAILABLE:
    app = Flask(__name__)
    CORS(app)

    # 설정 파일 로드 시도
    try:
        app.config.from_pyfile('editor.cfg', silent=True)
        app.config['UPLOAD_FOLDER'] = tempfile.gettempdir()
        app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB 제한

        # 이미지 저장 경로 설정 (절대 경로 사용)
        app.config['IMAGE_FOLDER'] = settings.EXTRACTED_IMAGES_DIR
        if not os.path.exists(app.config['IMAGE_FOLDER']):
            os.makedirs(app.config['IMAGE_FOLDER'])

        # 로깅 설정 - LOG가 사용 가능한 경우만
        if LOG_AVAILABLE and 'LOG_LEVEL' in app.config:
            LOG = Log
            log_level = app.config['LOG_LEVEL']
            log_file_path = app.config['LOG_FILE_PATH']
            log_root_path = os.path.join(app.root_path, log_file_path)
            Log.init(logger_name='editor', log_level=log_level, log_filepath=log_root_path)
            Log.debug("LOG_FILE_PATH=%s" % log_root_path)

        # MariaDB 설정 - MARIADB가 사용 가능한 경우만
        if MARIADB_AVAILABLE and 'DATABASE' in app.config:
            DATABASE = app.config['DATABASE']
            STATIC_PATH = app.config.get('STATIC_PATH', '')
            gDB = MariaDB(
                app=app,
                g=g,
                database=app.config["DATABASE"],
                user=app.config["DB_USER"],
                password=app.config["DB_PASSWORD"],
                host=app.config["DB_HOST"]
            )
    except Exception as e:
        # 설정 파일이 없어도 기본 기능은 동작하도록
        app.config['UPLOAD_FOLDER'] = tempfile.gettempdir()
        app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024
        app.config['IMAGE_FOLDER'] = settings.EXTRACTED_IMAGES_DIR
        if not os.path.exists(app.config['IMAGE_FOLDER']):
            os.makedirs(app.config['IMAGE_FOLDER'])
else:
    # Flask 없을 때 기본 설정
    app = None


def process_docx_file(file_path=None, file_content=None, filename=None, wzruleid=None):
    """
    DOCX 파일을 처리하는 독립적인 함수
    Flask 라우트 없이 직접 호출 가능

    Args:
        file_path: DOCX 파일 경로 (선택)
        file_content: DOCX 파일 바이너리 내용 (선택)
        filename: 파일명 (file_content 사용시 필수)
        wzruleid: 규정 ID (이미지 폴더명으로 사용, 선택)

    Returns:
        dict: 파싱 결과
    """
    import time
    import logging

    logger = logging.getLogger(__name__)
    start_time = time.time()

    logger.info(f"[DOCX2JSON] Starting DOCX processing")
    logger.info(f"[DOCX2JSON] File: {filename or os.path.basename(file_path) if file_path else 'unknown'}")
    logger.info(f"[DOCX2JSON] Content size: {len(file_content)/1024:.1f}KB" if file_content else f"File path: {file_path}")

    try:
        # 임시 파일 경로 결정
        if file_path:
            temp_path = file_path
            if not filename:
                filename = os.path.basename(file_path)
        elif file_content and filename:
            # 임시 파일로 저장
            temp_path = os.path.join(tempfile.gettempdir(), filename)
            with open(temp_path, 'wb') as f:
                f.write(file_content)
        else:
            return {'error': '파일 경로나 내용이 필요합니다'}

        # 문서 파싱
        doc_parse_start = time.time()
        logger.info(f"[DOCX2JSON] Creating docx.Document with path: {temp_path}")
        doc = docx.Document(temp_path)
        doc_parse_elapsed = time.time() - doc_parse_start
        logger.info(f"[DOCX2JSON] Document loaded in {doc_parse_elapsed:.2f}s")
        print(f"[DOCX2JSON] Document created successfully")

        metadata_start = time.time()
        metadata = extract_metadata(doc)
        metadata_elapsed = time.time() - metadata_start
        logger.info(f"[DOCX2JSON] Metadata extracted in {metadata_elapsed:.2f}s")

        # 순차적 번호 추출 기능을 사용하여 내용 추출
        extract_start = time.time()
        logger.info(f"[DOCX2JSON] Extracting numbers from docx...")
        extract_results = extract_numbers_from_docx(temp_path)
        extract_elapsed = time.time() - extract_start
        logger.info(f"[DOCX2JSON] Extracted {len(extract_results) if extract_results else 0} results in {extract_elapsed:.2f}s")

        if extract_elapsed > 10:
            logger.warning(f"[BOTTLENECK] DOCX extraction took {extract_elapsed:.2f}s (>10s threshold)")

        # 추출된 결과를 조문내용 형식으로 변환
        convert_start = time.time()
        sections, base_font_family = convert_to_sections_format(extract_results)
        # 문제가 있는 섹션들 수정
        sections = fix_problematic_sections(sections)
        # 빈 섹션들 제거
        sections = remove_empty_level1_sections(sections)
        convert_elapsed = time.time() - convert_start
        logger.info(f"[DOCX2JSON] Converted to sections format in {convert_elapsed:.2f}s")

        # 이미지 추출 (일반 이미지 + 표 이미지 변환 포함)
        image_info = {'images': [], 'image_count': 0}

        # 이미지를 세 곳에 모두 저장: fastapi/static, www/static, applib/static
        # Flask app이 있으면 설정 사용, 없으면 기본 경로 사용
        if app and app.config.get('IMAGE_FOLDER'):
            image_base_folders = [app.config['IMAGE_FOLDER']]
        else:
            # FastAPI 환경에서는 세 경로 모두 사용
            image_base_folders = [
                f'{settings.FASTAPI_DIR}/static/extracted_images',    # nginx /static 경로 (주요)
                settings.EXTRACTED_IMAGES_DIR,         # 기존 경로
                f'{settings.APPLIB_DIR}/static/extracted_images'  # applib 경로
            ]

        # 이미지 폴더명 결정
        # wzruleid가 있으면 wzruleid를 폴더명으로, 없으면 기존 filename 방식 사용
        folder_name = str(wzruleid) if wzruleid else os.path.splitext(filename)[0]

        # 첫 번째 폴더에 이미지 추출
        primary_image_folder = os.path.join(image_base_folders[0], folder_name)
        if not os.path.exists(primary_image_folder):
            os.makedirs(primary_image_folder)

        # 표 이미지 변환 기능 활성화 (전체 문서 범위)
        image_info = extract_images_for_app(temp_path, primary_image_folder,
                                            enable_table_conversion=True,
                                            article_range=None,  # 전체 범위
                                            use_enhanced_table_conversion=False,
                                            wzruleid=wzruleid)

        # 추출된 이미지를 다른 폴더에도 복사
        if len(image_base_folders) > 1 and image_info.get('image_count', 0) > 0:
            import shutil
            for additional_base in image_base_folders[1:]:
                additional_folder = os.path.join(additional_base, folder_name)
                if not os.path.exists(additional_folder):
                    os.makedirs(additional_folder)

                # 모든 이미지 파일 복사
                for img_file in os.listdir(primary_image_folder):
                    src = os.path.join(primary_image_folder, img_file)
                    dst = os.path.join(additional_folder, img_file)
                    if os.path.isfile(src):
                        shutil.copy2(src, dst)
                logger.info(f"[DOCX2JSON] Copied {image_info.get('image_count', 0)} images to {additional_folder}")

        # 문서 정보 생성
        doc_title = ""
        if doc.paragraphs and doc.paragraphs[0].text.strip():
            doc_title = doc.paragraphs[0].text.strip()

        # 관련기준 처리
        related_standards = []
        if "관련기준" in metadata and metadata["관련기준"]:
            related_standards = [item.strip() for item in metadata["관련기준"].split("\n") if item.strip()]

        # 소관부서: DOCX에서 추출 안 되면 PDF에서 추출
        dept = metadata.get("소관부서", "").strip()
        if not dept:
            try:
                import fitz, re as _re
                from utils.pdf_table_cropper import find_matching_pdf
                _pdf_path = find_matching_pdf(file_path or temp_path)
                if _pdf_path:
                    _pdf_doc = fitz.open(_pdf_path)
                    _page_text = _pdf_doc[0].get_text()
                    _pdf_doc.close()
                    _dept_match = _re.search(r'소관부서\s*[:：]\s*([^\n\)]+)', _page_text)
                    if _dept_match:
                        dept = _dept_match.group(1).strip()
            except Exception:
                pass

        # 문서 정보 구조화
        document_info = {
            "규정명": doc_title,
            "내규종류": "규정",
            "규정표기명": doc_title,
            "제정일": metadata.get("제정일", "").strip(),
            "최종개정일": metadata.get("최종개정일", "").strip(),
            "최종검토일": metadata.get("최종검토일", "").strip(),
            "소관부서": dept,
            "유관부서": metadata.get("유관부서", "").strip(),
            "관련기준": related_standards,
            "조문갯수": len(sections),
            "이미지개수": image_info.get("image_count", 0)
        }
        if base_font_family:
            document_info["기본글꼴"] = base_font_family

        # 최종 JSON 구조 생성
        document_structure = {
            "문서정보": document_info,
            "조문내용": sections
        }

        # 섹션과 이미지 연결
        if image_info and image_info.get("images"):
            document_structure = merge_json_with_images(document_structure, image_info)

        # 임시 파일 삭제 (file_content로 생성한 경우만)
        if file_content and os.path.exists(temp_path):
            try:
                os.unlink(temp_path)
            except:
                pass

        total_elapsed = time.time() - start_time
        logger.info(f"[DOCX2JSON] Total processing time: {total_elapsed:.2f}s")
        logger.info(f"[DOCX2JSON] Sections created: {len(sections)}")

        if total_elapsed > 30:
            logger.warning(f"[BOTTLENECK] Total DOCX processing took {total_elapsed:.2f}s (>30s threshold)")

        return {
            'success': True,
            'filename': filename,
            'document_info': document_info,
            'sections': sections,
            'images': image_info.get("images", []),
            'preview_data': document_structure
        }

    except Exception as e:
        elapsed = time.time() - start_time
        logger.error(f"[DOCX2JSON] Error after {elapsed:.2f}s: {str(e)}")
        return {'error': f'처리 중 오류가 발생했습니다: {str(e)}', 'success': False}


# Flask 라우트 - Flask가 사용 가능한 경우만
if FLASK_AVAILABLE and app:
    @app.route('/')
    def index():
        """메인 페이지 렌더링"""
        return render_template('index.html')

    @app.route('/dashboard')
    def dashboard():
        """메인 페이지 렌더링"""
        return render_template('dashboard.html')

    @app.route('/upload', methods=['POST'])
    def upload_file():
        """파일 업로드 및 처리"""
        if 'file' not in request.files:
            return jsonify({'error': '파일이 없습니다'}), 400

        file = request.files['file']

        if file.filename == '':
            return jsonify({'error': '선택된 파일이 없습니다'}), 400

        if not file.filename.endswith(('.docx', '.DOCX')):
            return jsonify({'error': 'DOCX 파일만 업로드 가능합니다'}), 400

        try:
            # 임시 파일로 저장
            temp_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
            file.save(temp_path)

            # 문서 파싱
            doc = docx.Document(temp_path)
            metadata = extract_metadata(doc)
            print(f"metadata = {metadata}")

            # 순차적 번호 추출 기능을 사용하여 내용 추출
            extract_results = extract_numbers_from_docx(temp_path)

            # 추출된 결과를 조문내용 형식으로 변환
            sections, base_font_family = convert_to_sections_format(extract_results)
            # 문제가 있는 섹션들 수정
            sections = fix_problematic_sections(sections)

            # 파일명에서 wzruleid 추출 시도 (파일명이 숫자로 시작하는 경우)
            import re
            wzruleid_match = re.match(r'^(\d+)', file.filename)
            extracted_wzruleid = wzruleid_match.group(1) if wzruleid_match else None

            # 이미지 추출 (일반 이미지 + 표 이미지 변환 포함)
            folder_name = extracted_wzruleid if extracted_wzruleid else os.path.splitext(file.filename)[0]
            image_folder = os.path.join(app.config['IMAGE_FOLDER'], folder_name)
            if not os.path.exists(image_folder):
                os.makedirs(image_folder)

            # 표 이미지 변환 기능 활성화 (전체 문서 범위)
            image_info = extract_images_for_app(temp_path, image_folder,
                                                enable_table_conversion=True,
                                                article_range=None,  # 전체 범위
                                                use_enhanced_table_conversion=False,
                                                wzruleid=extracted_wzruleid)

            # 문서 정보 생성
            doc_title = ""
            if doc.paragraphs and doc.paragraphs[0].text.strip():
                doc_title = doc.paragraphs[0].text.strip()

            # 관련기준 처리
            related_standards = []
            if "관련기준" in metadata and metadata["관련기준"]:
                related_standards = [item.strip() for item in metadata["관련기준"].split("\n") if item.strip()]

            # 문서 정보 구조화
                document_info = {
                "규정명": doc_title,
                "내규종류": "규정",
                "규정표기명": doc_title,
                "제정일": metadata.get("제정일", "").strip(),
                "최종개정일": metadata.get("최종개정일", "").strip(),
                "최종검토일": metadata.get("최종검토일", "").strip(),
                "소관부서": metadata.get("소관부서", "").strip(),
                "유관부서": metadata.get("유관부서", "").strip(),
                "관련기준": related_standards,
                "조문갯수": len(sections),
                "이미지개수": image_info.get("image_count", 0)
            }
            if base_font_family:
                document_info["기본글꼴"] = base_font_family

            # 최종 JSON 구조 생성
            document_structure = {
                "문서정보": document_info,
                "조문내용": sections
            }

            # 섹션과 이미지 연결
            if image_info and image_info.get("images"):
                document_structure = merge_json_with_images(document_structure, image_info)

            # JSON 변환 및 임시 파일 저장
            json_data = convert_to_json_string(document_structure)
            json_filename = os.path.splitext(file.filename)[0] + '.json'
            json_path = os.path.join(app.config['UPLOAD_FOLDER'], json_filename)
            save_to_json(document_structure, json_path)

            return jsonify({
                'success': True,
                'filename': file.filename,
                'json_filename': json_filename,
                'document_info': document_info,
                'sections': sections,
                'images': image_info.get("images", []),
                'preview_data': document_structure
            })

        except Exception as e:
            return jsonify({'error': f'처리 중 오류가 발생했습니다: {str(e)}'}), 500

    @app.route('/preview')
    def preview():
        """미리보기 페이지 렌더링"""
        filename = request.args.get('filename', '')
        if not filename:
            return jsonify({'error': '파일명이 제공되지 않았습니다'}), 400

        json_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)

        try:
            with open(json_path, 'r', encoding='utf-8') as f:
                document_data = f.read()
            return render_template('preview.html', filename=filename, document_data=document_data)
        except Exception as e:
            return jsonify({'error': f'파일을 읽는 중 오류가 발생했습니다: {str(e)}'}), 500


    @app.route('/download')
    def download_file():
        """JSON 파일 다운로드"""
        filename = request.args.get('filename', '')
        if not filename:
            return jsonify({'error': '파일명이 제공되지 않았습니다'}), 400

        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)

        try:
            return send_file(file_path, as_attachment=True)
        except Exception as e:
            return jsonify({'error': f'파일 다운로드 중 오류가 발생했습니다: {str(e)}'}), 500


    @app.route('/image/<path:filename>')
    def get_image(filename):
        """이미지 파일 제공"""
        return send_file(os.path.join(app.config['IMAGE_FOLDER'], filename))

def save_json_file(result, output_path, filename):
    """
    JSON 결과를 파일로 저장

    Args:
        result: 변환 결과 딕셔너리
        output_path: 출력 디렉토리 경로
        filename: 원본 파일명
    """
    # 출력 디렉토리 생성
    os.makedirs(output_path, exist_ok=True)

    # JSON 파일명 생성
    base_name = os.path.splitext(filename)[0]
    json_filename = f"{base_name}.json"
    json_filepath = os.path.join(output_path, json_filename)

    # JSON 파일 저장
    with open(json_filepath, 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    return json_filepath

def process_single_file(docx_file, output_dir):
    """
    단일 DOCX 파일을 처리하는 함수
    """
    print(f"Processing: {docx_file}")

    try:
        # DOCX 파일 처리
        result = process_docx_file(file_path=docx_file)

        if 'error' in result:
            print(f"  Error: {result['error']}")
            return False

        # JSON 파일 저장
        filename = os.path.basename(docx_file)
        json_filepath = save_json_file(result, output_dir, filename)

        print(f"  -> Successfully converted to: {json_filepath}")
        print(f"  -> Document info: {result.get('document_info', {})}")
        print(f"  -> Sections count: {len(result.get('sections', []))}")
        print(f"  -> Images count: {result.get('image_info', {}).get('image_count', 0)}")
        return True

    except Exception as e:
        print(f"  -> Error processing file: {str(e)}")
        return False


def main():
    """
    명령줄 실행을 위한 메인 함수
    Usage: python docx2json.py <docx_file_or_folder> [output_dir]
    """
    if len(sys.argv) < 2:
        print("Usage: python docx2json.py <docx_file_or_folder> [output_dir]")
        print("Examples:")
        print("  python docx2json.py file.docx /output/")
        print("  python docx2json.py file.docx")
        print("  python docx2json.py docx/ docx_json/")
        print("  python docx2json.py docx/")
        sys.exit(1)

    input_path = sys.argv[1]

    # 입력 경로 존재 확인
    if not os.path.exists(input_path):
        print(f"Error: Path not found - {input_path}")
        sys.exit(1)

    # 폴더인지 파일인지 확인
    if os.path.isfile(input_path):
        # 단일 파일 처리
        if not input_path.lower().endswith('.docx'):
            print(f"Error: Not a DOCX file - {input_path}")
            sys.exit(1)

        output_dir = sys.argv[2] if len(sys.argv) > 2 else os.path.dirname(input_path)
        print(f"Single file mode")
        print(f"Output directory: {output_dir}")

        success = process_single_file(input_path, output_dir)
        sys.exit(0 if success else 1)

    elif os.path.isdir(input_path):
        # 폴더 내 모든 DOCX 파일 처리
        # 출력 디렉토리 설정
        if len(sys.argv) > 2:
            output_dir = sys.argv[2]
        else:
            # 기본적으로 docx_json 폴더에 저장
            output_dir = "docx_json"

        # 출력 디렉토리 생성
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
            print(f"Created output directory: {output_dir}")

        # 입력 폴더에서 모든 DOCX 파일 찾기
        docx_pattern = os.path.join(input_path, "*.docx")
        docx_files = glob.glob(docx_pattern)

        if not docx_files:
            print(f"No DOCX files found in directory: {input_path}")
            sys.exit(1)

        print(f"Batch processing mode")
        print(f"Input directory: {input_path}")
        print(f"Output directory: {output_dir}")
        print(f"Found {len(docx_files)} DOCX files")
        print("-" * 50)

        success_count = 0
        error_count = 0

        for docx_file in sorted(docx_files):
            if process_single_file(docx_file, output_dir):
                success_count += 1
            else:
                error_count += 1
            print()  # 빈 줄 추가

        print("=" * 50)
        print(f"Batch processing completed!")
        print(f"Successfully processed: {success_count} files")
        print(f"Failed: {error_count} files")
        print(f"Total: {len(docx_files)} files")

        sys.exit(0 if error_count == 0 else 1)
    else:
        print(f"Error: Invalid path - {input_path}")
        sys.exit(1)

# 실행 모드 결정
if __name__ == '__main__':
    if len(sys.argv) > 1:
        # 명령줄 모드
        main()
    elif FLASK_AVAILABLE and app:
        # Flask 웹 서버 모드
        print("Starting Flask web server mode...")
        print("Access via: http://localhost:5002")
        app.run(host="0.0.0.0", port=5002, debug=True)
    else:
        print("Flask not available. Please install Flask or use command line mode.")
        print("Usage: python docx2json.py <docx_file> [output_dir]")
        sys.exit(1)

