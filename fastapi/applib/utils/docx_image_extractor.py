"""
Word 문서에서 이미지를 추출하는 유틸리티 모듈
"""
import os
import re
import docx
import base64
from io import BytesIO
from typing import Dict, Any, List, Tuple, Optional, BinaryIO
from docx.document import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
try:
    from utils.pdf_table_cropper import PdfTableCropper, find_matching_pdf
    PDF_CROPPER_AVAILABLE = True
except ImportError:
    PDF_CROPPER_AVAILABLE = False

try:
    from utils.docx_table_renderer import DocxTableRenderer as TableConverter
except ImportError:
    from utils.table_to_image import TableToImageConverter as TableConverter


class DocxImageExtractor:
    """
    Word 문서에서 이미지를 추출하는 클래스
    """
    
    def __init__(self, output_dir: str = "extracted_images",
                 enable_table_conversion: bool = True,
                 article_range: Optional[Tuple[int, int]] = None,
                 wzruleid: Optional[str] = None):
        """
        초기화

        Args:
            output_dir: 이미지를 저장할 디렉토리 경로
            enable_table_conversion: 표를 이미지로 변환할지 여부
            article_range: 추출할 조문 범위 (시작조, 끝조) 예: (1, 3)
            wzruleid: 규정 ID (이미지 파일명 접두사로 사용)
        """
        self.output_dir = output_dir
        self.enable_table_conversion = enable_table_conversion
        self.article_range = article_range
        self.wzruleid = wzruleid
        # 출력 디렉토리가 없으면 생성
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        # 표 이미지 변환기 초기화
        if enable_table_conversion:
            self.table_converter = TableConverter(output_dir, article_range=article_range)
    
    def extract_images_from_docx(self, docx_path: str) -> List[Dict[str, Any]]:
        """
        Word 문서에서 이미지를 추출하고 저장합니다.

        Args:
            docx_path: Word 문서 파일 경로

        Returns:
            추출된 이미지 정보 리스트
        """
        doc = docx.Document(docx_path)
        # wzruleid가 있으면 사용, 없으면 파일명 사용
        document_name = str(self.wzruleid) if self.wzruleid else os.path.basename(docx_path).replace('.docx', '')

        # 문서 내용을 순회하면서 이미지 및 관련 텍스트 추출
        image_info_list = self._process_document_content(doc, document_name)

        # 표를 이미지로 변환하여 추가
        if self.enable_table_conversion:
            table_images = self._extract_table_images(docx_path, document_name)
            image_info_list.extend(table_images)

        return image_info_list
    
    def _process_document_content(self, doc: Document, document_name: str) -> List[Dict[str, Any]]:
        """
        문서 내용을 처리하여 이미지 및 관련 텍스트를 추출합니다.

        Args:
            doc: Word 문서 객체
            document_name: 문서 이름

        Returns:
            이미지 정보 리스트
        """
        image_info_list = []
        image_index = 0

        # 단락에서 이미지 추출 (전체 문서 범위)
        for i, paragraph in enumerate(doc.paragraphs):
            para_text = paragraph.text.strip()

            # 단락에 이미지가 있는지 확인
            images = self._extract_images_from_paragraph(paragraph, document_name, image_index)

            if images:
                # 이미지 앞뒤의 텍스트 컨텍스트 가져오기
                prev_text = self._get_previous_text(doc.paragraphs, i, max_paragraphs=3)
                next_text = self._get_next_text(doc.paragraphs, i, max_paragraphs=3)
                
                for img_info in images:
                    img_info["previous_text"] = prev_text
                    img_info["next_text"] = next_text
                    
                    # 현재 단락의 텍스트가 캡션인지 확인
                    current_text = paragraph.text.strip()
                    if current_text and self._is_likely_caption(current_text):
                        img_info["caption"] = current_text
                    
                    image_info_list.append(img_info)
                    image_index += 1
        
        # 표에서 이미지 추출
        for table in doc.tables:
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        images = self._extract_images_from_paragraph(paragraph, document_name, image_index)
                        
                        if images:
                            # 표 내에서 컨텍스트 찾기
                            cell_text = cell.text.strip()
                            table_info = f"표 위치: {row_idx+1}행 {cell_idx+1}열"
                            
                            for img_info in images:
                                img_info["context"] = cell_text if cell_text else "표 내 이미지"
                                img_info["table_location"] = table_info
                                image_info_list.append(img_info)
                                image_index += 1
        
        return image_info_list
    
    def _extract_images_from_paragraph(self, paragraph: Paragraph, document_name: str, image_index: int) -> List[Dict[str, Any]]:
        """
        단락에서 이미지를 추출합니다.
        
        Args:
            paragraph: 단락 객체
            document_name: 문서 이름
            image_index: 이미지 인덱스
            
        Returns:
            이미지 정보 리스트
        """
        image_info_list = []
        
        if not hasattr(paragraph, '_element') or paragraph._element is None:
            return image_info_list
        
        # 이미지 관련 요소 찾기
        inline_shapes = paragraph._element.xpath('.//a:blip')
        if not inline_shapes:
            return image_info_list
        
        # 각 이미지 추출 및 저장
        for i, shape in enumerate(inline_shapes):
            try:
                # 이미지 참조 ID 가져오기
                rId = shape.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if not rId:
                    continue
                
                # 파일 확장자 결정
                part = paragraph.part
                image_part = part.related_parts[rId]
                image_content_type = image_part.content_type
                extension = self._get_file_extension_from_content_type(image_content_type)
                
                # 이미지 파일 이름 생성
                img_file_name = f"{document_name}_image_{image_index}_{i}{extension}"
                img_file_path = os.path.join(self.output_dir, img_file_name)
                
                # 이미지 저장
                with open(img_file_path, 'wb') as f:
                    f.write(image_part.blob)
                
                # 이미지 정보 저장
                # file_path를 웹 경로로 변환
                # /home/wizice/regulation/www/static/... → /static/...
                # /home/wizice/regulation/fastapi/static/... → /static/...
                web_path = img_file_path
                if '/www/static/' in img_file_path:
                    web_path = '/static/' + img_file_path.split('/www/static/')[1]
                elif '/fastapi/static/' in img_file_path:
                    web_path = '/static/' + img_file_path.split('/fastapi/static/')[1]

                image_info = {
                    "image_index": image_index + i,
                    "file_name": img_file_name,
                    "file_path": web_path,  # 웹 경로로 저장
                    "content_type": image_content_type,
                    "paragraph_text": paragraph.text.strip(),
                    "base64_data": self._get_base64_image(image_part.blob, image_content_type)
                }
                
                image_info_list.append(image_info)
            except Exception as e:
                print(f"이미지 추출 중 오류 발생: {e}")
        
        return image_info_list
    
    def _get_previous_text(self, paragraphs: List[Paragraph], current_idx: int, max_paragraphs: int = 3) -> str:
        """
        현재 단락 이전의 텍스트를 가져옵니다.
        
        Args:
            paragraphs: 단락 리스트
            current_idx: 현재 단락 인덱스
            max_paragraphs: 가져올 최대 단락 수
            
        Returns:
            이전 텍스트
        """
        prev_texts = []
        start_idx = max(0, current_idx - max_paragraphs)
        
        for i in range(start_idx, current_idx):
            text = paragraphs[i].text.strip()
            if text:
                prev_texts.append(text)
        
        return '\n'.join(prev_texts)
    
    def _get_next_text(self, paragraphs: List[Paragraph], current_idx: int, max_paragraphs: int = 3) -> str:
        """
        현재 단락 이후의 텍스트를 가져옵니다.
        
        Args:
            paragraphs: 단락 리스트
            current_idx: 현재 단락 인덱스
            max_paragraphs: 가져올 최대 단락 수
            
        Returns:
            이후 텍스트
        """
        next_texts = []
        end_idx = min(len(paragraphs), current_idx + max_paragraphs + 1)
        
        for i in range(current_idx + 1, end_idx):
            text = paragraphs[i].text.strip()
            if text:
                next_texts.append(text)
        
        return '\n'.join(next_texts)
    
    def _is_likely_caption(self, text: str) -> bool:
        """
        텍스트가 이미지 캡션인지 판단합니다.
        
        Args:
            text: 검사할 텍스트
            
        Returns:
            캡션 여부
        """
        # 캡션의 특징적인 패턴
        patterns = [
            r"^\s*<.+>\s*$",          # <그림 1>
            r"^\s*\[.+\]\s*$",        # [그림 1]
            r"^\s*그림\s*\d+",        # 그림 1
            r"^\s*Fig\.?\s*\d+",      # Fig. 1
            r"^\s*Figure\s*\d+",      # Figure 1
            r"^\s*표\s*\d+",          # 표 1
            r"^\s*Table\s*\d+",       # Table 1
            r"^\s*사진\s*\d+"         # 사진 1
        ]
        
        for pattern in patterns:
            if re.search(pattern, text):
                return True
        
        return False
    
    def _get_file_extension_from_content_type(self, content_type: str) -> str:
        """
        컨텐츠 타입에서 파일 확장자를 결정합니다.
        
        Args:
            content_type: 컨텐츠 타입
            
        Returns:
            파일 확장자
        """
        content_type_to_ext = {
            'image/png': '.png',
            'image/jpeg': '.jpg',
            'image/jpg': '.jpg',
            'image/gif': '.gif',
            'image/bmp': '.bmp',
            'image/tiff': '.tiff',
            'image/svg+xml': '.svg'
        }
        
        return content_type_to_ext.get(content_type, '.png')
    
    def _get_base64_image(self, image_blob: bytes, content_type: str) -> str:
        """
        이미지를 Base64 인코딩으로 변환합니다.
        
        Args:
            image_blob: 이미지 바이너리 데이터
            content_type: 컨텐츠 타입
            
        Returns:
            Base64 인코딩된 이미지 데이터
        """
        base64_data = base64.b64encode(image_blob).decode('utf-8')
        return f"data:{content_type};base64,{base64_data}"
    
    def get_image_contexts(self, doc_path: str) -> List[Dict[str, Any]]:
        """
        문서의 모든 이미지와 그 컨텍스트 정보를 가져옵니다.
        
        Args:
            doc_path: 문서 경로
            
        Returns:
            이미지 컨텍스트 정보 리스트
        """
        # 이미지 추출
        image_info_list = self.extract_images_from_docx(doc_path)
        
        # 컨텍스트 정보 구성
        result = []
        for img_info in image_info_list:
            context_info = {
                "image_index": img_info["image_index"],
                "file_name": img_info["file_name"],
                "file_path": img_info["file_path"],
                "previous_text": img_info.get("previous_text", ""),
                "next_text": img_info.get("next_text", ""),
                "caption": img_info.get("caption", ""),
                "paragraph_text": img_info.get("paragraph_text", ""),
                "context": img_info.get("context", ""),
                "table_location": img_info.get("table_location", "")
            }
            result.append(context_info)
        
        return result

    def _extract_table_images(self, docx_path: str, document_name: str) -> List[Dict[str, Any]]:
        """
        문서의 표들을 이미지로 변환하여 추출합니다.
        PDF 파일이 있으면 PDF 크롭(원본 100% 품질)을 우선 사용하고,
        없으면 기존 PIL 렌더링으로 fallback합니다.

        Args:
            docx_path: DOCX 파일 경로
            document_name: 문서 이름

        Returns:
            표 이미지 정보 리스트
        """
        table_images = None

        # 1순위: PDF 크롭 (원본 품질)
        if PDF_CROPPER_AVAILABLE:
            pdf_path = find_matching_pdf(docx_path)
            if pdf_path:
                try:
                    print(f"[DocxImageExtractor] PDF 발견: {pdf_path} → PDF 크롭 모드")
                    cropper = PdfTableCropper(
                        output_dir=self.output_dir,
                        article_range=self.article_range,
                    )
                    table_images = cropper.extract_tables_as_images(pdf_path, document_name)
                except Exception as e:
                    print(f"[DocxImageExtractor] PDF 크롭 실패, PIL fallback: {e}")
                    table_images = None

        # 2순위: 기존 PIL 렌더링 (fallback)
        if table_images is None:
            try:
                print(f"[DocxImageExtractor] PIL 렌더링 모드")
                table_images = self.table_converter.extract_tables_as_images(docx_path, document_name)
            except Exception as e:
                print(f"표 이미지 추출 중 오류 발생: {e}")
                return []

        # 이미지 정보를 기존 형식에 맞게 변환
        formatted_images = []
        for table_info in table_images:
            formatted_info = {
                "image_index": len(formatted_images),
                "file_name": table_info["file_name"],
                "file_path": table_info["file_path"],
                "content_type": table_info["content_type"],
                "paragraph_text": f"표 {table_info['table_index']} ({table_info['rows']}x{table_info['cols']})",
                "base64_data": table_info["base64_data"],
                "previous_text": table_info.get("previous_text", ""),
                "next_text": table_info.get("next_text", ""),
                "context": table_info.get("context", "표"),
                "table_location": table_info.get("table_location", "문서 내 표"),
                "table_data": table_info.get("table_data", []),
                "is_table": True  # 표 이미지임을 표시
            }
            formatted_images.append(formatted_info)

        return formatted_images

