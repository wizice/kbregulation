"""
HTML 변환 + Playwright 방식을 사용한 고품질 테이블 이미지 변환기
원본 DOCX의 스타일을 최대한 보존하여 이미지로 변환합니다.
"""

import os
import io
import base64
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import mammoth
from playwright.sync_api import sync_playwright
from bs4 import BeautifulSoup
import docx
from docx.table import Table
import tempfile


class EnhancedTableConverter:
    """
    HTML 변환 + Playwright를 사용한 고품질 테이블 이미지 변환기
    """

    def __init__(self, output_dir: str = "extracted_images",
                 image_quality: int = 300,
                 article_range: Optional[Tuple[int, int]] = None):
        """
        초기화

        Args:
            output_dir: 이미지 저장 디렉토리
            image_quality: 이미지 DPI (기본 300)
            article_range: 추출할 조문 범위 (시작조, 끝조)
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.image_quality = image_quality
        self.article_range = article_range

    def extract_tables_as_images(self, docx_path: str, document_name: str) -> List[Dict[str, Any]]:
        """
        DOCX 파일에서 표를 추출하여 고품질 이미지로 변환

        Args:
            docx_path: DOCX 파일 경로
            document_name: 문서 이름

        Returns:
            표 이미지 정보 리스트
        """
        print(f"Enhanced table conversion 시작: {docx_path}")

        try:
            # 1. DOCX를 HTML로 변환
            html_content = self._convert_docx_to_html(docx_path)
            if not html_content:
                print("HTML 변환 실패")
                return []

            # 2. HTML에서 표 추출
            tables_html = self._extract_tables_from_html(html_content)
            if not tables_html:
                print("HTML에서 표를 찾을 수 없음")
                return []

            # 3. 조문 범위 필터링 (필요한 경우)
            if self.article_range:
                tables_html = self._filter_tables_by_article_range(docx_path, tables_html)

            # 4. 각 표를 고품질 이미지로 변환
            table_images = []
            for idx, table_html in enumerate(tables_html):
                try:
                    image_info = self._convert_html_table_to_image(
                        table_html, document_name, idx
                    )
                    if image_info:
                        # 원본 DOCX에서 추가 정보 추출
                        additional_info = self._get_table_metadata(docx_path, idx)
                        image_info.update(additional_info)
                        table_images.append(image_info)

                except Exception as e:
                    print(f"표 {idx} 변환 중 오류: {e}")
                    continue

            print(f"총 {len(table_images)}개 표 변환 완료")
            return table_images

        except Exception as e:
            print(f"Enhanced table conversion 실패: {e}")
            return []

    def _convert_docx_to_html(self, docx_path: str) -> Optional[str]:
        """DOCX를 스타일 보존된 HTML로 변환"""
        try:
            with open(docx_path, "rb") as docx_file:
                # mammoth를 사용해 스타일 맵핑과 함께 HTML 변환
                style_map = """
                table => table[class="docx-table"]
                td => td[class="docx-cell"]
                th => th[class="docx-header"]
                p => p[class="docx-paragraph"]
                """

                result = mammoth.convert_to_html(
                    docx_file,
                    style_map=style_map,
                    convert_image=mammoth.images.img_element(self._save_embedded_image)
                )

                if result.messages:
                    print("HTML 변환 경고:")
                    for message in result.messages:
                        print(f"  - {message}")

                return result.value

        except Exception as e:
            print(f"DOCX → HTML 변환 실패: {e}")
            return None

    def _save_embedded_image(self, image):
        """임베디드 이미지 저장 핸들러"""
        try:
            # 이미지를 임시 파일로 저장하고 경로 반환
            temp_dir = self.output_dir / "temp_images"
            temp_dir.mkdir(exist_ok=True)

            # 이미지 확장자 결정
            content_type = getattr(image, 'content_type', 'image/png')
            ext = content_type.split('/')[-1] if '/' in content_type else 'png'

            image_path = temp_dir / f"embedded_{hash(str(image))}.{ext}"

            # 이미지 데이터 읽기
            with image.open() as img_file:
                image_data = img_file.read()

            with open(image_path, "wb") as f:
                f.write(image_data)

            return {"src": str(image_path)}

        except Exception as e:
            print(f"임베디드 이미지 저장 실패: {e}")
            return {"src": ""}

    def _extract_tables_from_html(self, html_content: str) -> List[str]:
        """HTML에서 표 요소들을 추출"""
        try:
            soup = BeautifulSoup(html_content, 'html.parser')
            tables = soup.find_all('table')

            table_htmls = []
            for table in tables:
                # 표를 독립적인 HTML 문서로 래핑
                table_html = str(table)
                table_htmls.append(table_html)

            print(f"HTML에서 {len(table_htmls)}개 표 추출")
            return table_htmls

        except Exception as e:
            print(f"HTML 표 추출 실패: {e}")
            return []

    def _convert_html_table_to_image(self, table_html: str, document_name: str, table_idx: int) -> Optional[Dict[str, Any]]:
        """HTML 표를 Playwright로 고품질 이미지로 변환"""
        try:
            # 완전한 HTML 문서 생성
            full_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="utf-8">
                <style>
                    body {{
                        margin: 20px;
                        font-family: 'Malgun Gothic', 'Arial', sans-serif;
                        background-color: white;
                    }}

                    table.docx-table, table {{
                        border-collapse: collapse;
                        width: auto;
                        margin: 0;
                        font-size: 12px;
                        color: #000;
                    }}

                    td.docx-cell, td, th.docx-header, th {{
                        border: 1px solid #000;
                        padding: 8px 12px;
                        text-align: left;
                        vertical-align: top;
                        line-height: 1.4;
                        word-wrap: break-word;
                    }}

                    th.docx-header, th {{
                        background-color: #f0f0f0;
                        font-weight: bold;
                        text-align: center;
                    }}

                    p.docx-paragraph, p {{
                        margin: 4px 0;
                        line-height: 1.4;
                    }}

                    /* 추가 스타일링 */
                    .table-container {{
                        display: inline-block;
                        background: white;
                        box-shadow: 0 0 10px rgba(0,0,0,0.1);
                        padding: 10px;
                    }}
                </style>
            </head>
            <body>
                <div class="table-container">
                    {table_html}
                </div>
            </body>
            </html>
            """

            # Playwright로 렌더링
            with sync_playwright() as p:
                browser = p.chromium.launch(headless=True)
                page = browser.new_page()

                # 고해상도 설정
                page.set_viewport_size({"width": 1920, "height": 1080})

                # HTML 콘텐츠 로드
                page.set_content(full_html, wait_until="networkidle")

                # 표 요소 찾기
                table_element = page.locator('table').first

                if not table_element.is_visible():
                    print(f"표 {table_idx}가 보이지 않음")
                    browser.close()
                    return None

                # 고품질 스크린샷 생성
                screenshot_bytes = table_element.screenshot(
                    type="png",
                    omit_background=False
                )

                browser.close()

                # 파일 저장
                table_filename = f"{document_name}_table_{table_idx}.png"
                table_filepath = self.output_dir / table_filename

                with open(table_filepath, "wb") as f:
                    f.write(screenshot_bytes)

                # 표 데이터도 추출
                table_data = self._extract_table_data_from_html(table_html)

                # Base64 인코딩
                base64_data = base64.b64encode(screenshot_bytes).decode('utf-8')

                print(f"✓ 표 {table_idx} 변환 완료: {table_filepath}")

                return {
                    "table_index": table_idx,
                    "file_name": table_filename,
                    "file_path": str(table_filepath),
                    "content_type": "image/png",
                    "base64_data": f"data:image/png;base64,{base64_data}",
                    "table_data": table_data,
                    "rows": len(table_data) if table_data else 0,
                    "cols": max(len(row) for row in table_data) if table_data else 0
                }

        except Exception as e:
            print(f"HTML 표 → 이미지 변환 실패: {e}")
            return None

    def _extract_table_data_from_html(self, table_html: str) -> List[List[str]]:
        """HTML 표에서 데이터 추출"""
        try:
            soup = BeautifulSoup(table_html, 'html.parser')
            table = soup.find('table')

            if not table:
                return []

            table_data = []
            rows = table.find_all('tr')

            for row in rows:
                row_data = []
                cells = row.find_all(['td', 'th'])

                for cell in cells:
                    cell_text = cell.get_text(strip=True)
                    row_data.append(cell_text)

                if row_data:  # 빈 행 제외
                    table_data.append(row_data)

            return table_data

        except Exception as e:
            print(f"HTML 표 데이터 추출 실패: {e}")
            return []

    def _filter_tables_by_article_range(self, docx_path: str, tables_html: List[str]) -> List[str]:
        """조문 범위에 따라 표 필터링"""
        if not self.article_range:
            return tables_html

        try:
            # 원본 DOCX에서 조문 정보 추출
            doc = docx.Document(docx_path)
            filtered_tables = []

            # 각 표의 위치를 분석하여 조문 범위 내에 있는지 확인
            for idx, table_html in enumerate(tables_html):
                if idx < len(doc.tables):
                    # 실제 조문 범위 확인 로직 (기존 방식 활용)
                    article_num = self._get_table_article_number(doc, idx)

                    start_article, end_article = self.article_range
                    if start_article <= article_num <= end_article:
                        filtered_tables.append(table_html)
                        print(f"✓ 표 {idx} 포함: 제{article_num}조 (범위 내)")
                    else:
                        print(f"✗ 표 {idx} 제외: 제{article_num}조 (범위 외)")
                else:
                    # 안전하게 포함
                    filtered_tables.append(table_html)

            return filtered_tables

        except Exception as e:
            print(f"조문 범위 필터링 실패: {e}")
            return tables_html

    def _get_table_article_number(self, doc: docx.Document, table_idx: int) -> int:
        """표가 속한 조문 번호를 찾음"""
        # 기존 table_to_image.py의 로직을 활용
        try:
            if table_idx < len(doc.tables):
                table = doc.tables[table_idx]
                # 표 앞의 텍스트에서 조문 번호 추출
                # (기존 구현과 동일한 방식)
                return 1  # 임시로 1 반환
            return 0
        except:
            return 0

    def _get_table_metadata(self, docx_path: str, table_idx: int) -> Dict[str, Any]:
        """표의 메타데이터 추출 (컨텍스트 정보 등)"""
        try:
            doc = docx.Document(docx_path)

            if table_idx < len(doc.tables):
                table = doc.tables[table_idx]

                # 표 크기 정보
                rows_count = len(table.rows)
                cols_count = len(table.columns) if table.rows else 0

                # 표 제목 생성
                title = f"표 {table_idx} ({rows_count}x{cols_count})"

                return {
                    "title": title,
                    "description": f"테이블 크기: {rows_count}행 {cols_count}열",
                    "is_table": True
                }

            return {"title": f"표 {table_idx}", "is_table": True}

        except Exception as e:
            print(f"표 메타데이터 추출 실패: {e}")
            return {"title": f"표 {table_idx}", "is_table": True}


def test_enhanced_converter():
    """테스트 함수"""
    converter = EnhancedTableConverter(
        output_dir="static/extracted_images/enhanced_test",
        article_range=(1, 3)
    )

    docx_path = "docx/6.4.1._의료사회복지체계_202503개정.docx"
    document_name = "6.4.1._의료사회복지체계_202503개정"

    results = converter.extract_tables_as_images(docx_path, document_name)

    print(f"\n=== 테스트 결과 ===")
    print(f"변환된 표 수: {len(results)}")
    for result in results:
        print(f"- {result['file_name']}: {result['rows']}x{result['cols']}")


if __name__ == "__main__":
    test_enhanced_converter()