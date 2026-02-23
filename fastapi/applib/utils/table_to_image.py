"""
표를 이미지로 변환하는 유틸리티 모듈
"""

import os
import io
import re
import base64
from typing import List, Dict, Any, Optional, Tuple
from PIL import Image, ImageDraw, ImageFont
import docx
from docx.table import Table
from docx.text.paragraph import Paragraph


class TableToImageConverter:
    """
    DOCX 표를 이미지로 변환하는 클래스
    """

    def __init__(self, output_dir: str = "extracted_images",
                 font_size: int = 12,
                 cell_padding: int = 10,
                 border_width: int = 1,
                 article_range: Optional[Tuple[int, int]] = None):
        """
        초기화

        Args:
            output_dir: 이미지를 저장할 디렉토리 경로
            font_size: 폰트 크기
            cell_padding: 셀 내부 패딩
            border_width: 테두리 두께
            article_range: 추출할 조문 범위 (시작조, 끝조) 예: (1, 3)
        """
        self.output_dir = output_dir
        self.font_size = font_size
        self.cell_padding = cell_padding
        self.border_width = border_width
        self.article_range = article_range

        # 출력 디렉토리가 없으면 생성
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        # 기본 폰트 설정
        self.font = self._get_font()

    def _download_korean_font(self, font_path: str) -> bool:
        """
        한글 폰트를 다운로드합니다.

        Args:
            font_path: 다운로드할 폰트 경로

        Returns:
            성공 여부
        """
        try:
            import urllib.request

            print(f"한글 폰트 다운로드 시도: {font_path}")

            # Noto Sans KR 폰트 다운로드
            font_url = "https://github.com/google/fonts/raw/main/ofl/notosanskr/NotoSansKR%5Bwght%5D.ttf"

            urllib.request.urlretrieve(font_url, font_path)

            if os.path.exists(font_path) and os.path.getsize(font_path) > 1000000:  # 1MB 이상
                print(f"한글 폰트 다운로드 성공: {font_path}")
                return True
            else:
                print(f"한글 폰트 다운로드 실패: 파일 크기 이상")
                return False

        except Exception as e:
            print(f"한글 폰트 다운로드 오류: {e}")
            return False

    def _get_font(self):
        """
        사용할 폰트를 가져옵니다.
        한글 지원 폰트를 우선적으로 찾습니다.

        Returns:
            ImageFont 객체
        """
        try:
            # 한글 지원 폰트 경로 (우선순위순)
            font_paths = [
                # 다운로드한 한글 폰트 (최우선)
                "/tmp/NotoSansKR.ttf",
                # Ubuntu/Linux 한글 폰트
                "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
                "/usr/share/fonts/truetype/nanum/NanumBarunGothic.ttf",
                "/usr/share/fonts/truetype/nanum/NanumMyeongjo.ttf",
                "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
                "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
                # CentOS/RHEL
                "/usr/share/fonts/nanum/NanumGothic.ttf",
                "/usr/share/fonts/liberation/LiberationSans-Regular.ttf",
                # Windows 한글 폰트
                "C:/Windows/Fonts/malgun.ttf",
                "C:/Windows/Fonts/malgunbd.ttf",
                "C:/Windows/Fonts/gulim.ttf",
                "C:/Windows/Fonts/batang.ttf",
                # macOS 한글 폰트
                "/Library/Fonts/AppleGothic.ttf",
                "/System/Library/Fonts/AppleSDGothicNeo.ttc",
                "/System/Library/Fonts/Arial.ttf",
                # 기타 공통 경로
                "/usr/share/fonts/TTF/DejaVuSans.ttf",
                "/System/Library/Fonts/Helvetica.ttc"
            ]

            print(f"폰트 검색 시작... (폰트 크기: {self.font_size})")

            # 한글 폰트가 없으면 자동 다운로드 시도
            if not os.path.exists("/tmp/NotoSansKR.ttf"):
                print("한글 폰트 없음. 자동 다운로드 시도...")
                self._download_korean_font("/tmp/NotoSansKR.ttf")

            for font_path in font_paths:
                if os.path.exists(font_path):
                    try:
                        font = ImageFont.truetype(font_path, self.font_size)
                        print(f"폰트 로드 성공: {font_path}")
                        return font
                    except Exception as e:
                        print(f"폰트 로드 실패: {font_path} - {e}")
                        continue

            print("특정 폰트를 찾을 수 없음. 기본 폰트 사용")

        except Exception as e:
            print(f"폰트 검색 오류: {e}")

        # 기본 폰트 사용 (최신 PIL 버전용)
        try:
            # PIL 10.0.0+ 버전
            font = ImageFont.load_default(size=self.font_size)
            print(f"기본 폰트 사용 (크기: {self.font_size})")
            return font
        except TypeError:
            try:
                # 이전 PIL 버전
                font = ImageFont.load_default()
                print("기본 폰트 사용 (기본 크기)")
                return font
            except Exception as e:
                print(f"기본 폰트 로드 실패: {e}")
                return None

    def extract_tables_as_images(self, docx_path: str, document_name: str) -> List[Dict[str, Any]]:
        """
        DOCX 파일에서 표를 추출하여 이미지로 변환합니다.
        article_range가 설정된 경우 해당 조문 범위 내의 표만 추출합니다.
        메타데이터 표(제정일, 승인란)는 자동으로 제외됩니다.

        Args:
            docx_path: DOCX 파일 경로
            document_name: 문서 이름

        Returns:
            표 이미지 정보 리스트
        """
        doc = docx.Document(docx_path)
        table_images = []

        # 조문 범위가 설정된 경우 필터링
        if self.article_range:
            filtered_tables = self._filter_tables_by_article_range(doc)
        else:
            # 모든 표 포함 (단, 메타데이터 표는 제외)
            filtered_tables = []
            for table_idx, table in enumerate(doc.tables):
                if not self._is_metadata_table(table):
                    filtered_tables.append((table_idx, table))
                else:
                    print(f"  ✗ 표 {table_idx} 제외: 메타데이터 표")

        print(f"추출 대상 표 개수: {len(filtered_tables)}")

        for table_idx, table in filtered_tables:
            try:
                # 표를 이미지로 변환
                image_info = self._convert_table_to_image(table, document_name, table_idx)
                if image_info:
                    # 표 주변 텍스트 컨텍스트 추가
                    context_info = self._get_table_context(doc, table)
                    image_info.update(context_info)
                    table_images.append(image_info)
                    print(f"  ✓ 표 {table_idx} 변환 완료")

            except Exception as e:
                print(f"표 {table_idx} 변환 중 오류 발생: {e}")
                continue

        return table_images

    def _is_metadata_table(self, table: Table) -> bool:
        """
        표가 메타데이터 표(제정일, 승인란 등)인지 확인합니다.

        Args:
            table: DOCX 표 객체

        Returns:
            메타데이터 표이면 True
        """
        if not table.rows:
            return False

        # 첫 번째 행의 텍스트 확인
        first_row_text = ""
        for cell in table.rows[0].cells:
            first_row_text += cell.text.strip() + " "
        first_row_text = first_row_text.strip()

        # 메타데이터 표 패턴
        metadata_patterns = [
            "제정일",
            "최종개정일",
            "최종검토일",
            "소관부서",
            "승인",
            "승  인",
            "병원장",
            "(인)"
        ]

        for pattern in metadata_patterns:
            if pattern in first_row_text:
                return True

        # 마지막 행도 확인 (승인란이 마지막에 있을 수 있음)
        if len(table.rows) > 1:
            last_row_text = ""
            for cell in table.rows[-1].cells:
                last_row_text += cell.text.strip() + " "
            last_row_text = last_row_text.strip()

            for pattern in ["승인", "승  인", "병원장", "(인)"]:
                if pattern in last_row_text:
                    return True

        return False

    def _convert_table_to_image(self, table: Table, document_name: str, table_idx: int) -> Optional[Dict[str, Any]]:
        """
        단일 표를 이미지로 변환합니다.

        Args:
            table: DOCX 표 객체
            document_name: 문서 이름
            table_idx: 표 인덱스

        Returns:
            표 이미지 정보
        """
        try:
            # 표 데이터 추출
            table_data = self._extract_table_data(table)
            if not table_data:
                return None

            # 이미지 생성
            image = self._render_table_image(table_data)
            if not image:
                return None

            # 이미지 파일 저장
            table_filename = f"{document_name}_table_{table_idx}.png"
            table_filepath = os.path.join(self.output_dir, table_filename)
            image.save(table_filepath, "PNG")

            # Base64 인코딩
            buffer = io.BytesIO()
            image.save(buffer, format="PNG")
            base64_data = base64.b64encode(buffer.getvalue()).decode('utf-8')

            # 웹 경로로 변환
            web_path = table_filepath
            if '/www/static/' in table_filepath:
                web_path = '/static/' + table_filepath.split('/www/static/')[1]
            elif '/fastapi/static/' in table_filepath:
                web_path = '/static/' + table_filepath.split('/fastapi/static/')[1]
            elif '/applib/static/' in table_filepath:
                web_path = '/static/' + table_filepath.split('/applib/static/')[1]

            return {
                "table_index": table_idx,
                "file_name": table_filename,
                "file_path": web_path,  # 웹 경로
                "content_type": "image/png",
                "base64_data": f"data:image/png;base64,{base64_data}",
                "table_data": table_data,
                "rows": len(table_data),
                "cols": max(len(row) for row in table_data) if table_data else 0
            }

        except Exception as e:
            print(f"표 변환 오류: {e}")
            return None

    def _extract_table_data(self, table: Table) -> List[List[str]]:
        """
        표에서 텍스트 데이터를 추출합니다.

        Args:
            table: DOCX 표 객체

        Returns:
            2차원 배열로 된 표 데이터
        """
        table_data = []

        for row in table.rows:
            row_data = []
            for cell in row.cells:
                # 셀 텍스트 추출 및 정리
                cell_text = cell.text.strip()
                # 긴 텍스트는 줄바꿈 처리
                if len(cell_text) > 50:
                    cell_text = self._wrap_text(cell_text, 50)
                row_data.append(cell_text)
            table_data.append(row_data)

        return table_data

    def _wrap_text(self, text: str, width: int) -> str:
        """
        긴 텍스트를 지정된 너비로 줄바꿈합니다.

        Args:
            text: 원본 텍스트
            width: 줄바꿈할 문자 수

        Returns:
            줄바꿈된 텍스트
        """
        if len(text) <= width:
            return text

        lines = []
        words = text.split()
        current_line = ""

        for word in words:
            if len(current_line + " " + word) <= width:
                current_line += " " + word if current_line else word
            else:
                if current_line:
                    lines.append(current_line)
                    current_line = word
                else:
                    # 단어가 너무 긴 경우 강제 줄바꿈
                    lines.append(word[:width])
                    current_line = word[width:]

        if current_line:
            lines.append(current_line)

        return "\n".join(lines)

    def _render_table_image(self, table_data: List[List[str]]) -> Optional[Image.Image]:
        """
        표 데이터를 이미지로 렌더링합니다.

        Args:
            table_data: 2차원 배열로 된 표 데이터

        Returns:
            PIL Image 객체
        """
        if not table_data:
            return None

        try:
            # 테이블 크기 계산
            max_cols = max(len(row) for row in table_data)
            rows = len(table_data)

            # 셀 크기 계산
            cell_heights = []
            cell_widths = [0] * max_cols

            for row_data in table_data:
                max_height = 0
                for col_idx, cell_text in enumerate(row_data):
                    # 텍스트 크기 측정
                    lines = cell_text.split('\n') if cell_text else ['']
                    text_width = max(self._get_text_width(line) for line in lines)
                    text_height = len(lines) * (self.font_size + 2)

                    # 최소/최대 크기 설정
                    text_width = max(100, min(text_width + self.cell_padding * 2, 300))
                    text_height = max(30, text_height + self.cell_padding * 2)

                    cell_widths[col_idx] = max(cell_widths[col_idx], text_width)
                    max_height = max(max_height, text_height)

                cell_heights.append(max_height)

            # 전체 이미지 크기
            total_width = sum(cell_widths) + (max_cols + 1) * self.border_width
            total_height = sum(cell_heights) + (rows + 1) * self.border_width

            # 이미지 생성
            image = Image.new('RGB', (total_width, total_height), 'white')
            draw = ImageDraw.Draw(image)

            # 표 그리기
            y_pos = 0
            for row_idx, row_data in enumerate(table_data):
                x_pos = 0
                cell_height = cell_heights[row_idx]

                for col_idx, cell_text in enumerate(row_data):
                    if col_idx >= len(cell_widths):
                        break

                    cell_width = cell_widths[col_idx]

                    # 셀 배경 (헤더는 회색)
                    if row_idx == 0:
                        draw.rectangle([x_pos, y_pos, x_pos + cell_width + self.border_width,
                                      y_pos + cell_height + self.border_width],
                                     fill='#f0f0f0', outline='black', width=self.border_width)
                    else:
                        draw.rectangle([x_pos, y_pos, x_pos + cell_width + self.border_width,
                                      y_pos + cell_height + self.border_width],
                                     fill='white', outline='black', width=self.border_width)

                    # 텍스트 그리기
                    if cell_text and self.font:
                        self._draw_text_in_cell(draw, cell_text,
                                               x_pos + self.cell_padding,
                                               y_pos + self.cell_padding,
                                               cell_width - self.cell_padding * 2)

                    x_pos += cell_width + self.border_width

                y_pos += cell_height + self.border_width

            return image

        except Exception as e:
            print(f"표 렌더링 오류: {e}")
            return None

    def _get_text_width(self, text: str) -> int:
        """
        텍스트 너비를 계산합니다.

        Args:
            text: 텍스트

        Returns:
            픽셀 단위 너비
        """
        if not text or not self.font:
            return 0

        try:
            # PIL의 textbbox 사용 (최신 버전)
            bbox = self.font.getbbox(text)
            return bbox[2] - bbox[0]
        except:
            # 이전 버전 호환성
            try:
                return self.font.getsize(text)[0]
            except:
                # 대략적인 계산
                return len(text) * (self.font_size // 2)

    def _draw_text_in_cell(self, draw: ImageDraw.Draw, text: str, x: int, y: int, max_width: int):
        """
        셀 내에 텍스트를 그립니다.
        한글 텍스트 렌더링을 안전하게 처리합니다.

        Args:
            draw: ImageDraw 객체
            text: 그릴 텍스트
            x: X 좌표
            y: Y 좌표
            max_width: 최대 너비
        """
        if not text:
            return

        lines = text.split('\n')
        line_height = self.font_size + 2

        for i, line in enumerate(lines):
            if not line.strip():
                continue

            try:
                # 한글 텍스트 인코딩 확인 및 정제
                if isinstance(line, str):
                    # UTF-8 인코딩 확인
                    clean_line = line.encode('utf-8').decode('utf-8')
                else:
                    clean_line = str(line)

                # 폰트가 있는 경우 폰트로 렌더링
                if self.font:
                    draw.text((x, y + i * line_height), clean_line, fill='black', font=self.font)
                else:
                    # 기본 폰트로 렌더링 (폰트 없음)
                    draw.text((x, y + i * line_height), clean_line, fill='black')

            except UnicodeDecodeError as e:
                print(f"텍스트 인코딩 오류: {e}, 원본 텍스트: {repr(line)}")
                # 오류 시 대체 텍스트 사용
                safe_text = line.encode('utf-8', errors='replace').decode('utf-8')
                if self.font:
                    draw.text((x, y + i * line_height), safe_text, fill='black', font=self.font)
                else:
                    draw.text((x, y + i * line_height), safe_text, fill='black')

            except Exception as e:
                print(f"텍스트 렌더링 오류: {e}, 원본 텍스트: {repr(line)}")
                # 최종 대안: 안전한 텍스트로 렌더링
                safe_text = ''.join(c if ord(c) < 128 else '?' for c in line)
                try:
                    if self.font:
                        draw.text((x, y + i * line_height), safe_text, fill='black', font=self.font)
                    else:
                        draw.text((x, y + i * line_height), safe_text, fill='black')
                except:
                    # 최종 대안도 실패시 무시
                    print(f"텍스트 렌더링 완전 실패: {line}")

    def _get_table_context(self, doc, table: Table) -> Dict[str, str]:
        """
        표 주변의 텍스트 컨텍스트를 가져옵니다.

        Args:
            doc: DOCX 문서 객체
            table: 표 객체

        Returns:
            컨텍스트 정보
        """
        context = {
            "previous_text": "",
            "next_text": "",
            "context": "표",
            "table_location": f"문서 내 표"
        }

        try:
            # 문서의 모든 요소를 순회하며 표의 위치를 찾음
            elements = []
            for element in doc.element.body:
                if element.tag.endswith('}p'):  # 단락
                    para = Paragraph(element, doc)
                    if para.text.strip():
                        elements.append(('paragraph', para.text.strip()))
                elif element.tag.endswith('}tbl'):  # 표
                    elements.append(('table', element))

            # 현재 표의 인덱스 찾기
            table_index = -1
            for i, (element_type, element_data) in enumerate(elements):
                if element_type == 'table' and element_data == table._element:
                    table_index = i
                    break

            if table_index >= 0:
                # 이전 텍스트 (최대 3개 단락)
                prev_texts = []
                for i in range(max(0, table_index - 3), table_index):
                    if elements[i][0] == 'paragraph':
                        prev_texts.append(elements[i][1])
                context["previous_text"] = "\n".join(prev_texts[-3:])

                # 다음 텍스트 (최대 3개 단락)
                next_texts = []
                for i in range(table_index + 1, min(len(elements), table_index + 4)):
                    if elements[i][0] == 'paragraph':
                        next_texts.append(elements[i][1])
                context["next_text"] = "\n".join(next_texts[:3])

        except Exception as e:
            print(f"표 컨텍스트 추출 오류: {e}")

        return context

    def _filter_tables_by_article_range(self, doc) -> List[Tuple[int, Any]]:
        """
        지정된 조문 범위 내의 표만 필터링합니다.
        예: article_range=(1, 3)이면 제1조, 제2조, 제3조 범위 내 표만 추출

        문서 구조에 따라 다양한 패턴을 지원:
        - "제X조" 패턴
        - "(목적)", "(정의)", "(절차)" 등의 괄호 패턴 (제1조, 제2조, 제3조로 매핑)

        Args:
            doc: DOCX 문서 객체

        Returns:
            (table_index, table_object) 튜플 리스트
        """
        if not self.article_range:
            return [(i, table) for i, table in enumerate(doc.tables)]

        start_article, end_article = self.article_range
        filtered_tables = []

        try:
            print(f"조문 범위 필터링: 제{start_article}조~제{end_article}조")

            # 섹션 헤더를 조문 번호로 매핑
            section_to_article = {
                "(목적)": 1,
                "(정의)": 2,
                "(절차)": 3,
                "(운영)": 4,
                "(기타)": 5,
                "(부칙)": 6
            }

            # 문서의 모든 요소를 순회하며 조문과 표의 위치 파악
            elements = []
            current_article = 0  # 현재 조문 번호

            # 모든 요소 가져오기
            for element in doc.element.body:
                if element.tag.endswith('}p'):  # 단락
                    para = Paragraph(element, doc)
                    text = para.text.strip()

                    # 1. 먼저 일반적인 제X조 패턴 확인
                    article_match = re.search(r'제\s*(\d+)\s*조', text)
                    if article_match:
                        new_article = int(article_match.group(1))
                        if new_article != current_article:
                            current_article = new_article
                            print(f"  조문 발견: 제{current_article}조 - '{text[:50]}...'")

                    # 2. 괄호 패턴 확인 (목적, 정의, 절차 등)
                    else:
                        for section_text, article_num in section_to_article.items():
                            if section_text in text:
                                if article_num != current_article:
                                    current_article = article_num
                                    print(f"  섹션 발견: {section_text} → 제{current_article}조 - '{text[:50]}...'")
                                break

                    elements.append(('paragraph', current_article, text))

                elif element.tag.endswith('}tbl'):  # 표
                    elements.append(('table', current_article, element))

            # 표 필터링
            for i, (element_type, article_num, element_data) in enumerate(elements):
                if element_type == 'table':
                    print(f"  표 발견: 제{article_num}조 범위에 있음")

                    # 현재 조문이 지정된 범위 내에 있는지 확인
                    if start_article <= article_num <= end_article:
                        # 실제 table 객체 찾기
                        for doc_table_idx, doc_table in enumerate(doc.tables):
                            if doc_table._element == element_data:
                                filtered_tables.append((len(filtered_tables), doc_table))
                                print(f"    ✓ 표 추가: 인덱스 {len(filtered_tables)-1} (제{article_num}조)")
                                break
                    else:
                        print(f"    ✗ 표 제외: 제{article_num}조는 범위 외 (필요: 제{start_article}조~제{end_article}조)")

            print(f"최종 결과: 제{start_article}조~제{end_article}조 범위에서 {len(filtered_tables)}개 표 발견")

        except Exception as e:
            print(f"조문 범위 필터링 오류: {e}")
            import traceback
            traceback.print_exc()
            # 오류 시 모든 표 반환
            filtered_tables = [(i, table) for i, table in enumerate(doc.tables)]

        return filtered_tables