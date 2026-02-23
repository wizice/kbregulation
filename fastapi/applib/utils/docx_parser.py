"""
KB신용정보 사규 문서를 파싱하기 위한 모듈
"""
import re
from typing import Dict, Any, List, Tuple, Optional, TYPE_CHECKING

if TYPE_CHECKING:
    import docx
else:
    try:
        import docx
    except ImportError:
        docx = None


def extract_formatted_text_from_paragraph(paragraph) -> str:
    """
    단락에서 윗첨자/아랫첨자/이탤릭체를 포함한 포맷된 텍스트를 추출합니다.

    Args:
        paragraph: docx 단락 객체

    Returns:
        포맷된 텍스트 (윗첨자는 <sup>, 아랫첨자는 <sub>, 이탤릭체는 <i> 태그로 표현)
    """
    result = []

    try:
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue

            # 윗첨자/아랫첨자/이탤릭체 확인
            is_superscript = run.font.superscript
            is_subscript = run.font.subscript

            if is_superscript:
                result.append(f"<sup>{text}</sup>")
            elif is_subscript:
                result.append(f"<sub>{text}</sub>")
            elif run.font.italic:
                result.append(f"<i>{text}</i>")
            else:
                result.append(text)
    except Exception:
        # 오류 발생 시 기본 텍스트 반환
        return paragraph.text

    return ''.join(result)


def extract_rich_text_from_paragraph(paragraph, skip_prefix_len=0) -> str:
    """
    단락에서 굵게(bold), 밑줄(underline), 윗첨자, 아랫첨자, 이탤릭 서식을 포함한 리치 텍스트를 추출합니다.

    Args:
        paragraph: docx 단락 객체
        skip_prefix_len: 건너뛸 접두 문자 수 (번호 제거용)

    Returns:
        서식 태그(<b>, <u>, <sup>, <sub>, <i>)가 포함된 텍스트
    """
    chars_fmt = []

    try:
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue

            is_bold = bool(run.font.bold)
            is_underline = bool(run.font.underline)
            is_sup = bool(run.font.superscript)
            is_sub = bool(run.font.subscript)
            is_italic = bool(run.font.italic)
            # 색상 추출 (검정색 000000 제외)
            color_hex = None
            if run.font.color and run.font.color.rgb:
                rgb_str = str(run.font.color.rgb)
                if rgb_str != '000000':
                    color_hex = rgb_str

            for c in text:
                chars_fmt.append((c, is_bold, is_underline, is_sup, is_sub, is_italic, color_hex))
    except Exception:
        return paragraph.text[skip_prefix_len:]

    # 접두 문자 건너뛰기
    chars_fmt = chars_fmt[skip_prefix_len:]

    if not chars_fmt:
        return ''

    # HTML 문자열 조립
    result = []
    in_b = in_u = in_sup = in_sub = in_i = False
    in_color = None  # 현재 적용 중인 색상

    for c, bold, underline, sup, sub, italic, color in chars_fmt:
        # 닫기 (역순)
        if in_i and not italic:
            result.append('</i>'); in_i = False
        if in_sub and not sub:
            result.append('</sub>'); in_sub = False
        if in_sup and not sup:
            result.append('</sup>'); in_sup = False
        if in_u and not underline:
            result.append('</u>'); in_u = False
        if in_color and in_color != color:
            result.append('</span>'); in_color = None
        if in_b and not bold:
            result.append('</b>'); in_b = False

        # 열기
        if bold and not in_b:
            result.append('<b>'); in_b = True
        if color and in_color != color:
            result.append(f'<span style="color:#{color}">'); in_color = color
        if underline and not in_u:
            result.append('<u>'); in_u = True
        if sup and not in_sup:
            result.append('<sup>'); in_sup = True
        if sub and not in_sub:
            result.append('<sub>'); in_sub = True
        if italic and not in_i:
            result.append('<i>'); in_i = True

        result.append(c)

    # 남은 태그 닫기
    if in_i: result.append('</i>')
    if in_sub: result.append('</sub>')
    if in_sup: result.append('</sup>')
    if in_u: result.append('</u>')
    if in_color: result.append('</span>')
    if in_b: result.append('</b>')

    return ''.join(result)


def extract_metadata(doc: Any) -> Dict[str, str]:
    """
    KB신용정보 사규 문서에서 메타데이터(제정일, 개정일 이력, 시행일)를 추출합니다.

    Args:
        doc: docx.Document 객체

    Returns:
        메타데이터를 포함하는 사전
    """
    metadata = {
        "제정일": "",
        "개정일_이력": [],
        "시행일": "",
        "최종개정일": "",
        "소관부서": ""
    }

    # 문서 제목 추출 (첫 번째 단락, 윗첨자/아랫첨자 포함)
    doc_title = ""
    if doc.paragraphs and extract_formatted_text_from_paragraph(doc.paragraphs[0]).strip():
        doc_title = extract_formatted_text_from_paragraph(doc.paragraphs[0]).strip()
        metadata["문서제목"] = doc_title

    # 처음 10개 단락을 검사하여 제정/개정 이력 및 시행일 추출
    import re
    for i, paragraph in enumerate(doc.paragraphs[:10]):
        text = extract_formatted_text_from_paragraph(paragraph).strip()
        if not text:
            continue

        # 소관부서 패턴 추출 (예: "(소관부서 : 경영전략부)")
        if "소관부서" in text:
            dept_match = re.search(r'\(소관부서\s*:\s*([^)]+)\)', text)
            if dept_match:
                metadata["소관부서"] = dept_match.group(1).strip()

        # 시행일 패턴 추출 (예: "[시행 2018. 12. 1]")
        if "[시행" in text:
            match = re.search(r'\[시행\s+([\d.\s]+)\s*\]', text)
            if match:
                metadata["시행일"] = match.group(1).strip()

        # 제정/개정 이력 패턴 추출 (예: "(1999. 10.  9. 제정)")
        if "제정" in text or "개정" in text:
            # 괄호로 묶인 제정/개정 이력 추출
            revisions = re.findall(r'\(([\d.\s]+)\s*(제정|개정|일부개정|전부개정)\)', text)
            for date, rev_type in revisions:
                date_cleaned = date.strip()

                # 제정일 추출
                if rev_type == "제정" and not metadata["제정일"]:
                    metadata["제정일"] = date_cleaned

                # 개정일 이력에 추가
                metadata["개정일_이력"].append({
                    "날짜": date_cleaned,
                    "유형": rev_type
                })

                # 최종 개정일 업데이트 (가장 마지막 개정일)
                if "개정" in rev_type:
                    metadata["최종개정일"] = date_cleaned

    # 개정일 이력을 문자열로 변환 (JSON 저장 용이성을 위해)
    if metadata["개정일_이력"]:
        revision_strings = [f"{item['날짜']} ({item['유형']})" for item in metadata["개정일_이력"]]
        metadata["개정일_이력_텍스트"] = ", ".join(revision_strings)

    return metadata


def extract_content_structure(doc: Any) -> List[Dict[str, Any]]:
    """
    문서에서 조문별 내용과 레벨 정보를 추출합니다.
    
    Args:
        doc: docx.Document 객체
        
    Returns:
        조문별 내용과 레벨 정보를 포함하는 리스트
    """
    content_structure = []
    current_section = None
    
    # 메타데이터 테이블 이후의 내용만 처리
    found_content_start = False
    
    for paragraph in doc.paragraphs:
        # 자동번호매기기 형식 확인 및 추출 (윗첨자/아랫첨자 포함)
        text = extract_formatted_text_from_paragraph(paragraph).strip()
        if not text:
            continue

        # 번호 매기기 패턴 확인
        numbering_prefix = ""
        numbering_level = detect_numbering_level(text)

        # 자동번호매기기로 생성된 번호가 없고, XML에서 번호매기기 속성이 있는 경우
        if numbering_level is None:
            # XML에서 번호매기기 확인
            numbering_prefix = extract_numbering_from_xml(paragraph)
            if numbering_prefix:
                # 접두사 추가하여 텍스트 재구성
                text = numbering_prefix + text
                # 내용 레벨 감지
                numbering_level = detect_numbering_level(text)
        
        # 메타데이터 테이블 이후의 내용만 처리
        if not found_content_start and ("목적" in text or "정의" in text) and re.search(r'^\d+\.|\(목적\)|\(정의\)', text):
            found_content_start = True
        
        if not found_content_start:
            continue
            
        # 내규의 제·개정 이력 섹션에 도달하면 처리 중단
        if "내규의 제·개정 이력" in text:
            break
            
        # 조문 레벨 및 내용 추출
        level, content = parse_paragraph(text)

        # 컨텍스트 기반 레벨 조정
        if level is not None and content:
            level = adjust_level_by_context(content, level, content_structure)
            section = {
                "level": level,
                "content": content,
                "subsections": []
            }
            
            # 레벨에 따라 계층 구조에 추가
            if level == 1:
                current_section = section
                content_structure.append(current_section)
            elif level == 2 and current_section:
                current_section["subsections"].append(section)
    
    return content_structure


def extract_numbering_from_xml(paragraph) -> str:
    """
    단락의 XML에서 자동번호매기기 속성을 추출합니다.
    
    Args:
        paragraph: docx 단락 객체
        
    Returns:
        추출된 번호 텍스트 (없으면 빈 문자열)
    """
    try:
        if hasattr(paragraph, '_element') and paragraph._element is not None:
            # numPr 요소 확인 (번호 매기기 속성)
            num_pr = paragraph._element.xpath('.//w:numPr')
            if num_pr:
                # numId와 ilvl 값을 확인
                num_id = paragraph._element.xpath('.//w:numId/@w:val')
                level_id = paragraph._element.xpath('.//w:ilvl/@w:val')
                
                if num_id and level_id:
                    # 레벨에 따라 적절한 번호 형식 생성
                    level = int(level_id[0])
                    # 더 세밀한 레벨 지원을 위해 확장 (최대 레벨 8까지)
                    if level == 0:
                        return "1. "       # 1수준 번호
                    elif level == 1:
                        return "1) "       # 2수준 번호
                    elif level == 2:
                        return "(가) "     # 3수준 번호
                    elif level == 3:
                        return "ⅰ) "      # 4수준 번호
                    elif level == 4:
                        return "① "       # 5수준 번호 (원형 번호)
                    elif level == 5:
                        return "㉮ "       # 6수준 번호
                    elif level == 6:
                        return "a) "       # 7수준 번호
                    elif level == 7:
                        return "(a) "      # 8수준 번호
                    else:
                        return "- "        # 기타 수준
    except Exception:
        pass
    
    return ""


def adjust_level_by_context(content: str, current_level: int, existing_structure: List[Dict[str, Any]]) -> int:
    """
    컨텍스트를 기반으로 레벨을 조정합니다.
    잘못된 레벨 할당을 수정하는 함수입니다.

    Args:
        content: 현재 섹션의 내용
        current_level: 현재 할당된 레벨
        existing_structure: 기존 구조 리스트

    Returns:
        조정된 레벨
    """
    # 특정 패턴의 텍스트가 잘못된 레벨로 할당된 경우 수정
    problematic_texts = [
        "정도관리 결과 허용범위를 초과하거나 이상치 발견 시",
        "영상 이상으로 환자에게 추가검사가 필요하거나"
    ]

    # 문제가 있는 텍스트인지 확인
    is_problematic = any(pattern in content for pattern in problematic_texts)

    if is_problematic and current_level == 1:
        # 바로 이전 섹션이 레벨 4인 경우, 이 섹션은 레벨 5여야 함
        if existing_structure:
            last_item = existing_structure[-1]
            if hasattr(last_item, 'get') and last_item.get('level') == 4:
                return 5
        # 기본적으로 레벨 5로 조정
        return 5

    return current_level


def detect_numbering_level(text: str) -> Optional[int]:
    """
    텍스트에서 번호매기기 패턴을 감지하고 레벨을 결정합니다.
    
    Args:
        text: 텍스트 문자열
        
    Returns:
        감지된 레벨 또는 None (감지 실패시)
    """
    # 1. 수준 (예: "1. (목적)" 또는 "제1조 (목적)")
    if re.match(r'^\d+\.\s*\([^)]+\)|^제\d+조\s*\([^)]+\)', text):
        return 1
        
    # 1-1. 수준 (예: "1-1. 정확한 환자 확인: ")
    if re.match(r'^\d+-\d+\.\s*', text):
        return 2
        
    # a. 수준 또는 추가 내용 (예: "a. 약물 투여 전")
    if re.match(r'^[a-z]\.\s*', text):
        return 3
        
    # 번호 매기기 형식 1), 2) 등 - 스킵
    # if re.match(r'^\s*\d+\)\s*', text):
    #     return 3
        
    # 번호 매기기 형식 ①, ②, ③ 등
    if re.match(r'^\s*[①②③④⑤⑥⑦⑧⑨⑩]\s*', text):
        return 3
    
    return None


def extract_numbering_text(paragraph) -> str:
    """
    단락에서 자동번호매기기로 생성된 번호를 추출합니다.
    
    Args:
        paragraph: docx 단락 객체
        
    Returns:
        추출된 번호 텍스트 (없으면 빈 문자열)
    """
    # 기본 텍스트에서 번호 형식을 추출하는 정규식 패턴
    numbering_patterns = [
        (r'^(\d+\.\s*)', 1),               # 1. 형식
        (r'^(\d+-\d+\.\s*)', 2),           # 1-1. 형식
        (r'^([a-z]\.\s*)', 3),             # a. 형식
        # (r'^(\d+\)\s*)', 3),               # 1) 형식 스킵
        (r'^([①②③④⑤⑥⑦⑧⑨⑩]\s*)', 3),      # ① 형식
    ]
    
    # 자동번호매기기 확인 시도
    try:
        # 텍스트에서 번호 패턴 확인
        text = paragraph.text.strip()
        for pattern, level in numbering_patterns:
            match = re.match(pattern, text)
            if match:
                # 이미 텍스트에 번호가 포함되어 있으므로 빈 문자열 반환
                return ""
        
        # XML에서 numbering 속성 확인
        if hasattr(paragraph, '_element') and paragraph._element is not None:
            # numPr 요소 확인 (번호 매기기 속성)
            num_pr = paragraph._element.xpath('.//w:numPr')
            if num_pr:
                # numId와 ilvl 값을 확인
                num_id = paragraph._element.xpath('.//w:numId/@w:val')
                level_id = paragraph._element.xpath('.//w:ilvl/@w:val')
                
                if num_id and level_id:
                    # 레벨에 따라 적절한 번호 형식 생성
                    level = int(level_id[0])
                    # 더 세밀한 레벨 지원을 위해 확장 (최대 레벨 8까지)
                    if level == 0:
                        return "1. "       # 1수준 번호
                    elif level == 1:
                        return "1) "       # 2수준 번호
                    elif level == 2:
                        return "(가) "     # 3수준 번호
                    elif level == 3:
                        return "ⅰ) "      # 4수준 번호
                    elif level == 4:
                        return "① "       # 5수준 번호 (원형 번호)
                    elif level == 5:
                        return "㉮ "       # 6수준 번호
                    elif level == 6:
                        return "a) "       # 7수준 번호
                    elif level == 7:
                        return "(a) "      # 8수준 번호
                    else:
                        return "- "        # 기타 수준
            
            # 또는 자동번호매기기에서 사용하는 다른 XML 태그 확인
            list_pr = paragraph._element.xpath('.//w:ilfo')
            if list_pr:
                # 목록 형식이 있으면 번호 추출 시도
                return "• "  # 기본 글머리 기호
    except Exception as e:
        # 오류가 발생하면 무시하고 빈 문자열 반환
        pass
    
    return ""


def parse_paragraph(text: str) -> Tuple[int, str]:
    """
    단락에서 레벨과 내용을 추출합니다.
    
    Args:
        text: 단락 텍스트
        
    Returns:
        (레벨, 내용) 튜플
    """
    # 공백 처리
    text = text.lstrip()
    
    # 1. 수준 (예: "1. (목적)" 또는 "제1조 (목적)")
    level1_match = re.match(r'^\d+\.\s*\(([^)]+)\)', text)
    level1_alt_match = re.match(r'^제\d+조\s*\(([^)]+)\)', text)
    if level1_match or level1_alt_match:
        return 1, text
        
    # 1-1. 수준 (예: "1-1. 정확한 환자 확인: ")
    level2_match = re.match(r'^\d+-\d+\.\s*(.*)', text)
    if level2_match:
        return 2, text
        
    # a. 수준 또는 추가 내용 (예: "a. 약물 투여 전")
    level3_match = re.match(r'^[a-z]\.\s*(.*)', text)
    if level3_match:
        return 3, text
        
    # 번호 매기기 형식 1), 2) 등 - 들여쓰기나 탭이 있을 수 있으므로 시작 부분의 공백 허용 - 스킵
    # level3_numbered_match = re.match(r'^\s*\d+\)\s*(.*)', text)
    # if level3_numbered_match:
    #     return 3, text
        
    # 번호 매기기 형식 ①, ②, ③ 등 - 들여쓰기나 탭이 있을 수 있으므로 시작 부분의 공백 허용
    level3_circle_match = re.match(r'^\s*[①②③④⑤⑥⑦⑧⑨⑩]\s*(.*)', text)
    if level3_circle_match:
        return 3, text
    
    # "제1조 (목적)" 형식이 한글로 된 경우
    if text.startswith("제") and "조" in text[:5] and "(" in text:
        return 1, text
        
    # 들여쓰기 수준 확인
    indent_level = 0
    for char in text:
        if char in [' ', '\t']:
            indent_level += 1
        else:
            break
            
    if indent_level >= 4:
        # 들여쓰기가 있는 텍스트는 보통 하위 항목
        if indent_level >= 8:
            return 3, text  # 더 깊은 들여쓰기는 레벨 3으로 간주
        else:
            return 2, text  # 중간 정도 들여쓰기는 레벨 2로 간주
        
    # 일반 텍스트 - 앞에 공백이 많으면 하위 레벨로 간주
    return None, text

