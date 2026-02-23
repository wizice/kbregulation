"""
DOCX 표를 원본 서식 보존하여 고품질 이미지로 변환하는 모듈

python-docx로 OoXML에서 서식 정보를 직접 읽고,
PIL/Pillow로 고품질 이미지 렌더링합니다.
"""

import os
import io
import re
import base64
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass, field
from PIL import Image, ImageDraw, ImageFont

import docx
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn


# ── 데이터 클래스 ──────────────────────────────────────────

@dataclass
class BorderSide:
    style: str = 'solid'
    width: float = 1.0
    color: str = '#000000'


@dataclass
class CellBorders:
    top: BorderSide = field(default_factory=BorderSide)
    bottom: BorderSide = field(default_factory=BorderSide)
    left: BorderSide = field(default_factory=BorderSide)
    right: BorderSide = field(default_factory=BorderSide)


@dataclass
class CellProps:
    grid_span: int = 1
    vmerge: Optional[str] = None   # 'restart', 'continue', None
    rowspan: int = 1
    width_twips: int = 0
    borders: CellBorders = field(default_factory=CellBorders)
    shading_fill: Optional[str] = None
    vertical_align: str = 'middle'
    margins: Dict[str, int] = field(default_factory=dict)


@dataclass
class RunFormat:
    font_name: str = ''
    font_name_east_asia: str = ''
    font_size_pt: float = 10.0
    bold: bool = False
    italic: bool = False
    underline: bool = False
    color: str = '#000000'


@dataclass
class ParaFormat:
    alignment: str = 'center'
    line_spacing_pt: Optional[float] = None
    space_before_pt: float = 0
    space_after_pt: float = 0


@dataclass
class CellInfo:
    row: int = 0
    col: int = 0
    grid_span: int = 1
    rowspan: int = 1
    is_merge_continuation: bool = False
    props: CellProps = field(default_factory=CellProps)
    paragraphs: List[Tuple[ParaFormat, List[Tuple[str, RunFormat]]]] = field(default_factory=list)
    text: str = ''


# ── 상수 ─────────────────────────────────────────────────

TWIP_TO_PX = 96.0 / 1440.0  # 1 twip → px (at 96dpi)

BORDER_STYLE_MAP = {
    'single': 'solid',
    'double': 'double',
    'dotted': 'dotted',
    'dashed': 'dashed',
    'dashSmallGap': 'dashed',
    'dotDash': 'dashed',
    'dotDotDash': 'dotted',
    'triple': 'double',
    'thick': 'solid',
    'thickThinSmallGap': 'double',
    'thinThickSmallGap': 'double',
    'nil': 'none',
    'none': 'none',
}

FONT_CSS_MAP = {
    '바탕체': "'Batangche', 'Batang', serif",
    '바탕': "'Batang', serif",
    '맑은 고딕': "'Malgun Gothic', sans-serif",
    '굴림': "'Gulim', sans-serif",
    '돋움': "'Dotum', sans-serif",
    '궁서': "'Gungsuh', serif",
    'HY신명조': "'Batang', serif",
}


# ── 메인 클래스 ──────────────────────────────────────────

class DocxTableRenderer:
    """
    DOCX 표를 원본 서식 그대로 이미지로 변환하는 클래스.
    TableToImageConverter와 동일한 인터페이스를 제공합니다.
    """

    def __init__(self, output_dir: str = "extracted_images",
                 article_range: Optional[Tuple[int, int]] = None,
                 dpi_scale: int = 2,
                 **kwargs):
        self.output_dir = output_dir
        self.article_range = article_range
        self.dpi_scale = dpi_scale
        self._font_cache = {}
        self._base_font_path = self._find_korean_font()

        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

    # ── 공개 API ─────────────────────────────────────────

    def extract_tables_as_images(self, docx_path: str, document_name: str) -> List[Dict[str, Any]]:
        """
        DOCX 파일에서 표를 추출하여 원본 서식 보존 이미지로 변환합니다.
        """
        doc = docx.Document(docx_path)
        doc_defaults = self._parse_document_defaults(doc)
        table_images = []

        # 조문 범위 필터링
        if self.article_range:
            filtered_tables = self._filter_tables_by_article_range(doc)
        else:
            filtered_tables = [
                (idx, tbl) for idx, tbl in enumerate(doc.tables)
                if not self._is_metadata_table(tbl)
            ]

        print(f"[DocxTableRenderer] 추출 대상 표 개수: {len(filtered_tables)}")

        for table_idx, table in filtered_tables:
            try:
                image_info = self._convert_table(table, doc, doc_defaults, document_name, table_idx)
                if image_info:
                    context = self._get_table_context(doc, table)
                    image_info.update(context)
                    table_images.append(image_info)
                    print(f"  ✓ 표 {table_idx} 변환 완료")
            except Exception as e:
                print(f"  ✗ 표 {table_idx} 변환 실패: {e}")
                continue

        return table_images

    # ── 테이블 변환 핵심 ─────────────────────────────────

    def _convert_table(self, table: Table, doc, doc_defaults: dict,
                       document_name: str, table_idx: int) -> Optional[Dict[str, Any]]:
        """단일 표를 이미지로 변환"""
        # 1) 그리드 너비 파싱
        grid_widths = self._parse_table_grid(table)

        # 2) 테이블 기본 테두리/마진
        tbl_borders = self._parse_table_borders(table)
        tbl_margins = self._parse_table_cell_margins(table)

        # 3) 셀 매트릭스 구축 (병합 처리 포함)
        cell_matrix = self._build_cell_matrix(table, grid_widths, doc_defaults, tbl_borders, tbl_margins)
        if not cell_matrix:
            return None

        # 4) 행 높이 파싱
        row_heights = self._parse_row_heights(table)

        # 5) 텍스트 데이터 추출 (기존 호환)
        table_data = self._extract_plain_table_data(table)

        # 6) PIL로 렌더링
        screenshot_bytes = self._render_table_with_pil(cell_matrix, grid_widths, row_heights, doc_defaults)
        if not screenshot_bytes:
            return None

        # 8) 파일 저장
        table_filename = f"{document_name}_table_{table_idx}.png"
        table_filepath = os.path.join(self.output_dir, table_filename)
        with open(table_filepath, "wb") as f:
            f.write(screenshot_bytes)

        # 9) Base64 인코딩
        base64_data = base64.b64encode(screenshot_bytes).decode('utf-8')

        # 10) 웹 경로 변환
        web_path = table_filepath
        for marker in ('/www/static/', '/fastapi/static/', '/applib/static/'):
            if marker in table_filepath:
                web_path = '/static/' + table_filepath.split(marker)[1]
                break

        return {
            "table_index": table_idx,
            "file_name": table_filename,
            "file_path": web_path,
            "content_type": "image/png",
            "base64_data": f"data:image/png;base64,{base64_data}",
            "table_data": table_data,
            "rows": len(table_data),
            "cols": max((len(r) for r in table_data), default=0),
        }

    # ── OoXML 파싱 ───────────────────────────────────────

    def _parse_document_defaults(self, doc) -> dict:
        """문서 기본 폰트/크기 추출"""
        defaults = {
            'ascii_font': 'Times New Roman',
            'east_asia_font': '',
            'font_size_pt': 10.0,
        }
        try:
            styles_el = doc.element.find(qn('w:body'))
            # docDefaults에서 rPrDefault 찾기
            root = doc.element
            doc_defaults_el = root.find(f'.//{qn("w:docDefaults")}')
            if doc_defaults_el is not None:
                rpr_default = doc_defaults_el.find(f'.//{qn("w:rPrDefault")}')
                if rpr_default is not None:
                    rpr = rpr_default.find(qn('w:rPr'))
                    if rpr is not None:
                        rfonts = rpr.find(qn('w:rFonts'))
                        if rfonts is not None:
                            defaults['ascii_font'] = rfonts.get(qn('w:ascii'), defaults['ascii_font'])
                            defaults['east_asia_font'] = rfonts.get(qn('w:eastAsia'), '')
                        sz = rpr.find(qn('w:sz'))
                        if sz is not None:
                            defaults['font_size_pt'] = int(sz.get(qn('w:val'), '20')) / 2.0
        except Exception as e:
            print(f"[DocxTableRenderer] 문서 기본값 파싱 오류: {e}")
        return defaults

    def _parse_table_grid(self, table: Table) -> List[int]:
        """테이블 그리드 컬럼 너비(twips) 추출"""
        grid = table._tbl.find(qn('w:tblGrid'))
        if grid is not None:
            return [int(gc.get(qn('w:w'), '0')) for gc in grid.findall(qn('w:gridCol'))]
        return []

    def _parse_table_borders(self, table: Table) -> CellBorders:
        """테이블 레벨 기본 테두리"""
        tbl_pr = table._tbl.find(qn('w:tblPr'))
        if tbl_pr is not None:
            borders_el = tbl_pr.find(qn('w:tblBorders'))
            if borders_el is not None:
                return self._parse_borders_element(borders_el)
        return CellBorders()

    def _parse_table_cell_margins(self, table: Table) -> Dict[str, int]:
        """테이블 레벨 셀 마진(twips)"""
        margins = {'top': 0, 'bottom': 0, 'left': 108, 'right': 108}
        tbl_pr = table._tbl.find(qn('w:tblPr'))
        if tbl_pr is not None:
            cm = tbl_pr.find(qn('w:tblCellMar'))
            if cm is not None:
                for side in ('top', 'bottom', 'left', 'right'):
                    el = cm.find(qn(f'w:{side}'))
                    if el is not None:
                        margins[side] = int(el.get(qn('w:w'), '0'))
        return margins

    def _parse_row_heights(self, table: Table) -> List[Optional[int]]:
        """각 행의 높이(twips) 추출"""
        heights = []
        for row in table.rows:
            tr_pr = row._tr.find(qn('w:trPr'))
            height = None
            if tr_pr is not None:
                trh = tr_pr.find(qn('w:trHeight'))
                if trh is not None:
                    height = int(trh.get(qn('w:val'), '0'))
            heights.append(height)
        return heights

    def _parse_borders_element(self, borders_el) -> CellBorders:
        """w:tblBorders 또는 w:tcBorders 요소에서 테두리 파싱"""
        result = CellBorders()
        for side_name in ('top', 'bottom', 'left', 'right'):
            el = borders_el.find(qn(f'w:{side_name}'))
            if el is not None:
                val = el.get(qn('w:val'), 'single')
                sz = int(el.get(qn('w:sz'), '4'))
                color = el.get(qn('w:color'), '000000')
                setattr(result, side_name, BorderSide(
                    style=BORDER_STYLE_MAP.get(val, 'solid'),
                    width=max(0.5, sz / 8.0),
                    color=f'#{color}' if color not in ('auto', '') else '#000000'
                ))
        return result

    # ── 셀 매트릭스 구축 ─────────────────────────────────

    def _build_cell_matrix(self, table: Table, grid_widths: List[int],
                           doc_defaults: dict, tbl_borders: CellBorders,
                           tbl_margins: Dict[str, int]) -> List[List[CellInfo]]:
        """
        병합을 고려한 셀 매트릭스 구축.
        각 행은 CellInfo 리스트로 구성됨.
        """
        num_grid_cols = len(grid_widths) if grid_widths else 0

        # Phase 1: 원시 셀 정보 수집
        raw_rows = []
        for row_idx, row in enumerate(table.rows):
            tcs = row._tr.findall(qn('w:tc'))
            row_cells = []
            grid_col = 0
            for tc in tcs:
                tc_pr = tc.find(qn('w:tcPr'))

                # gridSpan (colspan)
                gs_el = tc_pr.find(qn('w:gridSpan')) if tc_pr is not None else None
                grid_span = int(gs_el.get(qn('w:val'))) if gs_el is not None else 1

                # vMerge
                vm_el = tc_pr.find(qn('w:vMerge')) if tc_pr is not None else None
                if vm_el is not None:
                    vm_val = vm_el.get(qn('w:val'))
                    vmerge = vm_val if vm_val else 'continue'
                else:
                    vmerge = None

                # 셀 속성
                props = self._parse_cell_props(tc_pr, tbl_borders, tbl_margins, grid_widths, grid_col, grid_span)
                props.grid_span = grid_span
                props.vmerge = vmerge

                # 텍스트/서식
                paragraphs = self._parse_cell_content(tc, doc_defaults)
                cell_text = self._get_tc_text(tc)

                row_cells.append({
                    'grid_col': grid_col,
                    'grid_span': grid_span,
                    'vmerge': vmerge,
                    'props': props,
                    'paragraphs': paragraphs,
                    'text': cell_text,
                    'tc': tc,
                })
                grid_col += grid_span
            raw_rows.append(row_cells)

        # Phase 2: rowspan 계산
        for row_idx, row_cells in enumerate(raw_rows):
            for cell in row_cells:
                if cell['vmerge'] == 'restart':
                    span = 1
                    for next_row_idx in range(row_idx + 1, len(raw_rows)):
                        found = False
                        for next_cell in raw_rows[next_row_idx]:
                            if next_cell['grid_col'] == cell['grid_col'] and next_cell['vmerge'] == 'continue':
                                span += 1
                                found = True
                                break
                        if not found:
                            break
                    cell['rowspan'] = span
                elif cell['vmerge'] == 'continue':
                    cell['rowspan'] = 0  # 병합 연속 셀
                else:
                    cell['rowspan'] = 1

        # Phase 3: CellInfo 매트릭스 생성
        matrix = []
        for row_idx, row_cells in enumerate(raw_rows):
            row_infos = []
            for cell in row_cells:
                is_continuation = (cell['vmerge'] == 'continue')
                cell['props'].rowspan = cell['rowspan']
                info = CellInfo(
                    row=row_idx,
                    col=cell['grid_col'],
                    grid_span=cell['grid_span'],
                    rowspan=cell['rowspan'],
                    is_merge_continuation=is_continuation,
                    props=cell['props'],
                    paragraphs=cell['paragraphs'],
                    text=cell['text'],
                )
                row_infos.append(info)
            matrix.append(row_infos)

        return matrix

    def _parse_cell_props(self, tc_pr, tbl_borders: CellBorders,
                          tbl_margins: Dict[str, int],
                          grid_widths: List[int], grid_col: int, grid_span: int) -> CellProps:
        """셀 속성 파싱"""
        props = CellProps()

        # 너비: grid_widths에서 계산
        if grid_widths and grid_col < len(grid_widths):
            end_col = min(grid_col + grid_span, len(grid_widths))
            props.width_twips = sum(grid_widths[grid_col:end_col])

        if tc_pr is None:
            props.borders = CellBorders(
                top=BorderSide(tbl_borders.top.style, tbl_borders.top.width, tbl_borders.top.color),
                bottom=BorderSide(tbl_borders.bottom.style, tbl_borders.bottom.width, tbl_borders.bottom.color),
                left=BorderSide(tbl_borders.left.style, tbl_borders.left.width, tbl_borders.left.color),
                right=BorderSide(tbl_borders.right.style, tbl_borders.right.width, tbl_borders.right.color),
            )
            props.margins = dict(tbl_margins)
            return props

        # 테두리
        tc_borders = tc_pr.find(qn('w:tcBorders'))
        if tc_borders is not None:
            props.borders = self._parse_borders_element(tc_borders)
        else:
            # insideH, insideV 기본 테두리 적용
            props.borders = CellBorders(
                top=BorderSide(tbl_borders.top.style, tbl_borders.top.width, tbl_borders.top.color),
                bottom=BorderSide(tbl_borders.bottom.style, tbl_borders.bottom.width, tbl_borders.bottom.color),
                left=BorderSide(tbl_borders.left.style, tbl_borders.left.width, tbl_borders.left.color),
                right=BorderSide(tbl_borders.right.style, tbl_borders.right.width, tbl_borders.right.color),
            )

        # 배경색
        shd = tc_pr.find(qn('w:shd'))
        if shd is not None:
            fill = shd.get(qn('w:fill'))
            if fill and fill.lower() not in ('auto', 'ffffff', ''):
                props.shading_fill = fill

        # 수직 정렬
        v_align = tc_pr.find(qn('w:vAlign'))
        if v_align is not None:
            props.vertical_align = v_align.get(qn('w:val'), 'top')
        else:
            props.vertical_align = 'center'  # DOCX 기본값

        # 셀 마진
        tc_mar = tc_pr.find(qn('w:tcMar'))
        if tc_mar is not None:
            margins = {}
            for side in ('top', 'bottom', 'left', 'right'):
                el = tc_mar.find(qn(f'w:{side}'))
                margins[side] = int(el.get(qn('w:w'), '0')) if el is not None else tbl_margins.get(side, 0)
            props.margins = margins
        else:
            props.margins = dict(tbl_margins)

        return props

    def _parse_cell_content(self, tc, doc_defaults: dict) -> List[Tuple[ParaFormat, List[Tuple[str, RunFormat]]]]:
        """셀 내 단락/런 서식 파싱"""
        paragraphs = []
        for p_el in tc.findall(qn('w:p')):
            para_fmt = self._parse_paragraph_format(p_el, doc_defaults)
            runs = []
            for r_el in p_el.findall(qn('w:r')):
                run_fmt = self._parse_run_format(r_el, doc_defaults)
                text = ''
                for child in r_el:
                    if child.tag == qn('w:t'):
                        text += child.text or ''
                    elif child.tag == qn('w:br'):
                        text += '\n'
                    elif child.tag == qn('w:tab'):
                        text += '\t'
                if text:
                    runs.append((text, run_fmt))
            paragraphs.append((para_fmt, runs))
        return paragraphs

    def _parse_paragraph_format(self, p_el, doc_defaults: dict) -> ParaFormat:
        """단락 서식 파싱"""
        fmt = ParaFormat()
        p_pr = p_el.find(qn('w:pPr'))
        if p_pr is not None:
            jc = p_pr.find(qn('w:jc'))
            if jc is not None:
                val = jc.get(qn('w:val'), 'left')
                align_map = {'left': 'left', 'center': 'center', 'right': 'right',
                             'both': 'justify', 'distribute': 'justify'}
                fmt.alignment = align_map.get(val, 'left')

            spacing = p_pr.find(qn('w:spacing'))
            if spacing is not None:
                line_val = spacing.get(qn('w:line'))
                line_rule = spacing.get(qn('w:lineRule'))
                if line_val and line_rule == 'exact':
                    fmt.line_spacing_pt = int(line_val) / 20.0
                before = spacing.get(qn('w:before'))
                after = spacing.get(qn('w:after'))
                if before:
                    fmt.space_before_pt = int(before) / 20.0
                if after:
                    fmt.space_after_pt = int(after) / 20.0
        return fmt

    def _parse_run_format(self, r_el, doc_defaults: dict) -> RunFormat:
        """런 서식 파싱"""
        fmt = RunFormat(
            font_name=doc_defaults.get('ascii_font', ''),
            font_name_east_asia=doc_defaults.get('east_asia_font', ''),
            font_size_pt=doc_defaults.get('font_size_pt', 10.0),
        )
        r_pr = r_el.find(qn('w:rPr'))
        if r_pr is None:
            return fmt

        # 폰트
        rfonts = r_pr.find(qn('w:rFonts'))
        if rfonts is not None:
            ascii_font = rfonts.get(qn('w:ascii')) or rfonts.get(qn('w:hAnsi'))
            ea_font = rfonts.get(qn('w:eastAsia'))
            if ascii_font:
                fmt.font_name = ascii_font
            if ea_font:
                fmt.font_name_east_asia = ea_font

        # 크기
        sz = r_pr.find(qn('w:sz'))
        if sz is not None:
            val = sz.get(qn('w:val'))
            if val:
                fmt.font_size_pt = int(val) / 2.0

        # 볼드
        b = r_pr.find(qn('w:b'))
        if b is not None:
            val = b.get(qn('w:val'))
            fmt.bold = val != '0' if val is not None else True

        # 이탤릭
        i = r_pr.find(qn('w:i'))
        if i is not None:
            val = i.get(qn('w:val'))
            fmt.italic = val != '0' if val is not None else True

        # 밑줄
        u = r_pr.find(qn('w:u'))
        if u is not None:
            val = u.get(qn('w:val'))
            fmt.underline = val not in ('none', None, '')

        # 색상
        color_el = r_pr.find(qn('w:color'))
        if color_el is not None:
            c = color_el.get(qn('w:val'))
            if c and c.lower() != 'auto':
                fmt.color = f'#{c}'

        return fmt

    @staticmethod
    def _get_tc_text(tc) -> str:
        """tc 요소에서 순수 텍스트 추출"""
        texts = []
        for p in tc.findall(qn('w:p')):
            p_texts = []
            for r in p.findall(qn('w:r')):
                for t in r.findall(qn('w:t')):
                    if t.text:
                        p_texts.append(t.text)
            texts.append(''.join(p_texts))
        return '\n'.join(texts).strip()

    # ── PIL 렌더링 ──────────────────────────────────────

    def _find_korean_font(self) -> Optional[str]:
        """한글 지원 폰트 경로 탐색"""
        font_paths = [
            "/tmp/NotoSansKR.ttf",
            "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
            "/usr/share/fonts/truetype/nanum/NanumBarunGothic.ttf",
            "/usr/share/fonts/nanum/NanumGothic.ttf",
            "/usr/share/fonts/truetype/nanum/NanumMyeongjo.ttf",
            "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        ]
        for p in font_paths:
            if os.path.exists(p):
                return p
        # 자동 다운로드 시도
        try:
            import urllib.request
            noto_path = "/tmp/NotoSansKR.ttf"
            font_url = "https://github.com/google/fonts/raw/main/ofl/notosanskr/NotoSansKR%5Bwght%5D.ttf"
            urllib.request.urlretrieve(font_url, noto_path)
            if os.path.exists(noto_path) and os.path.getsize(noto_path) > 1000000:
                print(f"[DocxTableRenderer] 한글 폰트 다운로드 성공: {noto_path}")
                return noto_path
        except Exception as e:
            print(f"[DocxTableRenderer] 한글 폰트 다운로드 실패: {e}")
        return None

    def _get_pil_font(self, size_pt: float, bold: bool = False) -> ImageFont.FreeTypeFont:
        """지정된 크기의 PIL 폰트를 캐시에서 가져오거나 생성"""
        px_size = max(8, int(round(size_pt * 96.0 / 72.0 * self.dpi_scale)))
        key = (px_size, bold)
        if key not in self._font_cache:
            font = None
            if self._base_font_path:
                try:
                    font = ImageFont.truetype(self._base_font_path, px_size)
                except Exception:
                    pass
            if font is None:
                try:
                    font = ImageFont.load_default(size=px_size)
                except TypeError:
                    font = ImageFont.load_default()
            self._font_cache[key] = font
        return self._font_cache[key]

    def _render_table_with_pil(self, cell_matrix: List[List[CellInfo]],
                               grid_widths: List[int],
                               row_heights: List[Optional[int]],
                               doc_defaults: dict) -> Optional[bytes]:
        """셀 매트릭스를 PIL로 직접 렌더링하여 PNG bytes 반환"""
        scale = self.dpi_scale
        twip_px = TWIP_TO_PX * scale

        # 컬럼 너비 (px)
        if grid_widths:
            col_widths_px = [w * twip_px for w in grid_widths]
        else:
            col_widths_px = [200 * scale]
        num_cols = len(col_widths_px)
        num_rows = len(cell_matrix)

        if num_rows == 0:
            return None

        # 행별 최소 높이 계산 (텍스트 측정)
        min_row_heights = [0.0] * num_rows

        for row_idx, row_cells in enumerate(cell_matrix):
            for cell in row_cells:
                if cell.is_merge_continuation or cell.rowspan <= 0:
                    continue

                # 셀 너비 계산
                end_col = min(cell.col + cell.grid_span, num_cols)
                cell_w = sum(col_widths_px[cell.col:end_col])

                # 셀 마진
                pad_l = cell.props.margins.get('left', 108) * twip_px
                pad_r = cell.props.margins.get('right', 108) * twip_px
                pad_t = cell.props.margins.get('top', 0) * twip_px
                pad_b = cell.props.margins.get('bottom', 0) * twip_px
                inner_w = max(cell_w - pad_l - pad_r, 20 * scale)

                # 텍스트 높이 측정
                text_h = self._measure_cell_text_height(cell, inner_w, doc_defaults)
                needed_h = text_h + pad_t + pad_b + 4 * scale

                if cell.rowspan == 1:
                    min_row_heights[row_idx] = max(min_row_heights[row_idx], needed_h)
                else:
                    # 병합된 셀: 높이를 해당 행들에 분배
                    per_row = needed_h / cell.rowspan
                    for r in range(row_idx, min(row_idx + cell.rowspan, num_rows)):
                        min_row_heights[r] = max(min_row_heights[r], per_row)

        # DOCX 지정 행 높이와 비교하여 최종 높이 결정
        final_row_heights = []
        for i in range(num_rows):
            docx_h = row_heights[i] * twip_px if i < len(row_heights) and row_heights[i] else 0
            final_row_heights.append(max(min_row_heights[i], docx_h, 24 * scale))

        # 좌표 계산
        col_positions = [0.0]
        for w in col_widths_px:
            col_positions.append(col_positions[-1] + w)

        row_positions = [0.0]
        for h in final_row_heights:
            row_positions.append(row_positions[-1] + h)

        total_w = int(col_positions[-1]) + 2
        total_h = int(row_positions[-1]) + 2

        # 이미지 생성
        image = Image.new('RGB', (total_w, total_h), 'white')
        draw = ImageDraw.Draw(image)

        # 셀 렌더링
        for row_cells in cell_matrix:
            for cell in row_cells:
                if cell.is_merge_continuation or cell.rowspan <= 0:
                    continue
                self._draw_cell_pil(draw, cell, col_positions, row_positions,
                                    final_row_heights, doc_defaults)

        # PNG bytes
        buf = io.BytesIO()
        image.save(buf, format='PNG')
        return buf.getvalue()

    def _measure_cell_text_height(self, cell: CellInfo, inner_width: float,
                                  doc_defaults: dict) -> float:
        """셀 내 텍스트의 총 높이 측정"""
        scale = self.dpi_scale
        pt_to_px = scale * 96.0 / 72.0
        total_h = 0.0

        for para_fmt, runs in cell.paragraphs:
            if not runs:
                font = self._get_pil_font(doc_defaults.get('font_size_pt', 10.0))
                total_h += self._get_font_line_height(font) + 2 * scale
                continue

            total_h += para_fmt.space_before_pt * pt_to_px

            # 줄바꿈 계산
            para_lines = self._wrap_paragraph_runs(runs, inner_width, doc_defaults)
            for line_font, _ in para_lines:
                line_h = self._get_font_line_height(line_font)
                if para_fmt.line_spacing_pt:
                    line_h = max(line_h, para_fmt.line_spacing_pt * pt_to_px)
                total_h += line_h

            total_h += para_fmt.space_after_pt * pt_to_px

        return max(total_h, 12 * scale)

    def _wrap_paragraph_runs(self, runs: List[Tuple[str, RunFormat]],
                             max_width: float,
                             doc_defaults: dict) -> List[Tuple[ImageFont.FreeTypeFont, str]]:
        """런들을 줄바꿈하여 (font, text) 리스트 반환"""
        if not runs:
            return []

        first_run_fmt = runs[0][1]
        font = self._get_pil_font(first_run_fmt.font_size_pt, first_run_fmt.bold)

        # 전체 텍스트 합치기
        full_text = ''.join(text for text, _ in runs)

        lines = []
        for raw_line in full_text.split('\n'):
            if not raw_line:
                lines.append((font, ''))
                continue
            wrapped = self._wrap_text_to_width(raw_line, font, max_width)
            for wl in wrapped:
                lines.append((font, wl))

        return lines if lines else [(font, '')]

    def _wrap_text_to_width(self, text: str, font: ImageFont.FreeTypeFont,
                            max_width: float) -> List[str]:
        """텍스트를 지정된 픽셀 너비에 맞게 줄바꿈"""
        if max_width <= 0:
            return [text]

        try:
            bbox = font.getbbox(text)
            text_w = bbox[2] - bbox[0]
        except Exception:
            text_w = len(text) * 10

        if text_w <= max_width:
            return [text]

        # 글자 단위 줄바꿈 (한글 호환)
        lines = []
        current = ''
        for char in text:
            test = current + char
            try:
                bbox = font.getbbox(test)
                w = bbox[2] - bbox[0]
            except Exception:
                w = len(test) * 10

            if w > max_width and current:
                lines.append(current)
                current = char
            else:
                current = test

        if current:
            lines.append(current)

        return lines if lines else [text]

    def _get_font_line_height(self, font: ImageFont.FreeTypeFont) -> float:
        """폰트의 한 줄 높이 반환"""
        try:
            bbox = font.getbbox('가Ag')
            return (bbox[3] - bbox[1]) * 1.3
        except Exception:
            return 16 * self.dpi_scale

    def _draw_cell_pil(self, draw: ImageDraw.Draw, cell: CellInfo,
                       col_positions: List[float], row_positions: List[float],
                       row_heights: List[float], doc_defaults: dict):
        """단일 셀을 PIL로 그리기"""
        num_cols = len(col_positions) - 1
        num_rows = len(row_positions) - 1
        scale = self.dpi_scale
        twip_px = TWIP_TO_PX * scale
        pt_to_px = scale * 96.0 / 72.0

        # 셀 영역 계산
        c0 = cell.col
        c1 = min(cell.col + cell.grid_span, num_cols)
        r0 = cell.row
        r1 = min(cell.row + cell.rowspan, num_rows)

        x0 = col_positions[c0]
        x1 = col_positions[c1]
        y0 = row_positions[r0]
        y1 = row_positions[r1]

        # 배경색
        if cell.props.shading_fill:
            try:
                fill_color = f'#{cell.props.shading_fill}'
                draw.rectangle([x0, y0, x1, y1], fill=fill_color)
            except Exception:
                pass

        # 테두리 그리기
        self._draw_cell_borders(draw, cell, x0, y0, x1, y1, scale)

        # 패딩
        pad_l = cell.props.margins.get('left', 108) * twip_px
        pad_r = cell.props.margins.get('right', 108) * twip_px
        pad_t = cell.props.margins.get('top', 0) * twip_px + 2 * scale
        pad_b = cell.props.margins.get('bottom', 0) * twip_px + 2 * scale

        inner_x0 = x0 + pad_l
        inner_x1 = x1 - pad_r
        inner_y0 = y0 + pad_t
        inner_y1 = y1 - pad_b
        inner_w = max(inner_x1 - inner_x0, 1)
        inner_h = max(inner_y1 - inner_y0, 1)

        # 텍스트 높이 측정 (수직 정렬을 위해)
        text_total_h = self._measure_cell_text_height(cell, inner_w, doc_defaults)

        # 수직 정렬 오프셋
        valign = cell.props.vertical_align
        if valign in ('center', 'middle'):
            y_offset = max(0, (inner_h - text_total_h) / 2)
        elif valign == 'bottom':
            y_offset = max(0, inner_h - text_total_h)
        else:
            y_offset = 0

        # 텍스트 그리기
        cursor_y = inner_y0 + y_offset
        for para_fmt, runs in cell.paragraphs:
            cursor_y += para_fmt.space_before_pt * pt_to_px

            if not runs:
                font = self._get_pil_font(doc_defaults.get('font_size_pt', 10.0))
                cursor_y += self._get_font_line_height(font)
                cursor_y += para_fmt.space_after_pt * pt_to_px
                continue

            # 줄바꿈
            para_lines = self._wrap_paragraph_runs(runs, inner_w, doc_defaults)

            # 텍스트 색상
            first_run_fmt = runs[0][1]
            text_color = first_run_fmt.color if first_run_fmt.color else '#000000'

            for line_font, line_text in para_lines:
                line_h = self._get_font_line_height(line_font)
                if para_fmt.line_spacing_pt:
                    line_h = max(line_h, para_fmt.line_spacing_pt * pt_to_px)

                if cursor_y + line_h > y1:
                    break  # 셀 밖으로 나가면 중단

                if line_text.strip():
                    # 수평 정렬
                    try:
                        bbox = line_font.getbbox(line_text)
                        text_w = bbox[2] - bbox[0]
                    except Exception:
                        text_w = len(line_text) * 10

                    if para_fmt.alignment == 'center':
                        tx = inner_x0 + (inner_w - text_w) / 2
                    elif para_fmt.alignment == 'right':
                        tx = inner_x1 - text_w
                    else:
                        tx = inner_x0

                    try:
                        draw.text((tx, cursor_y), line_text, fill=text_color, font=line_font)
                        # 볼드 시뮬레이션: 1px 오프셋으로 재렌더링
                        if first_run_fmt.bold:
                            draw.text((tx + 1, cursor_y), line_text, fill=text_color, font=line_font)
                    except Exception:
                        pass

                cursor_y += line_h

            cursor_y += para_fmt.space_after_pt * pt_to_px

    def _draw_cell_borders(self, draw: ImageDraw.Draw, cell: CellInfo,
                           x0: float, y0: float, x1: float, y1: float,
                           scale: int):
        """셀 테두리 그리기"""
        borders = cell.props.borders
        for side_name in ('top', 'bottom', 'left', 'right'):
            b = getattr(borders, side_name)
            if b.style == 'none':
                continue

            width = max(1, int(b.width * scale))
            color = b.color if b.color else '#000000'

            if side_name == 'top':
                draw.line([(x0, y0), (x1, y0)], fill=color, width=width)
            elif side_name == 'bottom':
                draw.line([(x0, y1), (x1, y1)], fill=color, width=width)
            elif side_name == 'left':
                draw.line([(x0, y0), (x0, y1)], fill=color, width=width)
            elif side_name == 'right':
                draw.line([(x1, y0), (x1, y1)], fill=color, width=width)

    # ── 호환성 헬퍼 (TableToImageConverter에서 재사용) ────

    def _extract_plain_table_data(self, table: Table) -> List[List[str]]:
        """표에서 순수 텍스트 2차원 배열 추출"""
        data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            data.append(row_data)
        return data

    def _is_metadata_table(self, table: Table) -> bool:
        """메타데이터 표(제정일, 승인란 등) 여부 확인"""
        if not table.rows:
            return False

        first_row_text = " ".join(cell.text.strip() for cell in table.rows[0].cells)
        metadata_patterns = ["제정일", "최종개정일", "최종검토일", "소관부서", "승인", "승  인", "병원장", "(인)"]
        for pattern in metadata_patterns:
            if pattern in first_row_text:
                return True

        if len(table.rows) > 1:
            last_row_text = " ".join(cell.text.strip() for cell in table.rows[-1].cells)
            for pattern in ["승인", "승  인", "병원장", "(인)"]:
                if pattern in last_row_text:
                    return True

        return False

    def _get_table_context(self, doc, table: Table) -> Dict[str, str]:
        """표 주변 텍스트 컨텍스트 추출"""
        context = {
            "previous_text": "",
            "next_text": "",
            "context": "표",
            "table_location": "문서 내 표"
        }
        try:
            elements = []
            for element in doc.element.body:
                if element.tag.endswith('}p'):
                    para = Paragraph(element, doc)
                    if para.text.strip():
                        elements.append(('paragraph', para.text.strip()))
                elif element.tag.endswith('}tbl'):
                    elements.append(('table', element))

            table_index = -1
            for i, (etype, edata) in enumerate(elements):
                if etype == 'table' and edata == table._tbl:
                    table_index = i
                    break

            if table_index >= 0:
                prev_texts = [elements[i][1] for i in range(max(0, table_index - 3), table_index)
                              if elements[i][0] == 'paragraph']
                context["previous_text"] = "\n".join(prev_texts[-3:])

                next_texts = [elements[i][1] for i in range(table_index + 1, min(len(elements), table_index + 4))
                              if elements[i][0] == 'paragraph']
                context["next_text"] = "\n".join(next_texts[:3])

        except Exception as e:
            print(f"[DocxTableRenderer] 컨텍스트 추출 오류: {e}")

        return context

    def _filter_tables_by_article_range(self, doc) -> List[Tuple[int, Table]]:
        """조문 범위 내 표만 필터링"""
        if not self.article_range:
            return [(i, t) for i, t in enumerate(doc.tables)]

        start_article, end_article = self.article_range
        filtered = []
        current_article = 0

        elements = []
        for element in doc.element.body:
            if element.tag.endswith('}p'):
                para = Paragraph(element, doc)
                text = para.text.strip()
                match = re.search(r'제\s*(\d+)\s*조', text)
                if match:
                    current_article = int(match.group(1))
                elements.append(('paragraph', current_article, element))
            elif element.tag.endswith('}tbl'):
                elements.append(('table', current_article, element))

        for etype, article_num, edata in elements:
            if etype == 'table' and start_article <= article_num <= end_article:
                for doc_idx, doc_table in enumerate(doc.tables):
                    if doc_table._tbl == edata and not self._is_metadata_table(doc_table):
                        filtered.append((len(filtered), doc_table))
                        break

        return filtered
