"""
신구대비표 생성 라우터
- 현행 규정과 연혁 규정의 JSON 비교
- 변경 사항 추출 (추가, 삭제, 수정)
- 신구대비표 HTML/PDF 생성
- 파일 업로드를 통한 실시간 비교
"""
from fastapi import APIRouter, Depends, HTTPException, Query, File, UploadFile, Form
from fastapi.responses import JSONResponse, HTMLResponse, FileResponse
from typing import Dict, Any, List, Optional
from collections import OrderedDict
from api.auth_middleware import get_current_user
from api.timescaledb_manager_v2 import DatabaseConnectionManager
try:
    from api.router_rule_history import create_history_record
except ImportError:
    create_history_record = None
from pathlib import Path
import json
import logging
from datetime import datetime
from difflib import SequenceMatcher
import re
import os
import sys
import tempfile
import shutil

# PDF 생성용 reportlab 임포트
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# 한글 폰트 등록 (우선순위: KB금융체 > NanumGothic > DroidSansFallback)
_korean_font_registered = False
try:
    # 1차 시도: KB금융체
    kb_font_paths = [
        os.path.expanduser("~/.local/share/fonts/kb-finance/KBfgTextM.ttf"),
    ]
    for font_path in kb_font_paths:
        if os.path.exists(font_path):
            pdfmetrics.registerFont(TTFont('KoreanFont', font_path))
            bold_path = font_path.replace("KBfgTextM.ttf", "KBfgTextB.ttf")
            if os.path.exists(bold_path):
                pdfmetrics.registerFont(TTFont('KoreanFontBold', bold_path))
            else:
                pdfmetrics.registerFont(TTFont('KoreanFontBold', font_path))
            _korean_font_registered = True
            break

    # 2차 시도: NanumGothic
    if not _korean_font_registered:
        nanum_paths = [
            "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
            "/usr/share/fonts/naver-nanum/NanumGothic.ttf"
        ]
        for font_path in nanum_paths:
            if os.path.exists(font_path):
                pdfmetrics.registerFont(TTFont('KoreanFont', font_path))
                bold_path = font_path.replace("NanumGothic.ttf", "NanumGothicBold.ttf")
                if os.path.exists(bold_path):
                    pdfmetrics.registerFont(TTFont('KoreanFontBold', bold_path))
                else:
                    pdfmetrics.registerFont(TTFont('KoreanFontBold', font_path))
                _korean_font_registered = True
                break

    # 3차 시도: DroidSansFallbackFull (CJK 지원)
    if not _korean_font_registered:
        droid_path = "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf"
        if os.path.exists(droid_path):
            pdfmetrics.registerFont(TTFont('KoreanFont', droid_path))
            pdfmetrics.registerFont(TTFont('KoreanFontBold', droid_path))
            _korean_font_registered = True

except Exception as e:
    _logger = logging.getLogger(__name__)
    _logger.warning(f"Failed to register Korean font: {e}")

router = APIRouter(prefix="/api/v1/compare", tags=["compare"])
logger = logging.getLogger(__name__)

# 경로 설정
APPLIB_DIR = Path("/home/wizice/kbregulation/fastapi/applib")
MERGE_JSON_DIR = APPLIB_DIR / "merge_json"
MERGE_JSON_OLD_DIR = APPLIB_DIR / "merge_json_old"
DOCX_JSON_DIR = APPLIB_DIR / "docx_json"
WWW_JSON_DIR = Path("/home/wizice/kbregulation/www/static/file")

# DB 연결
db_manager = DatabaseConnectionManager()


def get_regulation_info(rule_id: int) -> Optional[Dict]:
    """규정 정보 조회"""
    try:
        with db_manager.get_connection() as conn:
            with conn.cursor() as cur:
                cur.execute("""
                    SELECT wzruleseq, wzruleid, wzname, wzpubno,
                           wzfilejson, wzcontent_path, wznewflag,
                           wzestabdate, wzlastrevdate, wzexecdate
                    FROM wz_rule
                    WHERE wzruleseq = %s
                """, (rule_id,))
                row = cur.fetchone()

                if row:
                    return {
                        'wzruleseq': row[0],
                        'wzruleid': row[1],
                        'wzname': row[2],
                        'wzpubno': row[3],
                        'wzfilejson': row[4],
                        'wzcontent_path': row[5],
                        'wznewflag': row[6],
                        'wzestabdate': row[7],
                        'wzlastrevdate': row[8],
                        'wzexecdate': row[9]
                    }
                return None
    except Exception as e:
        logger.error(f"Error getting regulation info: {e}")
        return None


def load_json_content(rule_info: Dict) -> Optional[Dict]:
    """규정의 JSON 내용 로드 (버전별 파일 우선)"""
    try:
        wzruleid = rule_info.get('wzruleid')
        wzruleseq = rule_info.get('wzruleseq')

        # 1. 버전별 파일 시도: merge_json/{wzruleid}_{wzruleseq}.json
        if wzruleid and wzruleseq:
            version_path = MERGE_JSON_DIR / f"{wzruleid}_{wzruleseq}.json"
            if version_path.exists():
                logger.info(f"Loading version-specific JSON: {version_path}")
                with open(version_path, 'r', encoding='utf-8') as f:
                    return json.load(f)

        # 2. wzfilejson 경로 시도 (DB에 저장된 경로)
        json_path = rule_info.get('wzfilejson') or rule_info.get('wzcontent_path')
        if json_path:
            if json_path.startswith('applib/'):
                full_path = APPLIB_DIR.parent / json_path
            elif json_path.startswith('/'):
                full_path = Path(json_path)
            else:
                full_path = APPLIB_DIR / json_path

            if full_path.exists():
                logger.info(f"Loading JSON from wzfilejson: {full_path}")
                with open(full_path, 'r', encoding='utf-8') as f:
                    return json.load(f)

        # 3. 현행 규정용: www/static/file/{wzruleid}.json
        if wzruleid and rule_info.get('wznewflag') == '현행':
            www_path = WWW_JSON_DIR / f"{wzruleid}.json"
            if www_path.exists():
                logger.info(f"Loading current JSON from www: {www_path}")
                with open(www_path, 'r', encoding='utf-8') as f:
                    return json.load(f)

        # 4. 기본 merge_json/{wzruleid}.json 시도
        if wzruleid:
            merge_path = MERGE_JSON_DIR / f"{wzruleid}.json"
            if merge_path.exists():
                logger.info(f"Loading default merge JSON: {merge_path}")
                with open(merge_path, 'r', encoding='utf-8') as f:
                    return json.load(f)

        return None
    except Exception as e:
        logger.error(f"Error loading JSON content: {e}")
        return None


def similarity_ratio(str1: str, str2: str) -> float:
    """두 문자열의 유사도 계산 (0.0 ~ 1.0)"""
    if not str1 or not str2:
        return 0.0
    return SequenceMatcher(None, str1, str2).ratio()


def normalize_text(text: str) -> str:
    """텍스트 정규화 (비교용)"""
    if not text:
        return ""
    text = re.sub(r'\s+', ' ', text.strip())
    return text


def _strip_html_tags(text: str) -> str:
    """HTML 태그 제거 (비교용)"""
    if not text:
        return ""
    return re.sub(r'</?(?:i|b|u|em|strong|font|strike|s|span|sup|sub)[^>]*>', '', text)


def highlight_diff(old_text: str, new_text: str, font_name: str = 'KoreanFont') -> tuple:
    """
    두 텍스트를 비교하여 변경된 부분만 표시

    Returns:
        (old_html, new_html):
            - old_html: 삭제된 부분에만 취소선
            - new_html: 추가된 부분에만 밑줄+파란색
    """
    old_plain = _strip_html_tags(old_text) if old_text else ""
    new_plain = _strip_html_tags(new_text) if new_text else ""

    if not old_plain and not new_plain:
        return "", ""

    if not old_plain:
        return "", f"<u><font face='{font_name}' color='blue'>{new_plain}</font></u>"

    if not new_plain:
        return f"<font face='{font_name}'><strike><font color='blue'>{old_plain}</font></strike></font>", ""

    # 단어 단위 SequenceMatcher로 차이 분석
    old_words = old_plain.split()
    new_words = new_plain.split()

    matcher = SequenceMatcher(None, old_words, new_words)
    opcodes = matcher.get_opcodes()

    old_parts = []
    new_parts = []

    for tag, i1, i2, j1, j2 in opcodes:
        old_chunk = ' '.join(old_words[i1:i2])
        new_chunk = ' '.join(new_words[j1:j2])

        if tag == 'equal':
            old_parts.append(old_chunk)
            new_parts.append(new_chunk)
        elif tag == 'delete':
            old_parts.append(f"<strike><font color='blue'>{old_chunk}</font></strike>")
        elif tag == 'insert':
            new_parts.append(f"<u><font color='blue'>{new_chunk}</font></u>")
        elif tag == 'replace':
            old_parts.append(f"<strike><font color='blue'>{old_chunk}</font></strike>")
            new_parts.append(f"<u><font color='blue'>{new_chunk}</font></u>")

    old_html = f"<font face='{font_name}'>{' '.join(old_parts)}</font>"
    new_html = f"<font face='{font_name}'>{' '.join(new_parts)}</font>"

    return old_html, new_html


def compare_articles(old_articles: List[Dict], new_articles: List[Dict]) -> List[Dict]:
    """
    두 조문 목록을 비교하여 변경 사항 추출

    Returns:
        변경 사항 리스트
    """
    changes = []

    # 레벨 1 기준으로 매칭
    old_by_number = {}
    new_by_number = {}

    for article in old_articles:
        if article.get('레벨') == 1:
            key = normalize_text(article.get('번호', ''))
            if key:
                old_by_number[key] = article

    for article in new_articles:
        if article.get('레벨') == 1:
            key = normalize_text(article.get('번호', ''))
            if key:
                new_by_number[key] = article

    all_keys = set(old_by_number.keys()) | set(new_by_number.keys())

    for key in sorted(all_keys, key=lambda x: extract_article_number(x)):
        old_art = old_by_number.get(key)
        new_art = new_by_number.get(key)

        if old_art and new_art:
            old_content = normalize_text(old_art.get('내용', ''))
            new_content = normalize_text(new_art.get('내용', ''))

            if old_content != new_content:
                sim = similarity_ratio(old_content, new_content)
                changes.append({
                    'type': 'modified',
                    'article_key': key,
                    'old': old_art,
                    'new': new_art,
                    'similarity': sim
                })
        elif old_art:
            changes.append({
                'type': 'deleted',
                'article_key': key,
                'old': old_art,
                'new': None,
                'similarity': 0
            })
        else:
            changes.append({
                'type': 'added',
                'article_key': key,
                'old': None,
                'new': new_art,
                'similarity': 0
            })

    # 하위 레벨 조문도 비교
    old_sub_map = build_sub_article_map(old_articles)
    new_sub_map = build_sub_article_map(new_articles)

    all_sub_keys = set(old_sub_map.keys()) | set(new_sub_map.keys())

    for key in all_sub_keys:
        old_sub = old_sub_map.get(key)
        new_sub = new_sub_map.get(key)

        if old_sub and new_sub:
            old_content = normalize_text(old_sub.get('내용', ''))
            new_content = normalize_text(new_sub.get('내용', ''))

            if old_content != new_content:
                sim = similarity_ratio(old_content, new_content)
                changes.append({
                    'type': 'modified',
                    'article_key': key,
                    'old': old_sub,
                    'new': new_sub,
                    'similarity': sim,
                    'level': old_sub.get('레벨', 2)
                })
        elif old_sub:
            changes.append({
                'type': 'deleted',
                'article_key': key,
                'old': old_sub,
                'new': None,
                'similarity': 0,
                'level': old_sub.get('레벨', 2)
            })
        elif new_sub:
            changes.append({
                'type': 'added',
                'article_key': key,
                'old': None,
                'new': new_sub,
                'similarity': 0,
                'level': new_sub.get('레벨', 2)
            })

    return changes


def build_path_map(articles: List[Dict]) -> OrderedDict:
    """각 조문에 계층 경로를 부여하고 OrderedDict로 반환"""
    path_map = OrderedDict()
    current_parents = {}
    current_parent_arts = {}

    for article in articles:
        level = article.get('레벨', 0)
        number = normalize_text(article.get('번호', ''))

        if level == 0:
            continue

        current_parents[level] = number
        current_parent_arts[level] = article

        for l in list(current_parents.keys()):
            if l > level:
                del current_parents[l]
        for l in list(current_parent_arts.keys()):
            if l > level:
                del current_parent_arts[l]

        path_parts = []
        for l in sorted(current_parents.keys()):
            if l <= level:
                path_parts.append(current_parents[l])
        path = '|'.join(path_parts)

        base_path = path
        counter = 1
        while path in path_map:
            path = f"{base_path}#{counter}"
            counter += 1

        parents = []
        for l in sorted(current_parent_arts.keys()):
            if l < level:
                parents.append(current_parent_arts[l])

        path_map[path] = {
            'article': article,
            'parents': parents.copy(),
            'level': level,
        }

    return path_map


def _get_parent_path(path: str) -> str:
    """경로에서 부모 경로 추출"""
    parts = path.split('|')
    return '|'.join(parts[:-1]) if len(parts) > 1 else ''


def _get_content_text(article: Dict) -> str:
    """비교용 내용 텍스트 추출"""
    return normalize_text(_strip_html_tags(article.get('내용', '')))


def _refine_matches_by_content(old_map: OrderedDict, new_map: OrderedDict):
    """
    경로 기반 매칭 후, 같은 부모 아래 자식 수가 다른 그룹에 대해
    내용 유사도 기반으로 재매칭하여 번호 변경(shift) 오류를 보정.
    """
    from collections import defaultdict

    old_groups = defaultdict(list)
    for path, entry in old_map.items():
        ppath = _get_parent_path(path)
        old_groups[ppath].append((path, entry))

    new_groups = defaultdict(list)
    for path, entry in new_map.items():
        ppath = _get_parent_path(path)
        new_groups[ppath].append((path, entry))

    refined_match = {}
    scope_old = set()
    scope_new = set()

    all_ppath = set(old_groups.keys()) | set(new_groups.keys())
    for ppath in all_ppath:
        old_children = old_groups.get(ppath, [])
        new_children = new_groups.get(ppath, [])

        if len(old_children) == len(new_children):
            continue
        if len(old_children) <= 1 or len(new_children) <= 1:
            continue

        for path, _ in old_children:
            scope_old.add(path)
        for path, _ in new_children:
            scope_new.add(path)

        old_contents = [_get_content_text(e['article']) for _, e in old_children]
        new_contents = [_get_content_text(e['article']) for _, e in new_children]

        matcher = SequenceMatcher(None, old_contents, new_contents)
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'equal':
                for oi, ni in zip(range(i1, i2), range(j1, j2)):
                    old_path = old_children[oi][0]
                    new_path = new_children[ni][0]
                    refined_match[old_path] = new_path
            elif tag == 'replace':
                old_indices = list(range(i1, i2))
                new_indices = list(range(j1, j2))
                if len(old_indices) == len(new_indices):
                    for oi, ni in zip(old_indices, new_indices):
                        refined_match[old_children[oi][0]] = new_children[ni][0]
                else:
                    sim_pairs = []
                    for oi in old_indices:
                        for ni in new_indices:
                            ratio = SequenceMatcher(None, old_contents[oi], new_contents[ni]).ratio()
                            sim_pairs.append((ratio, oi, ni))
                    sim_pairs.sort(reverse=True)
                    used_old = set()
                    used_new = set()
                    for ratio, oi, ni in sim_pairs:
                        if oi in used_old or ni in used_new:
                            continue
                        if ratio < 0.3:
                            break
                        refined_match[old_children[oi][0]] = new_children[ni][0]
                        used_old.add(oi)
                        used_new.add(ni)

    return refined_match, scope_old, scope_new


def compare_articles_with_hierarchy(old_articles: List[Dict], new_articles: List[Dict]) -> List[Dict]:
    """
    경로 기반 매칭 + 내용 유사도 보정을 사용한 신구대비표용 비교 함수
    """
    old_map = build_path_map(old_articles)
    new_map = build_path_map(new_articles)

    refined, scope_old, scope_new = _refine_matches_by_content(old_map, new_map)

    old_to_new = {}
    new_to_old = {}

    # 1) refined 매칭 우선 적용
    for old_path, new_path in refined.items():
        if old_path in old_map and new_path in new_map:
            old_to_new[old_path] = new_path
            new_to_old[new_path] = old_path

    # 2) 나머지는 경로 기반 매칭
    for path in new_map:
        if path in new_to_old:
            continue
        if path in scope_new:
            continue
        if path in old_map and path not in scope_old:
            old_to_new[path] = path
            new_to_old[path] = path

    # 3) 부모 경로가 재매칭된 경우 하위 항목도 전파
    propagated = {}
    for old_parent, new_parent in list(old_to_new.items()):
        if old_parent == new_parent:
            continue
        old_prefix = old_parent + '|'
        new_prefix = new_parent + '|'
        old_subs = {p[len(old_prefix):]: p for p in old_map
                    if p.startswith(old_prefix) and p not in old_to_new}
        new_subs = {p[len(new_prefix):]: p for p in new_map
                    if p.startswith(new_prefix) and p not in new_to_old}
        for suffix in old_subs:
            if suffix in new_subs:
                propagated[old_subs[suffix]] = new_subs[suffix]
    for old_p, new_p in propagated.items():
        old_to_new[old_p] = new_p
        new_to_old[new_p] = old_p

    matched_old = set(old_to_new.keys())
    matched_new = set(new_to_old.keys())

    changes_with_context = []
    shown_parents = set()
    output_paths = set()

    old_art_to_path = {id(entry['article']): path for path, entry in old_map.items()}
    new_art_to_path = {id(entry['article']): path for path, entry in new_map.items()}

    def add_parent_context(old_entry, new_entry):
        """변경된 항목의 부모를 컨텍스트로 추가"""
        old_parents = old_entry['parents'] if old_entry else []
        new_parents = new_entry['parents'] if new_entry else []

        old_parent_by_level = {p.get('레벨', 0): p for p in old_parents}
        new_parent_by_level = {p.get('레벨', 0): p for p in new_parents}
        all_parent_levels = sorted(set(old_parent_by_level.keys()) | set(new_parent_by_level.keys()))

        for parent_level in all_parent_levels:
            old_parent = old_parent_by_level.get(parent_level)
            new_parent = new_parent_by_level.get(parent_level)

            parent_paths = set()
            if old_parent:
                p = old_art_to_path.get(id(old_parent))
                if p:
                    parent_paths.add(p)
            if new_parent:
                p = new_art_to_path.get(id(new_parent))
                if p:
                    parent_paths.add(p)

            if parent_paths & (shown_parents | output_paths):
                continue

            shown_parents.update(parent_paths)
            ref_parent = old_parent or new_parent
            changes_with_context.append({
                'type': 'context',
                'old': old_parent if old_parent else new_parent,
                'new': new_parent if new_parent else old_parent,
                'level': ref_parent.get('레벨', 1),
                'is_parent': True
            })

    # 삭제 항목을 올바른 위치에 삽입하기 위한 준비
    old_path_list = list(old_map.keys())
    old_path_to_pos = {path: i for i, path in enumerate(old_path_list)}
    deleted_old_paths = [p for p in old_path_list if p not in matched_old]
    emitted_deleted = set()
    last_old_pos = -1

    def emit_deleted_between(from_pos, to_pos):
        """from_pos ~ to_pos 사이의 삭제 항목을 출력"""
        for del_path in deleted_old_paths:
            if del_path in emitted_deleted:
                continue
            del_pos = old_path_to_pos[del_path]
            if from_pos < del_pos < to_pos:
                del_entry = old_map[del_path]
                add_parent_context(del_entry, None)
                output_paths.add(del_path)
                changes_with_context.append({
                    'type': 'deleted',
                    'old': del_entry['article'],
                    'new': None,
                    'level': del_entry['level'],
                    'old_parents': del_entry['parents'],
                    'new_parents': [],
                    'is_parent': False
                })
                emitted_deleted.add(del_path)

    # new 문서 순서대로 순회
    for new_path, new_entry in new_map.items():
        if new_path in matched_new:
            old_path = new_to_old[new_path]
            old_entry = old_map[old_path]
            current_old_pos = old_path_to_pos.get(old_path, -1)

            emit_deleted_between(last_old_pos, current_old_pos)
            last_old_pos = current_old_pos

            old_art = old_entry['article']
            new_art = new_entry['article']
            old_content = normalize_text(_strip_html_tags(f"{old_art.get('번호', '')} {old_art.get('내용', '')}"))
            new_content = normalize_text(_strip_html_tags(f"{new_art.get('번호', '')} {new_art.get('내용', '')}"))

            if old_content != new_content:
                add_parent_context(old_entry, new_entry)
                output_paths.add(old_path)
                output_paths.add(new_path)
                changes_with_context.append({
                    'type': 'modified',
                    'old': old_art,
                    'new': new_art,
                    'level': new_art.get('레벨', 2),
                    'old_parents': old_entry['parents'],
                    'new_parents': new_entry['parents'],
                    'is_parent': False
                })
        else:
            add_parent_context(None, new_entry)
            output_paths.add(new_path)
            changes_with_context.append({
                'type': 'added',
                'old': None,
                'new': new_entry['article'],
                'level': new_entry['level'],
                'old_parents': [],
                'new_parents': new_entry['parents'],
                'is_parent': False
            })

    # 나머지 삭제 항목
    for del_path in deleted_old_paths:
        if del_path not in emitted_deleted:
            del_entry = old_map[del_path]
            add_parent_context(del_entry, None)
            output_paths.add(del_path)
            changes_with_context.append({
                'type': 'deleted',
                'old': del_entry['article'],
                'new': None,
                'level': del_entry['level'],
                'old_parents': del_entry['parents'],
                'new_parents': [],
                'is_parent': False
            })

    return changes_with_context


def build_sub_article_map(articles: List[Dict]) -> Dict[str, Dict]:
    """레벨 2 이상의 조문을 키로 매핑"""
    result = {}
    current_parent = ""

    for article in articles:
        level = article.get('레벨', 0)
        number = normalize_text(article.get('번호', ''))

        if level == 1:
            current_parent = number
        elif level >= 2 and number:
            key = f"{current_parent}_{number}"
            result[key] = article

    return result


def extract_article_number(key: str) -> tuple:
    """조문 번호에서 숫자 추출 (정렬용)"""
    numbers = re.findall(r'\d+', key)
    return tuple(int(n) for n in numbers) if numbers else (0,)


@router.get("/versions/{rule_id}")
async def get_rule_versions(
    rule_id: int,
    user: Dict[str, Any] = Depends(get_current_user)
):
    """특정 규정의 모든 버전 목록 조회 (현행 + 연혁)"""
    try:
        current_rule = get_regulation_info(rule_id)
        if not current_rule:
            raise HTTPException(status_code=404, detail="규정을 찾을 수 없습니다.")

        wzruleid = current_rule['wzruleid']

        with db_manager.get_connection() as conn:
            with conn.cursor() as cur:
                cur.execute("""
                    SELECT wzruleseq, wzruleid, wzname, wzpubno,
                           wznewflag, wzestabdate, wzlastrevdate, wzexecdate
                    FROM wz_rule
                    WHERE wzruleid = %s
                    ORDER BY
                        CASE WHEN wznewflag = '현행' THEN 0 ELSE 1 END,
                        wzlastrevdate DESC NULLS LAST,
                        wzruleseq DESC
                """, (wzruleid,))
                rows = cur.fetchall()

        versions = []
        for row in rows:
            versions.append({
                'wzruleseq': row[0],
                'wzruleid': row[1],
                'wzname': row[2],
                'wzpubno': row[3],
                'wznewflag': row[4],
                'wzestabdate': str(row[5]) if row[5] else None,
                'wzlastrevdate': str(row[6]) if row[6] else None,
                'wzexecdate': str(row[7]) if row[7] else None,
                'label': f"{row[4]} - {row[6] or row[5] or '날짜없음'}"
            })

        return JSONResponse(content={
            "success": True,
            "rule_id": rule_id,
            "wzruleid": wzruleid,
            "versions": versions,
            "total_count": len(versions)
        })

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting rule versions: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/diff")
async def compare_two_versions(
    old_rule_id: int = Query(..., description="구 규정 시퀀스"),
    new_rule_id: int = Query(..., description="신 규정 시퀀스"),
    user: Dict[str, Any] = Depends(get_current_user)
):
    """두 버전의 규정을 비교하여 신구대비표 데이터 생성"""
    try:
        old_rule = get_regulation_info(old_rule_id)
        new_rule = get_regulation_info(new_rule_id)

        if not old_rule:
            raise HTTPException(status_code=404, detail=f"구 규정을 찾을 수 없습니다: {old_rule_id}")
        if not new_rule:
            raise HTTPException(status_code=404, detail=f"신 규정을 찾을 수 없습니다: {new_rule_id}")

        old_json = load_json_content(old_rule)
        new_json = load_json_content(new_rule)

        if not old_json:
            raise HTTPException(status_code=404, detail="구 규정의 JSON 파일을 찾을 수 없습니다.")
        if not new_json:
            raise HTTPException(status_code=404, detail="신 규정의 JSON 파일을 찾을 수 없습니다.")

        old_articles = old_json.get('조문내용', [])
        new_articles = new_json.get('조문내용', [])

        changes = compare_articles(old_articles, new_articles)

        stats = {
            'total_changes': len(changes),
            'modified': len([c for c in changes if c['type'] == 'modified']),
            'added': len([c for c in changes if c['type'] == 'added']),
            'deleted': len([c for c in changes if c['type'] == 'deleted'])
        }

        return JSONResponse(content={
            "success": True,
            "old_rule": {
                "wzruleseq": old_rule['wzruleseq'],
                "wzname": old_rule['wzname'],
                "wzpubno": old_rule['wzpubno'],
                "wznewflag": old_rule['wznewflag'],
                "wzlastrevdate": str(old_rule['wzlastrevdate']) if old_rule['wzlastrevdate'] else None
            },
            "new_rule": {
                "wzruleseq": new_rule['wzruleseq'],
                "wzname": new_rule['wzname'],
                "wzpubno": new_rule['wzpubno'],
                "wznewflag": new_rule['wznewflag'],
                "wzlastrevdate": str(new_rule['wzlastrevdate']) if new_rule['wzlastrevdate'] else None
            },
            "changes": changes,
            "stats": stats
        })

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error comparing versions: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/html")
async def generate_comparison_html(
    old_rule_id: int = Query(..., description="구 규정 시퀀스"),
    new_rule_id: int = Query(..., description="신 규정 시퀀스"),
    user: Dict[str, Any] = Depends(get_current_user)
):
    """신구대비표 HTML 생성"""
    try:
        old_rule = get_regulation_info(old_rule_id)
        new_rule = get_regulation_info(new_rule_id)

        if not old_rule or not new_rule:
            raise HTTPException(status_code=404, detail="규정을 찾을 수 없습니다.")

        old_json = load_json_content(old_rule)
        new_json = load_json_content(new_rule)

        if not old_json or not new_json:
            raise HTTPException(status_code=404, detail="JSON 파일을 찾을 수 없습니다.")

        old_articles = old_json.get('조문내용', [])
        new_articles = new_json.get('조문내용', [])
        changes = compare_articles(old_articles, new_articles)

        html = generate_comparison_table_html(old_rule, new_rule, changes)

        return HTMLResponse(content=html)

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error generating comparison HTML: {e}")
        raise HTTPException(status_code=500, detail=str(e))


def generate_comparison_table_html(old_rule: Dict, new_rule: Dict, changes: List[Dict]) -> str:
    """신구대비표 HTML 테이블 생성"""

    rows_html = ""

    for change in changes:
        change_type = change['type']
        old_art = change.get('old')
        new_art = change.get('new')

        if change_type == 'modified':
            row_class = 'modified'
            old_style = 'background: #fff3cd;'
            new_style = 'background: #d4edda;'
        elif change_type == 'added':
            row_class = 'added'
            old_style = 'background: #f8f9fa; color: #999;'
            new_style = 'background: #d4edda;'
        else:
            row_class = 'deleted'
            old_style = 'background: #f8d7da;'
            new_style = 'background: #f8f9fa; color: #999;'

        if old_art:
            old_number = old_art.get('번호', '')
            old_content = old_art.get('내용', '')
            old_cell = f"<strong>{old_number}</strong> {old_content}"
        else:
            old_cell = "<em>(신설)</em>"

        if new_art:
            new_number = new_art.get('번호', '')
            new_content = new_art.get('내용', '')
            new_cell = f"<strong>{new_number}</strong> {new_content}"
        else:
            new_cell = "<em>(삭제)</em>"

        rows_html += f"""
        <tr class="{row_class}">
            <td style="{old_style} padding: 12px; border: 1px solid #dee2e6; vertical-align: top; width: 50%;">{old_cell}</td>
            <td style="{new_style} padding: 12px; border: 1px solid #dee2e6; vertical-align: top; width: 50%;">{new_cell}</td>
        </tr>
        """

    stats = {
        'modified': len([c for c in changes if c['type'] == 'modified']),
        'added': len([c for c in changes if c['type'] == 'added']),
        'deleted': len([c for c in changes if c['type'] == 'deleted'])
    }

    html = f"""
    <!DOCTYPE html>
    <html lang="ko">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>신구대비표 - {new_rule['wzname']}</title>
        <style>
            * {{ margin: 0; padding: 0; box-sizing: border-box; }}
            body {{ font-family: 'KB금융체Text', 'Malgun Gothic', sans-serif; padding: 20px; background: #f5f5f5; }}
            .container {{ max-width: 1200px; margin: 0 auto; background: #fff; padding: 30px; border-radius: 8px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }}
            h1 {{ text-align: center; margin-bottom: 10px; color: #333; font-size: 24px; }}
            .subtitle {{ text-align: center; color: #666; margin-bottom: 30px; font-size: 14px; }}
            .info-box {{ display: flex; justify-content: space-between; margin-bottom: 20px; padding: 15px; background: #f8f9fa; border-radius: 6px; }}
            .info-item {{ flex: 1; text-align: center; }}
            .info-label {{ font-size: 12px; color: #666; }}
            .info-value {{ font-size: 14px; font-weight: bold; color: #333; margin-top: 5px; }}
            .stats {{ display: flex; gap: 20px; justify-content: center; margin-bottom: 20px; }}
            .stat-item {{ padding: 10px 20px; border-radius: 4px; font-size: 13px; }}
            .stat-modified {{ background: #fff3cd; color: #856404; }}
            .stat-added {{ background: #d4edda; color: #155724; }}
            .stat-deleted {{ background: #f8d7da; color: #721c24; }}
            table {{ width: 100%; border-collapse: collapse; }}
            th {{ background: #343a40; color: #fff; padding: 15px; text-align: center; font-size: 15px; }}
            .legend {{ margin-top: 20px; padding: 15px; background: #f8f9fa; border-radius: 6px; font-size: 12px; }}
            .legend-item {{ display: inline-block; margin-right: 20px; }}
            .legend-color {{ display: inline-block; width: 16px; height: 16px; vertical-align: middle; margin-right: 5px; border-radius: 2px; }}
            @media print {{
                body {{ background: #fff; padding: 0; }}
                .container {{ box-shadow: none; padding: 20px; }}
                .no-print {{ display: none; }}
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <h1>신구대비표</h1>
            <p class="subtitle">{new_rule['wzpubno']} {new_rule['wzname']}</p>

            <div class="info-box">
                <div class="info-item">
                    <div class="info-label">구 규정</div>
                    <div class="info-value">{old_rule['wznewflag']} ({old_rule['wzlastrevdate'] or '-'})</div>
                </div>
                <div class="info-item">
                    <div class="info-label">신 규정</div>
                    <div class="info-value">{new_rule['wznewflag']} ({new_rule['wzlastrevdate'] or '-'})</div>
                </div>
            </div>

            <div class="stats">
                <span class="stat-item stat-modified">수정 {stats['modified']}건</span>
                <span class="stat-item stat-added">추가 {stats['added']}건</span>
                <span class="stat-item stat-deleted">삭제 {stats['deleted']}건</span>
            </div>

            <table>
                <thead>
                    <tr>
                        <th style="width: 50%;">구 조문</th>
                        <th style="width: 50%;">신 조문</th>
                    </tr>
                </thead>
                <tbody>
                    {rows_html if rows_html else '<tr><td colspan="2" style="text-align: center; padding: 40px; color: #999;">변경 사항이 없습니다.</td></tr>'}
                </tbody>
            </table>

            <div class="legend">
                <span class="legend-item"><span class="legend-color" style="background: #fff3cd;"></span> 수정</span>
                <span class="legend-item"><span class="legend-color" style="background: #d4edda;"></span> 추가</span>
                <span class="legend-item"><span class="legend-color" style="background: #f8d7da;"></span> 삭제</span>
            </div>

            <p style="text-align: center; margin-top: 20px; color: #999; font-size: 12px;">
                생성일시: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
            </p>
        </div>
    </body>
    </html>
    """

    return html


@router.get("/list")
async def get_comparable_regulations(
    search: str = Query(default="", description="검색어"),
    user: Dict[str, Any] = Depends(get_current_user)
):
    """비교 가능한 규정 목록 조회 (연혁이 있는 규정만)"""
    try:
        with db_manager.get_connection() as conn:
            with conn.cursor() as cur:
                query = """
                    SELECT DISTINCT r1.wzruleseq, r1.wzruleid, r1.wzname, r1.wzpubno,
                           r1.wzestabdate, r1.wzlastrevdate,
                           (SELECT COUNT(*) FROM wz_rule r2
                            WHERE r2.wzruleid = r1.wzruleid AND r2.wznewflag = '연혁') as history_count
                    FROM wz_rule r1
                    WHERE r1.wznewflag = '현행'
                    AND EXISTS (
                        SELECT 1 FROM wz_rule r2
                        WHERE r2.wzruleid = r1.wzruleid AND r2.wznewflag = '연혁'
                    )
                """

                params = []
                if search:
                    query += " AND (r1.wzname LIKE %s OR r1.wzpubno LIKE %s)"
                    params.extend([f'%{search}%', f'%{search}%'])

                query += " ORDER BY r1.wzpubno"

                cur.execute(query, params)
                rows = cur.fetchall()

        regulations = []
        for row in rows:
            regulations.append({
                'wzruleseq': row[0],
                'wzruleid': row[1],
                'wzname': row[2],
                'wzpubno': row[3],
                'wzestabdate': str(row[4]) if row[4] else None,
                'wzlastrevdate': str(row[5]) if row[5] else None,
                'history_count': row[6]
            })

        return JSONResponse(content={
            "success": True,
            "regulations": regulations,
            "total_count": len(regulations)
        })

    except Exception as e:
        logger.error(f"Error getting comparable regulations: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/current-list")
async def get_current_regulations(
    search: str = Query(default="", description="검색어"),
    user: Dict[str, Any] = Depends(get_current_user)
):
    """현행 규정 목록 조회 (파일 업로드 비교용)"""
    try:
        with db_manager.get_connection() as conn:
            with conn.cursor() as cur:
                query = """
                    SELECT wzruleseq, wzruleid, wzname, wzpubno,
                           wzestabdate, wzlastrevdate
                    FROM wz_rule
                    WHERE wznewflag = '현행'
                """

                params = []
                if search:
                    query += " AND (wzname LIKE %s OR wzpubno LIKE %s)"
                    params.extend([f'%{search}%', f'%{search}%'])

                query += " ORDER BY wzpubno"

                cur.execute(query, params)
                rows = cur.fetchall()

        regulations = []
        for row in rows:
            regulations.append({
                'wzruleseq': row[0],
                'wzruleid': row[1],
                'wzname': row[2],
                'wzpubno': row[3],
                'wzestabdate': str(row[4]) if row[4] else None,
                'wzlastrevdate': str(row[5]) if row[5] else None
            })

        return JSONResponse(content={
            "success": True,
            "regulations": regulations,
            "total_count": len(regulations)
        })

    except Exception as e:
        logger.error(f"Error getting current regulations: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/parsed-json-list/{rule_id}")
async def get_parsed_json_files(
    rule_id: int,
    user: Dict[str, Any] = Depends(get_current_user)
):
    """특정 규정의 파싱된 JSON 파일 목록 조회 (docx_json 폴더)"""
    try:
        rule_info = get_regulation_info(rule_id)
        if not rule_info:
            raise HTTPException(status_code=404, detail="규정을 찾을 수 없습니다.")

        wzpubno = rule_info.get('wzpubno', '').strip()
        if not wzpubno:
            raise HTTPException(status_code=400, detail="규정번호가 없습니다.")

        json_files = []

        if DOCX_JSON_DIR.exists():
            for json_path in DOCX_JSON_DIR.glob("*.json"):
                filename = json_path.name
                if filename.startswith(wzpubno):
                    stat = json_path.stat()
                    json_files.append({
                        'filename': filename,
                        'filepath': str(json_path),
                        'size': stat.st_size,
                        'modified': datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M:%S')
                    })

        json_files.sort(key=lambda x: x['modified'], reverse=True)

        return JSONResponse(content={
            "success": True,
            "rule_id": rule_id,
            "wzpubno": wzpubno,
            "wzname": rule_info.get('wzname'),
            "files": json_files,
            "total_count": len(json_files)
        })

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting parsed JSON files: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/compare-with-json")
async def compare_with_parsed_json(
    rule_id: int = Query(..., description="현행 규정 시퀀스"),
    json_filename: str = Query(..., description="비교할 JSON 파일명"),
    user: Dict[str, Any] = Depends(get_current_user)
):
    """현행 규정과 파싱된 JSON 파일 비교"""
    try:
        current_rule = get_regulation_info(rule_id)
        if not current_rule:
            raise HTTPException(status_code=404, detail=f"규정을 찾을 수 없습니다: {rule_id}")

        current_json = load_json_content(current_rule)
        if not current_json:
            raise HTTPException(status_code=404, detail="현행 규정의 JSON 파일을 찾을 수 없습니다.")

        json_path = DOCX_JSON_DIR / json_filename
        if not json_path.exists():
            raise HTTPException(status_code=404, detail=f"JSON 파일을 찾을 수 없습니다: {json_filename}")

        with open(json_path, 'r', encoding='utf-8') as f:
            new_json = json.load(f)

        old_articles = current_json.get('조문내용', [])
        new_articles = new_json.get('조문내용', [])

        changes = compare_articles(old_articles, new_articles)

        stats = {
            'total_changes': len(changes),
            'modified': len([c for c in changes if c['type'] == 'modified']),
            'added': len([c for c in changes if c['type'] == 'added']),
            'deleted': len([c for c in changes if c['type'] == 'deleted'])
        }

        return JSONResponse(content={
            "success": True,
            "old_rule": {
                "wzruleseq": current_rule['wzruleseq'],
                "wzname": current_rule['wzname'],
                "wzpubno": current_rule['wzpubno'],
                "wznewflag": current_rule['wznewflag'],
                "wzlastrevdate": str(current_rule['wzlastrevdate']) if current_rule['wzlastrevdate'] else None
            },
            "new_rule": {
                "wzruleseq": None,
                "wzname": current_rule['wzname'],
                "wzpubno": current_rule['wzpubno'],
                "wznewflag": "개정안",
                "wzlastrevdate": None,
                "filename": json_filename
            },
            "changes": changes,
            "stats": stats
        })

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error comparing with JSON: {e}")
        raise HTTPException(status_code=500, detail=str(e))


def parse_pdf_to_json(file_content: bytes, filename: str) -> Dict:
    """PDF 파일을 JSON으로 파싱 (pdf2txt + txt2json)"""
    try:
        import importlib.util

        applib_path = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'applib'))

        pdf2txt_spec = importlib.util.spec_from_file_location(
            "pdf2txt_applib",
            os.path.join(applib_path, "pdf2txt.py")
        )
        pdf2txt_module = importlib.util.module_from_spec(pdf2txt_spec)
        pdf2txt_spec.loader.exec_module(pdf2txt_module)

        txt2json_spec = importlib.util.spec_from_file_location(
            "txt2json_applib",
            os.path.join(applib_path, "txt2json.py")
        )
        txt2json_module = importlib.util.module_from_spec(txt2json_spec)
        txt2json_spec.loader.exec_module(txt2json_module)

        extract_text_from_pdf_bytes = pdf2txt_module.extract_text_from_pdf_bytes
        convert_txt_to_json = txt2json_module.convert_txt_to_json

        logger.info(f"[COMPARE] Parsing PDF file: {filename}")

        text_content = extract_text_from_pdf_bytes(file_content)

        if not text_content:
            return {'success': False, 'error': 'PDF에서 텍스트를 추출할 수 없습니다.'}

        json_data = convert_txt_to_json(text_content)

        if json_data:
            if 'sections' in json_data and '조문내용' not in json_data:
                json_data['조문내용'] = json_data['sections']

            return {
                'success': True,
                'data': json_data,
                'articles': json_data.get('조문내용', [])
            }
        else:
            return {'success': False, 'error': 'JSON 변환 실패'}

    except ImportError as e:
        logger.warning(f"PDF parsing modules not available: {e}")
        return {'success': False, 'error': f'PDF 파싱 모듈을 불러올 수 없습니다: {e}'}
    except Exception as e:
        logger.error(f"Error parsing PDF: {e}")
        return {'success': False, 'error': str(e)}


def correct_number_format(number: str, level: int) -> str:
    """레벨에 따라 번호 형식을 보정"""
    if not number:
        return number

    num_match = re.search(r'\d+', number)
    if not num_match:
        return number

    num = num_match.group()

    if level == 1:
        if re.match(r'^제\d+조', number):
            return number
        return f"제{num}조"
    elif level == 2:
        if re.match(r'^\d+\.$', number):
            return number
        if re.match(r'^\d+\)', number):
            return f"{num}."
        return f"{num}."
    elif level == 3:
        if re.match(r'^\d+\)$', number):
            return number
        if re.match(r'^\(\d+\)', number):
            return f"{num})"
        return f"{num})"
    elif level == 4:
        if re.match(r'^\(\d+\)$', number):
            return number
        return f"({num})"

    return number


def merge_pdf_docx_data(pdf_data: Dict, docx_data: Dict) -> Dict:
    """PDF와 DOCX 파싱 결과를 병합"""
    try:
        pdf_articles = pdf_data.get('articles', []) or pdf_data.get('data', {}).get('조문내용', [])
        docx_articles = docx_data.get('articles', []) or docx_data.get('data', {}).get('조문내용', [])

        logger.info(f"[MERGE] PDF articles: {len(pdf_articles)}, DOCX articles: {len(docx_articles)}")

        pdf_map = {article.get('seq', i): article for i, article in enumerate(pdf_articles)}
        docx_map = {article.get('seq', i): article for i, article in enumerate(docx_articles)}

        all_seqs = set(pdf_map.keys()) | set(docx_map.keys())

        merged_articles = []
        for seq in sorted(all_seqs):
            pdf_article = pdf_map.get(seq, {})
            docx_article = docx_map.get(seq, {})

            level = max(pdf_article.get('레벨', 0), docx_article.get('레벨', 0))

            raw_number = pdf_article.get('번호', '') or docx_article.get('번호', '')
            corrected_number = correct_number_format(raw_number, level) if level > 0 else raw_number

            merged_item = {
                'seq': seq,
                '레벨': level,
                '번호': corrected_number,
                '내용': docx_article.get('내용', '') or pdf_article.get('내용', ''),
                '관련이미지': pdf_article.get('관련이미지', []) or docx_article.get('관련이미지', [])
            }

            merged_articles.append(merged_item)

        merged_doc_info = {}
        pdf_doc_info = pdf_data.get('data', {}).get('문서정보', {})
        docx_doc_info = docx_data.get('data', {}).get('문서정보', {})

        merged_doc_info.update(pdf_doc_info)
        merged_doc_info.update(docx_doc_info)

        logger.info(f"[MERGE] Merged articles: {len(merged_articles)}")

        return {
            'success': True,
            'data': {
                '문서정보': merged_doc_info,
                '조문내용': merged_articles
            },
            'articles': merged_articles
        }

    except Exception as e:
        logger.error(f"Error merging PDF and DOCX data: {e}")
        return {'success': False, 'error': str(e)}


def parse_docx_to_json(file_content: bytes, filename: str, wzruleid: int) -> Dict:
    """DOCX 파일을 JSON으로 파싱"""
    try:
        applib_path = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'applib'))
        if applib_path not in sys.path:
            sys.path.insert(0, applib_path)

        from docx2json import process_docx_file

        logger.info(f"[COMPARE] Parsing DOCX file: {filename}")

        parse_result = process_docx_file(
            file_content=file_content,
            filename=filename,
            wzruleid=wzruleid
        )

        if parse_result.get('success'):
            sections = parse_result.get('sections', [])
            logger.info(f"[COMPARE] DOCX parsed successfully, sections: {len(sections)}")
            return {
                'success': True,
                'data': {
                    '문서정보': parse_result.get('document_info', {}),
                    '조문내용': sections
                },
                'articles': sections
            }
        else:
            return {
                'success': False,
                'error': parse_result.get('error', 'Unknown error')
            }

    except ImportError as e:
        logger.error(f"Failed to import docx2json: {e}")
        return {'success': False, 'error': f'docx2json 모듈을 불러올 수 없습니다: {e}'}
    except Exception as e:
        logger.error(f"Error parsing DOCX: {e}")
        return {'success': False, 'error': str(e)}


@router.post("/upload-compare")
async def upload_and_compare(
    rule_id: int = Form(..., description="현행 규정 시퀀스 (wzruleseq)"),
    pdf_file: UploadFile = File(..., description="비교할 PDF 파일 (필수)"),
    docx_file: UploadFile = File(..., description="비교할 DOCX 파일 (필수)"),
    user: Dict[str, Any] = Depends(get_current_user)
):
    """현행 규정과 업로드된 파일 비교 (PDF+DOCX 병합)"""
    try:
        if not pdf_file.filename.lower().endswith('.pdf'):
            raise HTTPException(status_code=400, detail="PDF 파일만 업로드 가능합니다.")

        if not docx_file.filename.lower().endswith('.docx'):
            raise HTTPException(status_code=400, detail="DOCX 파일만 업로드 가능합니다.")

        current_rule = get_regulation_info(rule_id)
        if not current_rule:
            raise HTTPException(status_code=404, detail=f"규정을 찾을 수 없습니다: {rule_id}")

        current_json = load_json_content(current_rule)
        if not current_json:
            raise HTTPException(status_code=404, detail="현행 규정의 JSON 파일을 찾을 수 없습니다.")

        docx_content = await docx_file.read()
        docx_result = parse_docx_to_json(
            file_content=docx_content,
            filename=docx_file.filename,
            wzruleid=current_rule['wzruleid']
        )

        if not docx_result.get('success'):
            raise HTTPException(
                status_code=400,
                detail=f"DOCX 파싱 실패: {docx_result.get('error')}"
            )

        pdf_content = await pdf_file.read()
        pdf_result = parse_pdf_to_json(
            file_content=pdf_content,
            filename=pdf_file.filename
        )

        if not pdf_result.get('success'):
            raise HTTPException(
                status_code=400,
                detail=f"PDF 파싱 실패: {pdf_result.get('error')}"
            )

        merged_result = merge_pdf_docx_data(pdf_result, docx_result)
        if not merged_result.get('success'):
            raise HTTPException(
                status_code=400,
                detail=f"PDF+DOCX 병합 실패: {merged_result.get('error')}"
            )

        new_articles = merged_result.get('articles', [])
        parse_mode = "PDF+DOCX 병합"
        logger.info(f"[COMPARE] Merged PDF+DOCX: {len(new_articles)} articles")

        old_articles = current_json.get('조문내용', [])
        changes = compare_articles(old_articles, new_articles)

        stats = {
            'total_changes': len(changes),
            'modified': len([c for c in changes if c['type'] == 'modified']),
            'added': len([c for c in changes if c['type'] == 'added']),
            'deleted': len([c for c in changes if c['type'] == 'deleted'])
        }

        filenames = f"{pdf_file.filename}, {docx_file.filename}"

        return JSONResponse(content={
            "success": True,
            "old_rule": {
                "wzruleseq": current_rule['wzruleseq'],
                "wzname": current_rule['wzname'],
                "wzpubno": current_rule['wzpubno'],
                "wznewflag": current_rule['wznewflag'],
                "wzlastrevdate": str(current_rule['wzlastrevdate']) if current_rule['wzlastrevdate'] else None
            },
            "new_rule": {
                "wzruleseq": None,
                "wzname": current_rule['wzname'],
                "wzpubno": current_rule['wzpubno'],
                "wznewflag": f"업로드 ({parse_mode})",
                "wzlastrevdate": datetime.now().strftime('%Y-%m-%d'),
                "filename": filenames
            },
            "changes": changes,
            "stats": stats,
            "parse_mode": parse_mode
        })

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error in upload and compare: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# 신구대비표 PDF 저장 경로
COMPARISON_TABLE_DIR = Path("/home/wizice/kbregulation/www/static/pdf/comparisonTable")
COMPARISON_TABLE_BACKUP_DIR = Path("/home/wizice/kbregulation/www/static/pdf/comparisonTable_backup")


def _get_kb_font_for_pdf():
    """PDF 생성용 KB금융체 폰트 등록 및 이름 반환"""
    font_name = 'Helvetica'
    font_name_bold = 'Helvetica-Bold'

    font_paths = [
        os.path.expanduser('~/.local/share/fonts/kb-finance/KBfgTextM.ttf'),
        '/usr/share/fonts/truetype/nanum/NanumGothic.ttf',
        '/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf'
    ]
    bold_paths = [
        os.path.expanduser('~/.local/share/fonts/kb-finance/KBfgTextB.ttf'),
        '/usr/share/fonts/truetype/nanum/NanumGothicBold.ttf',
        '/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf'
    ]

    for i, font_path in enumerate(font_paths):
        if os.path.exists(font_path):
            try:
                pdfmetrics.registerFont(TTFont('PDFKoreanFont', font_path))
                font_name = 'PDFKoreanFont'
                bold_path = bold_paths[i] if i < len(bold_paths) else font_path
                if os.path.exists(bold_path):
                    pdfmetrics.registerFont(TTFont('PDFKoreanFontBold', bold_path))
                    font_name_bold = 'PDFKoreanFontBold'
                else:
                    font_name_bold = 'PDFKoreanFont'
                logger.info(f"PDF 폰트 등록 성공: {font_path}")
                break
            except Exception as e:
                logger.warning(f"폰트 등록 실패 {font_path}: {e}")

    if font_name == 'Helvetica':
        logger.warning("한글 폰트를 찾을 수 없어 기본 폰트 사용")

    return font_name, font_name_bold


def generate_comparison_pdf(old_rule: Dict, new_rule: Dict, changes: List[Dict], output_path: str, remarks: str = "") -> bool:
    """KB신용정보 사규 개정(안) 신구대비표 PDF 생성 (reportlab 직접)"""
    from reportlab.lib.pagesizes import A4, landscape

    try:
        font_name, font_name_bold = _get_kb_font_for_pdf()

        page_width, page_height = landscape(A4)

        doc = SimpleDocTemplate(
            output_path,
            pagesize=landscape(A4),
            rightMargin=15*mm,
            leftMargin=15*mm,
            topMargin=15*mm,
            bottomMargin=15*mm
        )

        elements = []
        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            'Title',
            parent=styles['Title'],
            fontName=font_name_bold,
            fontSize=16,
            alignment=1,
            spaceAfter=8*mm
        )

        header_style = ParagraphStyle(
            'Header',
            parent=styles['Normal'],
            fontName=font_name_bold,
            fontSize=10,
            alignment=1,
            leading=14
        )

        cell_style = ParagraphStyle(
            'Cell',
            parent=styles['Normal'],
            fontName=font_name,
            fontSize=9,
            leading=13,
            wordWrap='CJK'
        )

        elements.append(Paragraph("KB신용정보 사규 개정(안) 신구대비표", title_style))

        table_data = []

        table_data.append([
            Paragraph("<b>사규명</b>", header_style),
            Paragraph("<b>현행</b>", header_style),
            Paragraph("<b>개정</b>", header_style),
            Paragraph("<b>비고</b>", header_style)
        ])

        rule_name = f"{new_rule.get('wzpubno', '')} {new_rule.get('wzname', '')}"

        first_row = True
        for change in changes:
            change_type = change['type']
            old_art = change.get('old')
            new_art = change.get('new')
            level = change.get('level', 2)

            old_number = old_art.get('번호', '') if old_art else ''
            old_content = old_art.get('내용', '') if old_art else ''
            old_number = correct_number_format(old_number, level) if old_number else ''
            old_full = f"{old_number} {old_content}".strip()

            new_number = new_art.get('번호', '') if new_art else ''
            new_content = new_art.get('내용', '') if new_art else ''
            new_number = correct_number_format(new_number, level) if new_number else ''
            new_full = f"{new_number} {new_content}".strip()

            old_clean = _strip_html_tags(old_full) if old_full else ""
            new_clean = _strip_html_tags(new_full) if new_full else ""

            if change_type == 'deleted':
                old_text = f"<font face='{font_name}'><strike>{old_clean}</strike></font>"
                new_text = ""
            elif change_type == 'added':
                old_text = ""
                new_text = f"<u><font face='{font_name}' color='blue'>{new_clean}</font></u>"
            elif change_type == 'modified':
                old_text, new_text = highlight_diff(old_full, new_full, font_name)
            else:
                old_text = old_clean
                new_text = new_clean

            if first_row and remarks:
                remark_text = remarks
                first_row = False
            else:
                remark_text = ""

            if len(table_data) == 1:
                rule_cell = Paragraph(rule_name, cell_style)
            else:
                rule_cell = Paragraph("", cell_style)

            table_data.append([
                rule_cell,
                Paragraph(old_text, cell_style),
                Paragraph(new_text, cell_style),
                Paragraph(remark_text, cell_style)
            ])

        if len(table_data) == 1:
            table_data.append([
                Paragraph(rule_name, cell_style),
                Paragraph("변경 사항 없음", cell_style),
                Paragraph("변경 사항 없음", cell_style),
                Paragraph("", cell_style)
            ])

        available_width = page_width - 30*mm
        col_widths = [
            available_width * 0.12,
            available_width * 0.38,
            available_width * 0.38,
            available_width * 0.12
        ]

        table = Table(table_data, colWidths=col_widths, repeatRows=1)

        table_style = TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E0E0E0')),
            ('FONTNAME', (0, 0), (-1, 0), font_name_bold),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ])

        if len(table_data) > 2:
            table_style.add('SPAN', (0, 1), (0, len(table_data) - 1))

        if remarks and len(table_data) > 2:
            table_style.add('SPAN', (3, 1), (3, len(table_data) - 1))

        table.setStyle(table_style)
        elements.append(table)

        doc.build(elements)
        logger.info(f"PDF generated successfully: {output_path}")
        return True

    except Exception as e:
        logger.error(f"Error generating PDF: {e}", exc_info=True)
        return False


def generate_comparison_pdf_v2(old_rule: Dict, new_rule: Dict, old_articles: List[Dict], new_articles: List[Dict], output_path: str, remarks: str = "") -> bool:
    """
    KB신용정보 사규 개정(안) 신구대비표 PDF 생성 (개선 버전)
    DOCX를 먼저 생성한 후 LibreOffice를 사용하여 PDF로 변환합니다.
    """
    import subprocess

    try:
        pdf_dir = os.path.dirname(output_path)
        pdf_basename = os.path.basename(output_path)
        docx_basename = pdf_basename.replace('.pdf', '.docx')
        docx_path = os.path.join(pdf_dir, docx_basename)

        logger.info(f"Generating DOCX first: {docx_path}")
        docx_success = generate_comparison_docx(
            old_rule, new_rule,
            old_articles, new_articles,
            docx_path, remarks
        )

        if not docx_success:
            logger.error("Failed to generate DOCX for PDF conversion")
            return False

        logger.info(f"Converting DOCX to PDF using LibreOffice: {output_path}")
        try:
            result = subprocess.run(
                [
                    'libreoffice', '--headless', '--convert-to', 'pdf',
                    '--outdir', pdf_dir,
                    docx_path
                ],
                capture_output=True,
                text=True,
                timeout=60
            )

            if result.returncode != 0:
                logger.error(f"LibreOffice conversion failed: {result.stderr}")
                return False

            logger.info(f"PDF generated successfully via LibreOffice: {output_path}")
            return True

        except subprocess.TimeoutExpired:
            logger.error("LibreOffice conversion timed out")
            return False
        except FileNotFoundError:
            logger.error("LibreOffice not found. Please install libreoffice-writer.")
            return False

    except Exception as e:
        logger.error(f"Error generating PDF v2: {e}", exc_info=True)
        return False


def generate_comparison_docx(old_rule: Dict, new_rule: Dict, old_articles: List[Dict], new_articles: List[Dict], output_path: str, remarks: str = "") -> bool:
    """KB신용정보 사규 개정(안) 신구대비표 DOCX 생성"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.section import WD_ORIENT
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn, nsdecls
        from docx.oxml import parse_xml

        doc = Document()

        # KB금융체 사용 (시스템에 없으면 fallback)
        DOCX_FONT_NAME = 'KB금융체Text'
        style = doc.styles['Normal']
        style.font.name = DOCX_FONT_NAME
        style.font.size = Pt(9)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), DOCX_FONT_NAME)

        def set_run_font(run, font_size=Pt(9)):
            """run에 KB금융체 폰트 설정"""
            run.font.name = DOCX_FONT_NAME
            run.font.size = font_size
            r = run._element
            rPr = r.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                rFonts.set(qn('w:eastAsia'), DOCX_FONT_NAME)

        # 가로 방향 A4 설정
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

        # 제목
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run("KB신용정보 사규 개정(안) 신구대비표")
        title_run.bold = True
        set_run_font(title_run, Pt(16))

        # 규정 정보
        rule_pubno = new_rule.get('wzpubno', '') or old_rule.get('wzpubno', '')
        rule_name_full = new_rule.get('wzname', '') or old_rule.get('wzname', '')
        rule_name = re.sub(r'^[\d.]+\s*', '', rule_name_full).strip()

        # "제·개정 이력" 이후 조문 제외
        def filter_articles_before_history(articles):
            filtered = []
            for article in articles:
                content = article.get('내용', '')
                if '제·개정 이력' in content or '제개정 이력' in content or '제·개정이력' in content:
                    break
                filtered.append(article)
            return filtered

        old_articles_filtered = filter_articles_before_history(old_articles)
        new_articles_filtered = filter_articles_before_history(new_articles)

        # 다중 레벨 계층 기반 비교 수행
        hierarchy_changes = compare_articles_with_hierarchy(old_articles_filtered, new_articles_filtered)

        synchronized_rows = []

        for change in hierarchy_changes:
            change_type = change['type']
            old_art = change.get('old')
            new_art = change.get('new')
            level = change.get('level', 2)

            old_number = old_art.get('번호', '') if old_art else ''
            old_content = old_art.get('내용', '') if old_art else ''
            old_number = correct_number_format(old_number, level) if old_number else ''
            old_full = f"{old_number} {old_content}".strip()

            new_number = new_art.get('번호', '') if new_art else ''
            new_content = new_art.get('내용', '') if new_art else ''
            new_number = correct_number_format(new_number, level) if new_number else ''
            new_full = f"{new_number} {new_content}".strip()

            old_clean = _strip_html_tags(old_full) if old_full else ""
            new_clean = _strip_html_tags(new_full) if new_full else ""

            if change_type == 'context':
                is_bold = (level == 1)
                old_data = ([('normal', old_clean)], level, is_bold) if old_clean else None
                new_data = ([('normal', new_clean)], level, is_bold) if new_clean else None
                synchronized_rows.append((old_data, new_data))
            elif change_type == 'deleted':
                old_data = ([('deleted', old_clean)], level, False)
                synchronized_rows.append((old_data, None))
            elif change_type == 'added':
                old_data = ([('normal', '<신설>')], level, False)
                new_data = ([('added', new_clean)], level, False)
                synchronized_rows.append((old_data, new_data))
            elif change_type == 'modified':
                old_plain = _strip_html_tags(old_full) if old_full else ""
                new_plain = _strip_html_tags(new_full) if new_full else ""
                old_words = old_plain.split()
                new_words = new_plain.split()
                matcher = SequenceMatcher(None, old_words, new_words)
                opcodes = matcher.get_opcodes()

                old_segments = []
                new_segments = []
                for tag, i1, i2, j1, j2 in opcodes:
                    old_chunk = ' '.join(old_words[i1:i2])
                    new_chunk = ' '.join(new_words[j1:j2])
                    if tag == 'equal':
                        if old_chunk:
                            old_segments.append(('normal', old_chunk))
                        if new_chunk:
                            new_segments.append(('normal', new_chunk))
                    elif tag == 'delete':
                        if old_chunk:
                            old_segments.append(('deleted', old_chunk))
                    elif tag == 'insert':
                        if new_chunk:
                            new_segments.append(('added', new_chunk))
                    elif tag == 'replace':
                        if old_chunk:
                            old_segments.append(('deleted', old_chunk))
                        if new_chunk:
                            new_segments.append(('added', new_chunk))

                old_data = (old_segments, level, False) if old_segments else None
                new_data = (new_segments, level, False) if new_segments else None
                synchronized_rows.append((old_data, new_data))
            else:
                old_data = ([('normal', old_clean)], level, False) if old_clean else None
                new_data = ([('normal', new_clean)], level, False) if new_clean else None
                synchronized_rows.append((old_data, new_data))

        # 테이블 생성 (5열)
        table = doc.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        # 테이블 수준에서 내부 가로선 제거
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        if tbl.tblPr is None:
            tbl.insert(0, tblPr)

        ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        for existing in list(tblPr):
            if existing.tag == f'{ns}tblBorders':
                tblPr.remove(existing)

        tblBorders = parse_xml(f'''<w:tblBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:insideH w:val="nil"/>
            <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>
        </w:tblBorders>''')
        tblPr.append(tblBorders)

        # 열 너비 설정 (가로 A4 = 29.7cm - 3cm 마진 = 26.7cm)
        available_cm = 26.7
        widths = [available_cm * 0.07, available_cm * 0.09, available_cm * 0.355, available_cm * 0.355, available_cm * 0.13]
        for i, width in enumerate(widths):
            table.columns[i].width = Cm(width)

        # 헤더 행
        header_cells = table.rows[0].cells
        headers = ['사규번호', '사규제목', '현행', '개정', '비고']
        for i, header_text in enumerate(headers):
            cell = header_cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(header_text)
            run.bold = True
            set_run_font(run, Pt(10))
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E0E0E0"/>')
            cell._tc.get_or_add_tcPr().append(shading)

        # 레벨별 들여쓰기 매핑 (cm)
        level_indent = {1: 0, 2: 0.4, 3: 0.8, 4: 1.2, 5: 1.6, 6: 2.0}

        def add_segments_to_cell(cell, line_data):
            """셀에 단일 라인 서식 적용된 내용 추가"""
            if line_data is None:
                return

            segments, level, is_bold = line_data

            para = cell.paragraphs[0]
            para.clear()

            indent = level_indent.get(level, 0)
            if indent > 0:
                para.paragraph_format.left_indent = Cm(indent)

            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.space_before = Pt(0)

            for seg_type, text in segments:
                run = para.add_run(text)
                set_run_font(run, Pt(9))

                if is_bold:
                    run.bold = True

                if seg_type == 'deleted':
                    run.font.strike = True
                    run.font.color.rgb = RGBColor(0, 0, 255)
                elif seg_type == 'added':
                    run.underline = True
                    run.font.color.rgb = RGBColor(0, 0, 255)

                space_run = para.add_run(' ')
                set_run_font(space_run, Pt(9))

        # 데이터 행 추가
        if synchronized_rows:
            first_data_row = True
            for old_data, new_data in synchronized_rows:
                row = table.add_row()
                cells = row.cells

                if first_data_row:
                    p = cells[0].paragraphs[0]
                    run = p.add_run(rule_pubno)
                    set_run_font(run, Pt(9))

                    p = cells[1].paragraphs[0]
                    run = p.add_run(rule_name)
                    set_run_font(run, Pt(9))

                    p = cells[4].paragraphs[0]
                    run = p.add_run(remarks if remarks else "")
                    set_run_font(run, Pt(9))

                    first_data_row = False

                add_segments_to_cell(cells[2], old_data)
                add_segments_to_cell(cells[3], new_data)

            # 모든 열의 내부 가로선 제거
            if len(synchronized_rows) > 1:
                ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

                total_rows = len(table.rows)
                for row_idx in range(1, total_rows):
                    for col_idx in [0, 1, 2, 3, 4]:
                        cell = table.rows[row_idx].cells[col_idx]
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()

                        for existing in list(tcPr):
                            if existing.tag == f'{ns}tcBorders':
                                tcPr.remove(existing)

                        if row_idx == 1 and total_rows > 2:
                            borders = parse_xml(f'''<w:tcBorders {nsdecls("w")}>
                                <w:bottom w:val="nil"/>
                            </w:tcBorders>''')
                            tcPr.append(borders)
                        elif row_idx == total_rows - 1:
                            borders = parse_xml(f'''<w:tcBorders {nsdecls("w")}>
                                <w:top w:val="nil"/>
                            </w:tcBorders>''')
                            tcPr.append(borders)
                        elif row_idx > 1 and row_idx < total_rows - 1:
                            borders = parse_xml(f'''<w:tcBorders {nsdecls("w")}>
                                <w:top w:val="nil"/>
                                <w:bottom w:val="nil"/>
                            </w:tcBorders>''')
                            tcPr.append(borders)

                for row_idx in range(1, total_rows):
                    for col_idx in [0, 1, 2, 3, 4]:
                        cell = table.rows[row_idx].cells[col_idx]
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()

                        for existing in list(tcPr):
                            if existing.tag == f'{ns}vAlign':
                                tcPr.remove(existing)

                        vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="top"/>')
                        tcPr.append(vAlign)
        else:
            row = table.add_row()
            cells = row.cells
            for idx, txt in enumerate([rule_pubno, rule_name, "변경 사항 없음", "변경 사항 없음", ""]):
                p = cells[idx].paragraphs[0]
                run = p.add_run(txt)
                set_run_font(run, Pt(9))

        # 셀 너비 강제 적용
        ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        processed_tcs = set()
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                tc = cell._tc
                tc_id = id(tc)
                if tc_id in processed_tcs:
                    continue
                processed_tcs.add(tc_id)

                tcPr = tc.get_or_add_tcPr()
                for existing in list(tcPr):
                    if existing.tag == f'{ns}tcW':
                        tcPr.remove(existing)
                tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{int(widths[i] * 567)}" w:type="dxa"/>')
                tcPr.append(tcW)

        doc.save(output_path)
        logger.info(f"DOCX generated successfully: {output_path}")
        return True

    except Exception as e:
        logger.error(f"Error generating DOCX: {e}", exc_info=True)
        return False


@router.post("/save-comparison")
async def save_comparison_table(
    rule_id: int = Form(..., description="현행 규정 시퀀스 (wzruleseq)"),
    revision_date: str = Form(..., description="개정일 (YYYYMMDD 형식)"),
    remarks: str = Form(default="", description="비고 (개정 사유)"),
    pdf_file: UploadFile = File(..., description="비교할 PDF 파일 (필수)"),
    docx_file: UploadFile = File(..., description="비교할 DOCX 파일 (필수)"),
    user: Dict[str, Any] = Depends(get_current_user)
):
    """신구대비표 PDF 생성 및 저장"""
    try:
        if not pdf_file.filename.lower().endswith('.pdf'):
            raise HTTPException(status_code=400, detail="PDF 파일만 업로드 가능합니다.")

        if not docx_file.filename.lower().endswith('.docx'):
            raise HTTPException(status_code=400, detail="DOCX 파일만 업로드 가능합니다.")

        if not re.match(r'^\d{8}$', revision_date):
            raise HTTPException(status_code=400, detail="개정일은 YYYYMMDD 형식이어야 합니다.")

        current_rule = get_regulation_info(rule_id)
        if not current_rule:
            raise HTTPException(status_code=404, detail=f"규정을 찾을 수 없습니다: {rule_id}")

        current_json = load_json_content(current_rule)
        if not current_json:
            raise HTTPException(status_code=404, detail="현행 규정의 JSON 파일을 찾을 수 없습니다.")

        docx_content = await docx_file.read()
        docx_result = parse_docx_to_json(
            file_content=docx_content,
            filename=docx_file.filename,
            wzruleid=current_rule['wzruleid']
        )

        if not docx_result.get('success'):
            raise HTTPException(
                status_code=400,
                detail=f"DOCX 파싱 실패: {docx_result.get('error')}"
            )

        pdf_content = await pdf_file.read()
        pdf_result = parse_pdf_to_json(
            file_content=pdf_content,
            filename=pdf_file.filename
        )

        if not pdf_result.get('success'):
            raise HTTPException(
                status_code=400,
                detail=f"PDF 파싱 실패: {pdf_result.get('error')}"
            )

        merged_result = merge_pdf_docx_data(pdf_result, docx_result)
        if not merged_result.get('success'):
            raise HTTPException(
                status_code=400,
                detail=f"PDF+DOCX 병합 실패: {merged_result.get('error')}"
            )

        new_articles = merged_result.get('articles', [])
        logger.info(f"[SAVE-COMPARE] Merged PDF+DOCX: {len(new_articles)} articles")

        old_articles = current_json.get('조문내용', [])
        changes = compare_articles(old_articles, new_articles)

        new_rule = {
            'wzruleseq': None,
            'wzname': current_rule['wzname'],
            'wzpubno': current_rule['wzpubno'],
            'wznewflag': "개정안",
            'wzlastrevdate': f"{revision_date[:4]}-{revision_date[4:6]}-{revision_date[6:8]}"
        }

        COMPARISON_TABLE_DIR.mkdir(parents=True, exist_ok=True)

        filename = f"{current_rule['wzruleid']}_{rule_id}_{revision_date}.pdf"
        save_path = COMPARISON_TABLE_DIR / filename

        if save_path.exists():
            COMPARISON_TABLE_BACKUP_DIR.mkdir(parents=True, exist_ok=True)
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            backup_path = COMPARISON_TABLE_BACKUP_DIR / f"{current_rule['wzruleid']}_{rule_id}_{revision_date}_{timestamp}.pdf"
            shutil.copy2(save_path, backup_path)
            logger.info(f"Backed up existing file to: {backup_path}")

        if not generate_comparison_pdf_v2(current_rule, new_rule, old_articles, new_articles, str(save_path), remarks):
            raise HTTPException(status_code=500, detail="PDF 생성 실패")

        docx_filename = f"{current_rule['wzruleid']}_{rule_id}_{revision_date}.docx"
        docx_save_path = COMPARISON_TABLE_DIR / docx_filename
        docx_relative_path = None
        if generate_comparison_docx(current_rule, new_rule, old_articles, new_articles, str(docx_save_path), remarks):
            docx_relative_path = f"comparisonTable/{docx_filename}"
            logger.info(f"DOCX comparison table generated: {docx_relative_path}")
        else:
            logger.warning("DOCX 신구대비표 생성 실패 (PDF는 정상 생성됨)")

        relative_path = f"comparisonTable/{filename}"

        with db_manager.get_connection() as conn:
            with conn.cursor() as cur:
                cur.execute("""
                    UPDATE wz_rule
                    SET wzFileComparison = %s,
                        wzFileComparisonDocx = %s
                    WHERE wzruleseq = %s
                """, (relative_path, docx_relative_path, rule_id))
                conn.commit()

        logger.info(f"Saved comparison table: {relative_path}")

        stats = {
            'total_changes': len(changes),
            'modified': len([c for c in changes if c['type'] == 'modified']),
            'added': len([c for c in changes if c['type'] == 'added']),
            'deleted': len([c for c in changes if c['type'] == 'deleted'])
        }

        return JSONResponse(content={
            "success": True,
            "message": "신구대비표가 저장되었습니다.",
            "file_path": relative_path,
            "filename": filename,
            "rule_id": rule_id,
            "wzruleid": current_rule['wzruleid'],
            "stats": stats,
            "view_url": f"/static/pdf/{relative_path}"
        })

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error saving comparison table: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/comparison-file/{rule_id}")
async def get_comparison_file_info(
    rule_id: int,
    user: Dict[str, Any] = Depends(get_current_user)
):
    """규정의 신구대비표 파일 정보 조회"""
    try:
        with db_manager.get_connection() as conn:
            with conn.cursor() as cur:
                cur.execute("""
                    SELECT wzFileComparison, wzFileComparisonDocx, wzruleid, wzpubno, wzname, wzlastrevdate
                    FROM wz_rule
                    WHERE wzruleseq = %s
                """, (rule_id,))

                row = cur.fetchone()
                if not row:
                    raise HTTPException(status_code=404, detail="규정을 찾을 수 없습니다.")

                wzfilecomparison, wzfilecomparisondocx, wzruleid, wzpubno, wzname, wzlastrevdate = row

                if not wzfilecomparison and not wzfilecomparisondocx:
                    return JSONResponse(content={
                        "success": False,
                        "message": "신구대비표가 등록되지 않았습니다.",
                        "has_file": False,
                        "rule_id": rule_id,
                        "wzruleid": wzruleid,
                        "wzname": wzname
                    })

                pdf_exists = False
                docx_exists = False

                if wzfilecomparison:
                    pdf_path = Path("/home/wizice/kbregulation/www/static/pdf") / wzfilecomparison
                    pdf_exists = pdf_path.exists()

                if wzfilecomparisondocx:
                    docx_path = Path("/home/wizice/kbregulation/www/static/pdf") / wzfilecomparisondocx
                    docx_exists = docx_path.exists()

                return JSONResponse(content={
                    "success": True,
                    "has_file": True,
                    "file_path": wzfilecomparison,
                    "file_path_docx": wzfilecomparisondocx,
                    "file_exists": pdf_exists,
                    "docx_exists": docx_exists,
                    "view_url": f"/static/pdf/{wzfilecomparison}" if pdf_exists else None,
                    "docx_url": f"/static/pdf/{wzfilecomparisondocx}" if docx_exists else None,
                    "rule_id": rule_id,
                    "wzruleid": wzruleid,
                    "wzname": wzname,
                    "wzpubno": wzpubno,
                    "wzlastrevdate": str(wzlastrevdate) if wzlastrevdate else None
                })

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting comparison file info: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/download/{rule_id}")
async def download_comparison_pdf(
    rule_id: int,
    user: Dict[str, Any] = Depends(get_current_user)
):
    """신구대비표 PDF 다운로드"""
    try:
        with db_manager.get_connection() as conn:
            with conn.cursor() as cur:
                cur.execute("""
                    SELECT wzFileComparison, wzruleid, wzpubno, wzname
                    FROM wz_rule
                    WHERE wzruleseq = %s
                """, (rule_id,))

                row = cur.fetchone()
                if not row:
                    raise HTTPException(status_code=404, detail="규정을 찾을 수 없습니다.")

                wzfilecomparison, wzruleid, wzpubno, wzname = row

                if not wzfilecomparison:
                    raise HTTPException(status_code=404, detail="신구대비표가 등록되지 않았습니다.")

                file_path = Path("/home/wizice/kbregulation/www/static/pdf") / wzfilecomparison

                if not file_path.exists():
                    raise HTTPException(status_code=404, detail="신구대비표 파일을 찾을 수 없습니다.")

                download_filename = f"신구대비표_{wzpubno}_{wzname}.pdf"

                return FileResponse(
                    path=str(file_path),
                    filename=download_filename,
                    media_type="application/pdf"
                )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error downloading comparison PDF: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/generate-download")
async def generate_and_download_comparison(
    rule_id: int = Form(..., description="현행 규정 시퀀스 (wzruleseq)"),
    remarks: str = Form(default="", description="비고 (개정 사유)"),
    output_format: str = Form(default="pdf", description="출력 형식 (pdf 또는 docx)"),
    pdf_file: UploadFile = File(..., description="비교할 PDF 파일 (필수)"),
    docx_file: UploadFile = File(..., description="비교할 DOCX 파일 (필수)"),
    user: Dict[str, Any] = Depends(get_current_user)
):
    """신구대비표 PDF/DOCX 생성 후 바로 다운로드 (DB 저장 안함)"""
    try:
        if not pdf_file.filename.lower().endswith('.pdf'):
            raise HTTPException(status_code=400, detail="PDF 파일만 업로드 가능합니다.")

        if not docx_file.filename.lower().endswith('.docx'):
            raise HTTPException(status_code=400, detail="DOCX 파일만 업로드 가능합니다.")

        current_rule = get_regulation_info(rule_id)
        if not current_rule:
            raise HTTPException(status_code=404, detail=f"규정을 찾을 수 없습니다: {rule_id}")

        current_json = load_json_content(current_rule)
        if not current_json:
            raise HTTPException(status_code=404, detail="현행 규정의 JSON 파일을 찾을 수 없습니다.")

        docx_content = await docx_file.read()
        docx_result = parse_docx_to_json(
            file_content=docx_content,
            filename=docx_file.filename,
            wzruleid=current_rule['wzruleid']
        )

        if not docx_result.get('success'):
            raise HTTPException(
                status_code=400,
                detail=f"DOCX 파싱 실패: {docx_result.get('error')}"
            )

        pdf_content = await pdf_file.read()
        pdf_result = parse_pdf_to_json(
            file_content=pdf_content,
            filename=pdf_file.filename
        )

        if not pdf_result.get('success'):
            raise HTTPException(
                status_code=400,
                detail=f"PDF 파싱 실패: {pdf_result.get('error')}"
            )

        merged_result = merge_pdf_docx_data(pdf_result, docx_result)
        if not merged_result.get('success'):
            raise HTTPException(
                status_code=400,
                detail=f"PDF+DOCX 병합 실패: {merged_result.get('error')}"
            )

        new_articles = merged_result.get('articles', [])
        old_articles = current_json.get('조문내용', [])

        new_rule = {
            'wzruleseq': None,
            'wzname': current_rule['wzname'],
            'wzpubno': current_rule['wzpubno'],
            'wznewflag': "개정안",
            'wzlastrevdate': datetime.now().strftime('%Y-%m-%d')
        }

        if output_format not in ['pdf', 'docx']:
            output_format = 'pdf'

        if output_format == 'docx':
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
                tmp_path = tmp_file.name

            if not generate_comparison_docx(current_rule, new_rule, old_articles, new_articles, tmp_path, remarks):
                raise HTTPException(status_code=500, detail="DOCX 생성 실패")

            download_filename = f"신구대비표_{current_rule['wzpubno']}_{current_rule['wzname']}_미리보기.docx"
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp_file:
                tmp_path = tmp_file.name

            if not generate_comparison_pdf_v2(current_rule, new_rule, old_articles, new_articles, tmp_path, remarks):
                raise HTTPException(status_code=500, detail="PDF 생성 실패")

            download_filename = f"신구대비표_{current_rule['wzpubno']}_{current_rule['wzname']}_미리보기.pdf"
            media_type = "application/pdf"

        return FileResponse(
            path=tmp_path,
            filename=download_filename,
            media_type=media_type,
            background=None
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error generating comparison PDF for download: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/replace-comparison")
async def replace_comparison_table(
    wzruleseq: int = Form(..., description="규정 시퀀스 (wzruleseq)"),
    docx_file: UploadFile = File(..., description="수정된 DOCX 파일"),
    user: Dict[str, Any] = Depends(get_current_user)
):
    """수정된 신구대비표 DOCX 업로드 및 교체"""
    import subprocess

    try:
        if not docx_file.filename.lower().endswith('.docx'):
            raise HTTPException(status_code=400, detail="DOCX 파일만 업로드 가능합니다.")

        content = await docx_file.read()
        if len(content) > 10 * 1024 * 1024:
            raise HTTPException(status_code=400, detail="파일 크기가 10MB를 초과합니다.")

        with db_manager.get_connection() as conn:
            with conn.cursor() as cur:
                cur.execute("""
                    SELECT wzruleseq, wzruleid, wzname, wzpubno,
                           wzFileComparison, wzFileComparisonDocx, wzlastrevdate
                    FROM wz_rule
                    WHERE wzruleseq = %s
                """, (wzruleseq,))
                row = cur.fetchone()

        if not row:
            raise HTTPException(status_code=404, detail="규정을 찾을 수 없습니다.")

        rule_seq, wzruleid, wzname, wzpubno, current_pdf_path, current_docx_path, wzlastrevdate = row

        COMPARISON_TABLE_DIR.mkdir(parents=True, exist_ok=True)
        COMPARISON_TABLE_BACKUP_DIR.mkdir(parents=True, exist_ok=True)

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        docx_filename = f"{wzruleid}_{wzruleseq}_{timestamp}.docx"
        pdf_filename = f"{wzruleid}_{wzruleseq}_{timestamp}.pdf"

        docx_save_path = COMPARISON_TABLE_DIR / docx_filename
        pdf_save_path = COMPARISON_TABLE_DIR / pdf_filename

        with open(docx_save_path, 'wb') as f:
            f.write(content)
        logger.info(f"Saved new DOCX: {docx_save_path}")

        pdf_converted = False
        try:
            result = subprocess.run(
                [
                    'libreoffice', '--headless', '--convert-to', 'pdf',
                    '--outdir', str(COMPARISON_TABLE_DIR),
                    str(docx_save_path)
                ],
                capture_output=True,
                text=True,
                timeout=60
            )

            if result.returncode == 0:
                pdf_converted = True
                logger.info(f"PDF converted successfully: {pdf_save_path}")
            else:
                logger.error(f"LibreOffice conversion failed: {result.stderr}")

        except subprocess.TimeoutExpired:
            logger.error("LibreOffice conversion timed out")
        except FileNotFoundError:
            logger.error("LibreOffice not found. Please install libreoffice-writer.")

        pdf_relative_path = f"comparisonTable/{pdf_filename}" if pdf_converted else current_pdf_path
        docx_relative_path = f"comparisonTable/{docx_filename}"

        with db_manager.get_connection() as conn:
            with conn.cursor() as cur:
                cur.execute("""
                    UPDATE wz_rule
                    SET wzFileComparison = %s,
                        wzFileComparisonDocx = %s
                    WHERE wzruleseq = %s
                """, (pdf_relative_path, docx_relative_path, wzruleseq))
                conn.commit()

        logger.info(f"DB updated for wzruleseq={wzruleseq}: PDF={pdf_relative_path}, DOCX={docx_relative_path}")

        # wz_rule_history에 수정 이력 추가
        try:
            if create_history_record is None:
                raise ImportError("router_rule_history not available")
            modification_date = datetime.now().strftime('%Y.%m.%d')
            comp_pdf_abs = str(COMPARISON_TABLE_DIR / pdf_filename) if pdf_converted else None
            comp_docx_abs = str(COMPARISON_TABLE_DIR / docx_filename)

            history_id = create_history_record(
                wzruleseq=wzruleseq,
                wzruleid=wzruleid,
                wzpubno=wzpubno,
                wzname=wzname,
                action_type='modification',
                modification_date=modification_date,
                docx_path=None,
                pdf_path=None,
                comparison_path=comp_pdf_abs,
                comparison_docx_path=comp_docx_abs,
                note=f"신구대비표 교체: {docx_file.filename}",
                changed_by=user.get('username', 'unknown')
            )
            logger.info(f"History record created: {history_id}")
        except Exception as hist_err:
            logger.warning(f"Failed to create history record: {hist_err}")

        # summary JSON 갱신
        try:
            from api.service_rule_editor import run_json_merge_and_summary
            run_json_merge_and_summary()
            logger.info(f"Summary JSON updated after comparison table replacement")
        except Exception as summary_err:
            logger.warning(f"Failed to update summary JSON: {summary_err}")

        return JSONResponse(content={
            "success": True,
            "message": "신구대비표가 교체되었습니다." + ("" if pdf_converted else " (PDF 변환 실패, DOCX만 저장됨)"),
            "wzruleseq": wzruleseq,
            "wzruleid": wzruleid,
            "wzname": wzname,
            "docx_path": docx_relative_path,
            "pdf_path": pdf_relative_path if pdf_converted else None,
            "pdf_converted": pdf_converted
        })

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error replacing comparison table: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))
