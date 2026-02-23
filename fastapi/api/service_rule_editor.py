# -*- coding: utf-8 -*-
"""
    service_rule_editor.py
    ~~~~~~~~~~~~~~~~~~~~~~

    규정 편집 관련 API
    - 규정 수정, 삭제
    - DOCX/PDF 파일 파싱

    :copyright: (c) 2024 by wizice.
    :license:  wizice.com
"""

from fastapi import APIRouter, HTTPException, Depends, UploadFile, File, Form, Request, BackgroundTasks
from typing import Dict, Any, Optional
import os
import io
import json
from datetime import datetime
import logging
import shutil
from pathlib import Path

# 기존 파싱 도구 임포트
import sys
import tempfile
import subprocess

from .auth_middleware import get_current_user
from .timescaledb_manager_v2 import DatabaseConnectionManager
from . import query_download_logs
from settings import settings
from app_logger import get_logger

# settings import 후 applib 경로 추가
sys.path.append(settings.APPLIB_DIR)

logger = get_logger(__name__)

# JSON 병합 및 요약 스크립트 경로
JSON_ALL_SCRIPT = f'{settings.APPLIB_DIR}/JSON_ALL.py'
CREATE_SUMMARY_SCRIPT = f'{settings.APPLIB_DIR}/create_summary.py'
PYTHON_EXECUTABLE = sys.executable  # 현재 Python 실행 파일

# DB로부터 summary 생성 모듈 임포트
sys.path.insert(0, str(Path(__file__).parent.parent))
from generate_summary_from_db import generate_summary_from_db


def extract_document_name_from_merged_data(merged_data):
    """
    병합된 JSON 데이터에서 문서명을 추출하고 파일명에 적합하게 변환
    merge_json.py의 extract_document_name 로직과 동일
    """
    import re

    # 문서명 추출 우선순위
    document_name = None

    # 병합된 데이터에서 문서정보 확인
    if merged_data and "문서정보" in merged_data and merged_data["문서정보"]:
        doc_info = merged_data["문서정보"]
        document_name = doc_info.get("규정명") or doc_info.get("규정표기명")

    # 기본값 설정
    if not document_name:
        document_name = "unknown_document"

    # 파일명에 적합하게 정제
    # 공백을 언더스코어로 변환
    document_name = re.sub(r'\s+', '_', document_name.strip())

    # 파일명에 사용할 수 없는 문자 제거/변환
    document_name = re.sub(r'[<>:"/\\|?*]', '', document_name)

    # 연속된 언더스코어를 하나로 통합
    document_name = re.sub(r'_+', '_', document_name)

    # 앞뒤 언더스코어 제거
    document_name = document_name.strip('_')

    # 파일명 길이 제한 (Windows 호환성)
    if len(document_name) > 100:
        document_name = document_name[:100]

    return document_name

# 파싱 도구 임포트
try:
    # applib 모듈에서 import
    import sys
    import os
    import importlib.util

    # applib 경로를 명시적으로 지정
    applib_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'applib')

    # pdf2txt.py를 직접 로드 (venv의 pdf2txt와 충돌 방지)
    pdf2txt_spec = importlib.util.spec_from_file_location(
        "pdf2txt_local",
        os.path.join(applib_path, "pdf2txt.py")
    )
    pdf2txt_local = importlib.util.module_from_spec(pdf2txt_spec)
    pdf2txt_spec.loader.exec_module(pdf2txt_local)
    extract_text_from_pdf = pdf2txt_local.extract_text_from_pdf

    # txt2json.py도 같은 방식으로 로드
    txt2json_spec = importlib.util.spec_from_file_location(
        "txt2json_local",
        os.path.join(applib_path, "txt2json.py")
    )
    txt2json_local = importlib.util.module_from_spec(txt2json_spec)
    txt2json_spec.loader.exec_module(txt2json_local)
    MentalHealthRegulationParser = txt2json_local.MentalHealthRegulationParser

    PARSERS_AVAILABLE = True
    logger.info(f"PDF parsing tools loaded successfully from {applib_path}")
except ImportError as e:
    logger.warning(f"PDF parsing tools not available: {e}")
    PARSERS_AVAILABLE = False

    # 폴백 함수 정의 (import 실패 시)
    def extract_text_from_pdf(pdf_path, txt_path=None):
        """PDF에서 텍스트 추출 (폴백 버전)"""
        import pdfplumber
        import time
        import os

        start_time = time.time()
        file_size = os.path.getsize(pdf_path) / (1024 * 1024)  # MB
        logger.info(f"[PDF_PARSING] Starting PDF text extraction")
        logger.info(f"[PDF_PARSING] PDF file: {os.path.basename(pdf_path)}, Size: {file_size:.2f}MB")

        text_content = []
        page_times = []

        with pdfplumber.open(pdf_path) as pdf:
            total_pages = len(pdf.pages)
            logger.info(f"[PDF_PARSING] Total pages to process: {total_pages}")

            for i, page in enumerate(pdf.pages, 1):
                page_start = time.time()
                page_text = page.extract_text()
                page_elapsed = time.time() - page_start
                page_times.append(page_elapsed)

                if page_text:
                    text_content.append(page_text)

                # 매 10페이지마다 진행 상황 로그
                if i % 10 == 0 or i == total_pages:
                    avg_time = sum(page_times) / len(page_times)
                    elapsed = time.time() - start_time
                    remaining = avg_time * (total_pages - i)
                    logger.info(f"[PDF_PARSING] Progress: {i}/{total_pages} pages, "
                              f"Elapsed: {elapsed:.1f}s, Est. remaining: {remaining:.1f}s")

                # 페이지당 1초 이상 걸리면 경고
                if page_elapsed > 1.0:
                    logger.warning(f"[PDF_PARSING] Page {i} took {page_elapsed:.2f}s (slow)")

        full_text = '\n'.join(text_content)
        total_elapsed = time.time() - start_time

        logger.info(f"[PDF_PARSING] Extraction completed: {total_pages} pages in {total_elapsed:.2f}s")
        logger.info(f"[PDF_PARSING] Average time per page: {total_elapsed/total_pages:.2f}s")
        logger.info(f"[PDF_PARSING] Text size: {len(full_text)/1024:.1f}KB")

        if total_elapsed > 30:
            logger.warning(f"[BOTTLENECK] PDF extraction took {total_elapsed:.2f}s (>30s threshold)")

        # txt_path가 제공되면 파일로 저장
        if txt_path:
            save_start = time.time()
            with open(txt_path, 'w', encoding='utf-8') as f:
                f.write(full_text)
            save_elapsed = time.time() - save_start
            logger.info(f"[PDF_PARSING] Text saved to file in {save_elapsed:.2f}s")
            return txt_path
        else:
            return full_text

# DOCX 파싱 도구 설정
# applib 경로를 sys.path에 추가
if applib_path not in sys.path:
    sys.path.insert(0, applib_path)

# DOCX 파싱 가능 여부 플래그 (런타임에 체크)
DOCX_MODULE_AVAILABLE = False
DOCX_AVAILABLE = True  # True로 변경하여 첫 번째 경로 시도

# utils 모듈 import는 실제 사용 시점에 시도
logger.info("DOCX parsing will be attempted at runtime")

# DOCX 파싱을 위한 헬퍼 함수들 정의 (런타임에 사용)
def runtime_extract_metadata(doc):
    """DOCX 문서에서 메타데이터 추출 (런타임 버전)"""
    metadata = {}
    try:
        for para in doc.paragraphs:
            text = para.text.strip()
            if "제정일" in text or "제 정 일" in text:
                metadata["제정일"] = text.split(":")[-1].strip() if ":" in text else text
            elif "최종개정일" in text or "최 종 개 정 일" in text:
                metadata["최종개정일"] = text.split(":")[-1].strip() if ":" in text else text
            elif "담당부서" in text or "담 당 부 서" in text:
                metadata["담당부서"] = text.split(":")[-1].strip() if ":" in text else text
            elif "관련기준" in text or "관 련 기 준" in text:
                metadata["관련기준"] = text.split(":")[-1].strip() if ":" in text else text
    except:
        pass
    return metadata

def runtime_extract_numbers_from_docx(file_path):
    """DOCX에서 조문 추출 (런타임 버전)"""
    sections = []
    try:
        import docx as python_docx
        doc = python_docx.Document(file_path)
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                sections.append({
                    'seq': i + 1,
                    'text': text,
                    'level': 1
                })
    except:
        pass
    return sections

def runtime_convert_to_sections_format(extract_results):
    """섹션 형식으로 변환 (런타임 버전)"""
    sections = []
    for i, item in enumerate(extract_results):
        sections.append({
            "seq": i + 1,
            "레벨": item.get('level', 1),
            "번호": "",
            "내용": item.get('text', ''),
            "관련이미지": []
        })
    return sections


def generate_unique_wzruleid(cursor, max_attempts=100):
    """
    중복되지 않는 4자리 wzruleid 생성

    Args:
        cursor: DB cursor 객체
        max_attempts: 최대 시도 횟수 (기본값 100)

    Returns:
        int: 1000~9999 범위의 고유한 wzruleid

    Raises:
        Exception: max_attempts 내에 고유한 ID 생성 실패 시
    """
    import random

    for attempt in range(max_attempts):
        wzruleid = random.randint(1000, 9999)

        # 중복 확인
        cursor.execute("SELECT 1 FROM wz_rule WHERE wzruleid = %s LIMIT 1", (wzruleid,))
        if not cursor.fetchone():
            logger.info(f"Generated unique wzruleid: {wzruleid} (attempt {attempt + 1})")
            return wzruleid

    # 모든 시도 실패
    raise Exception(f"Failed to generate unique wzruleid after {max_attempts} attempts")


router = APIRouter(
    prefix="/api/v1/rule",
    tags=["rule-editor"],
    responses={404: {"description": "Not found"}},
)

# 공개 API 라우터 (인증 불필요 - 사용자 화면에서 접근)
public_router = APIRouter(
    prefix="/api/v1/rule-public",
    tags=["rule-public"],
    responses={404: {"description": "Not found"}},
)


def get_db_connection():
    """데이터베이스 연결 생성"""
    db_config = {
        'database': settings.DB_NAME,
        'user': settings.DB_USER,
        'password': settings.DB_PASSWORD,
        'host': settings.DB_HOST,
        'port': settings.DB_PORT
    }
    return DatabaseConnectionManager(**db_config)


@router.put("/update")
async def update_regulation(
    regulation: Dict[str, Any],
    user: Dict[str, Any] = Depends(get_current_user)
):
    """규정 수정"""
    try:
        db_manager = get_db_connection()

        with db_manager.get_connection() as conn:
            with conn.cursor() as cur:
                # 규정 업데이트
                cur.execute("""
                    UPDATE wz_rule
                    SET wzname = %s,
                        wzpubno = %s,
                        wzmgrdptnm = %s,
                        wzestabdate = %s,
                        wzexecdate = %s,
                        content_text = %s,
                        wzlastrevdate = %s,
                        wzNewFlag = %s
                    WHERE wzruleseq = %s
                """, (
                    regulation.get('wzname'),
                    regulation.get('wzpubno'),
                    regulation.get('wzmgrdptnm'),
                    regulation.get('wzestabdate'),
                    regulation.get('wzexecdate'),
                    regulation.get('content_text'),
                    datetime.now().strftime('%Y-%m-%d'),
                    regulation.get('wzNewFlag', '현행'),  # 기본값은 '현행'
                    regulation.get('wzruleseq')
                ))

                conn.commit()

                return {
                    "success": True,
                    "message": "규정이 수정되었습니다."
                }

    except Exception as e:
        logger.error(f"Error updating regulation: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/delete/{rule_id}")
async def delete_regulation(
    rule_id: int,
    background_tasks: BackgroundTasks,
    user: Dict[str, Any] = Depends(get_current_user)
):
    """규정 삭제 (JSON 파일 포함)"""
    try:
        db_manager = get_db_connection()

        with db_manager.get_connection() as conn:
            with conn.cursor() as cur:
                # 규정 정보 및 JSON 파일 경로 조회 (wzruleid 추가)
                cur.execute("""
                    SELECT wzname, wzFileJson, wzpubno, wzruleid FROM wz_rule
                    WHERE wzruleseq = %s
                """, (rule_id,))

                result = cur.fetchone()
                if not result:
                    raise HTTPException(status_code=404, detail="규정을 찾을 수 없습니다.")

                rule_name = result[0]
                json_path = result[1]
                rule_pubno = result[2]
                wzruleid = result[3]

                logger.info(f"Deleting regulation: name={rule_name}, pubno={rule_pubno}, wzruleid={wzruleid}")

                # JSON 파일 삭제
                deleted_files = []
                if json_path and os.path.exists(json_path):
                    try:
                        os.remove(json_path)
                        deleted_files.append(json_path)
                        logger.info(f"Deleted JSON file: {json_path}")
                    except Exception as e:
                        logger.warning(f"Failed to delete JSON file {json_path}: {e}")

                # merge_json 폴더에서 관련 파일 추가 검색 및 삭제
                import glob
                merge_json_dir = os.path.join(
                    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                    'applib', 'merge_json'
                )

                # 우선 wzruleid 기반 파일 삭제 (새 파일명 형식)
                if wzruleid:
                    wzruleid_file = os.path.join(merge_json_dir, f"{wzruleid}.json")
                    if os.path.exists(wzruleid_file):
                        try:
                            os.remove(wzruleid_file)
                            deleted_files.append(wzruleid_file)
                            logger.info(f"Deleted JSON file by wzruleid: {wzruleid_file}")
                        except Exception as e:
                            logger.warning(f"Failed to delete file {wzruleid_file}: {e}")

                # 추가로 wzpubno나 기타 패턴으로도 검색 (레거시 파일명 대응)
                if rule_pubno:
                    pattern = os.path.join(merge_json_dir, f"*{rule_pubno}*")
                    matching_files = glob.glob(pattern)

                    for file_path in matching_files:
                        if os.path.exists(file_path) and file_path not in deleted_files:
                            try:
                                os.remove(file_path)
                                deleted_files.append(file_path)
                                logger.info(f"Deleted related JSON file: {file_path}")
                            except Exception as e:
                                logger.warning(f"Failed to delete file {file_path}: {e}")

                # merge_json_old 폴더에서도 검색 및 삭제
                merge_json_old_dir = os.path.join(
                    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                    'applib', 'merge_json_old'
                )
                if os.path.exists(merge_json_old_dir):
                    pattern_old = os.path.join(merge_json_old_dir, f"*{rule_pubno}*")
                    matching_old_files = glob.glob(pattern_old)

                    for file_path in matching_old_files:
                        if os.path.exists(file_path):
                            try:
                                os.remove(file_path)
                                deleted_files.append(file_path)
                                logger.info(f"Deleted history JSON file: {file_path}")
                            except Exception as e:
                                logger.warning(f"Failed to delete history file {file_path}: {e}")

                # www/static/file 폴더에서도 검색 및 삭제 (사용자 화면용)
                www_static_dir = settings.WWW_STATIC_FILE_DIR
                if os.path.exists(www_static_dir):
                    # 우선 wzruleid 기반 파일 삭제
                    if wzruleid:
                        wzruleid_www_file = os.path.join(www_static_dir, f"{wzruleid}.json")
                        if os.path.exists(wzruleid_www_file):
                            try:
                                os.remove(wzruleid_www_file)
                                deleted_files.append(wzruleid_www_file)
                                logger.info(f"Deleted www JSON file by wzruleid: {wzruleid_www_file}")
                            except Exception as e:
                                logger.warning(f"Failed to delete www file {wzruleid_www_file}: {e}")

                    # 레거시 파일명도 검색
                    pattern_www = os.path.join(www_static_dir, f"*{rule_pubno}*")
                    matching_www_files = glob.glob(pattern_www)

                    for file_path in matching_www_files:
                        if os.path.exists(file_path) and file_path not in deleted_files:
                            try:
                                os.remove(file_path)
                                deleted_files.append(file_path)
                                logger.info(f"Deleted www/static/file JSON file: {file_path}")
                            except Exception as e:
                                logger.warning(f"Failed to delete www file {file_path}: {e}")

                # DB에서 규정 삭제
                cur.execute("""
                    DELETE FROM wz_rule
                    WHERE wzruleseq = %s
                """, (rule_id,))

                conn.commit()

                logger.info(f"Regulation deleted: {rule_name} (ID: {rule_id}) by {user['username']}")
                logger.info(f"Deleted {len(deleted_files)} JSON files")

                # 백그라운드에서 JSON 병합 및 요약 생성 (사용자 화면 업데이트용)
                background_tasks.add_task(run_json_merge_and_summary)
                logger.info(f"✅ 삭제 완료, 백그라운드 병합 작업 예약됨")

                return {
                    "success": True,
                    "message": f"'{rule_name}' 규정이 삭제되었습니다.",
                    "deleted_files": deleted_files,
                    "files_count": len(deleted_files)
                }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting regulation: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/create")
async def create_regulation(
    request: Request,
    background_tasks: BackgroundTasks,
    user: Dict[str, Any] = Depends(get_current_user)
) -> Dict[str, Any]:
    """
    신규 내규 생성

    Args:
        request: HTTP 요청 객체 (JSON 데이터 포함)
        user: 현재 사용자

    Returns:
        생성된 규정 정보
    """
    try:
        # Request body 파싱
        data = await request.json()
        logger.info(f"Creating new regulation with data: {data}")

        # 필수 필드 검증
        name = data.get('name')
        publication_no = data.get('publication_no')
        department = data.get('department')

        if not name or not publication_no:
            raise HTTPException(status_code=400, detail="제목과 분류번호는 필수 입력사항입니다.")

        # DB 연결
        db_config = {
            'host': settings.DB_HOST,
            'port': settings.DB_PORT,
            'database': settings.DB_NAME,
            'user': settings.DB_USER,
            'password': settings.DB_PASSWORD,
        }

        db_manager = DatabaseConnectionManager(**db_config)

        # 새 규정 데이터 준비
        established_date = data.get('established_date') or datetime.now().strftime('%Y-%m-%d')
        execution_date = data.get('execution_date') or datetime.now().strftime('%Y-%m-%d')
        category = data.get('category') or '규정'
        status = data.get('status', '현행')  # 현행/연혁 구분, 기본값은 '현행'

        # 분류번호에서 wzCateSeq 자동 추출
        # 예: "1.6.1." -> wzCateSeq = 1, "2.3.1." -> wzCateSeq = 2
        wz_cate_seq = None
        if publication_no:
            # 첫 번째 숫자 추출 (1.6.1. -> 1)
            import re
            match = re.match(r'^(\d+)\.', publication_no)
            if match:
                wz_cate_seq = int(match.group(1))
                logger.info(f"Extracted wzCateSeq={wz_cate_seq} from publication_no={publication_no}")

                # 분류가 유효한지 검증
                with db_manager.get_connection() as conn:
                    with conn.cursor() as cursor:
                        cursor.execute("SELECT wzCateName FROM wz_cate WHERE wzCateSeq = %s", (wz_cate_seq,))
                        cate_result = cursor.fetchone()

                        if not cate_result:
                            # 분류가 존재하지 않으면 에러 발생
                            logger.error(f"❌ 분류번호 {wz_cate_seq}번이 wz_cate 테이블에 존재하지 않습니다. (규정번호: {publication_no})")

                            # 사용 가능한 분류 목록 조회
                            cursor.execute("SELECT wzCateSeq, wzCateName FROM wz_cate ORDER BY wzCateSeq")
                            available_cates = cursor.fetchall()
                            cate_list = ', '.join([f"{row[0]}번({row[1].strip()})" for row in available_cates[:5]])

                            raise HTTPException(
                                status_code=400,
                                detail=f"유효하지 않은 분류번호입니다. {wz_cate_seq}번 분류가 존재하지 않습니다. 먼저 분류를 등록해주세요. (사용 가능한 분류 예시: {cate_list}...)"
                            )
                        else:
                            cate_name = cate_result[0].strip() if cate_result[0] else ''
                            logger.info(f"✅ 유효한 분류: {wz_cate_seq}번 - {cate_name}")
            else:
                logger.warning(f"Could not extract wzCateSeq from publication_no={publication_no}")

        # DB 연결 및 wzruleid 생성
        with db_manager.get_connection() as conn:
            cursor = conn.cursor()

            # ===== 분류번호 중복 체크 =====
            cursor.execute(
                "SELECT wzRuleSeq, wzName FROM wz_rule WHERE wzPubNo = %s AND wzNewFlag = 'N'",
                (publication_no,)
            )
            existing_rule = cursor.fetchone()

            if existing_rule:
                existing_seq = existing_rule[0]
                existing_name = existing_rule[1].strip() if existing_rule[1] else ''
                logger.warning(f"⚠️ Duplicate publication_no detected: {publication_no} (existing: seq={existing_seq}, name={existing_name})")
                raise HTTPException(
                    status_code=400,
                    detail=f"분류번호 '{publication_no}'는 이미 등록되어 있습니다. (기존 규정: {existing_name})"
                )

            logger.info(f"✅ Publication number {publication_no} is available")

            # 고유한 wzruleid 생성
            wzruleid = generate_unique_wzruleid(cursor)
            logger.info(f"Generated wzruleid={wzruleid} for new regulation: {name}")

            # INSERT 쿼리 실행
            insert_query = """
                INSERT INTO wz_rule (
                    wzRuleId, wzName, wzPubNo, wzMgrDptNm,
                    wzEstabDate, wzExecDate, wzLKndName,
                    wzLevel, wzEditType, wzNewFlag, wzCateSeq,
                    wzCreatedBy, wzModifiedBy,
                    wzFileDocx, wzFilePdf
                ) VALUES (
                    %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s
                ) RETURNING wzRuleSeq
            """

            params = (
                wzruleid,  # wzRuleId - 새로 생성된 고유 ID
                name, publication_no, department,
                established_date, execution_date, category,
                1,  # wzlevel - 기본값 1
                '제정',  # wzedittype - 신규는 '제정'
                status,  # wznewflag - '현행' 또는 '연혁'
                wz_cate_seq,  # wzCateSeq - 분류번호에서 자동 추출
                user.get('username'),  # wzcreatedby
                user.get('username'),  # wzmodifiedby
                '',  # wzFileDocx - 빈 문자열로 초기화
                ''   # wzFilePdf - 빈 문자열로 초기화
            )

            cursor.execute(insert_query, params)
            result = cursor.fetchone()
            conn.commit()
            cursor.close()

        # result는 튜플로 반환될 수 있음
        if result:
            if isinstance(result, tuple):
                new_id = result[0]
            elif isinstance(result, dict):
                # 딕셔너리일 때 키를 찾아서 값 추출
                if 'wzRuleSeq' in result:
                    new_id = result['wzRuleSeq']
                elif 'wzruleseq' in result:
                    new_id = result['wzruleseq']
                else:
                    # 첫 번째 값 사용
                    new_id = list(result.values())[0] if result else None
            else:
                new_id = result

            # new_id가 또다시 dict인 경우 처리
            if isinstance(new_id, dict):
                new_id = new_id.get('wzruleseq') or new_id.get('wzRuleSeq') or list(new_id.values())[0]

            logger.info(f"Successfully created new regulation with ID: {new_id}")

            # 백그라운드에서 JSON 병합 및 요약 생성 (사용자 화면 업데이트용)
            background_tasks.add_task(run_json_merge_and_summary)
            logger.info(f"✅ 생성 완료, 백그라운드 병합 작업 예약됨")

            return {
                "success": True,
                "rule_id": new_id,
                "message": "신규 내규가 생성되었습니다.",
                "data": {
                    "wzRuleSeq": new_id,
                    "wzName": name,
                    "wzPubNo": publication_no,
                    "wzMgrDptNm": department,
                    "wzEstabDate": established_date,
                    "wzExecDate": execution_date,
                    "wzLKndName": category,
                    "wzNewFlag": status,
                    "wzCateSeq": wz_cate_seq
                }
            }
        else:
            raise Exception("Failed to create new regulation - no ID returned")

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error creating regulation: {e}")
        raise HTTPException(status_code=500, detail=f"신규 내규 생성 실패: {str(e)}")


@router.post("/parse-revision")
async def parse_revision_file(
    pdf_file: Optional[UploadFile] = File(None),
    docx_file: Optional[UploadFile] = File(None),
    rule_id: int = Form(...),
    reason: Optional[str] = Form(None),
    revision_date: Optional[str] = Form(None),
    user: Dict[str, Any] = Depends(get_current_user)
):
    """개정 파일 파싱 (DOCX/PDF - 둘 다 업로드 가능)"""
    logger.info(f"[PARSE-REVISION] Starting parse-revision request")
    logger.info(f"[PARSE-REVISION] User: {user.get('username')}, Rule ID: {rule_id}")
    logger.info(f"[PARSE-REVISION] Files - PDF: {pdf_file.filename if pdf_file else 'None'}, DOCX: {docx_file.filename if docx_file else 'None'}")

    if not PARSERS_AVAILABLE:
        logger.error("[PARSE-REVISION] Parsers not available!")
        raise HTTPException(
            status_code=500,
            detail="문서 파싱 라이브러리를 불러올 수 없습니다."
        )

    # 두 파일 모두 필수
    if not pdf_file or not docx_file:
        missing_files = []
        if not pdf_file:
            missing_files.append("PDF")
        if not docx_file:
            missing_files.append("DOCX")

        raise HTTPException(
            status_code=400,
            detail=f"{', '.join(missing_files)} 파일을 업로드해주세요. PDF와 DOCX 파일 모두 필요합니다."
        )

    try:
        logger.info("[PARSE-REVISION] Starting file processing...")
        from datetime import datetime

        # 타임스탬프 생성
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        logger.info(f"[PARSE-REVISION] Timestamp: {timestamp}")

        # wzruleid 조회
        db_config = {
            'host': settings.DB_HOST,
            'port': settings.DB_PORT,
            'database': settings.DB_NAME,
            'user': settings.DB_USER,
            'password': settings.DB_PASSWORD,
        }
        db_manager = DatabaseConnectionManager(**db_config)

        with db_manager.get_connection() as conn:
            with conn.cursor() as cur:
                cur.execute("SELECT wzruleid FROM wz_rule WHERE wzruleseq = %s", (rule_id,))
                result = cur.fetchone()
                wzruleid = result[0] if result and result[0] else 0

        logger.info(f"[PARSE-REVISION] Retrieved wzruleid={wzruleid} for rule_id={rule_id}")

        # 저장 폴더 생성 (통합 경로 체계)
        pdf_folder = os.path.join(applib_path, "pdf")           # PDF 원본
        docx_folder = os.path.join(applib_path, "docx")         # DOCX 원본
        pdf_txt_folder = os.path.join(applib_path, "pdf_txt")   # PDF→TXT 결과
        docx_json_folder = os.path.join(applib_path, "docx_json") # DOCX→JSON 결과
        txt_json_folder = os.path.join(applib_path, "txt_json") # TXT→JSON 결과
        merge_json_folder = os.path.join(applib_path, "merge_json") # 병합 결과

        for folder in [pdf_folder, docx_folder, pdf_txt_folder, docx_json_folder, txt_json_folder, merge_json_folder]:
            os.makedirs(folder, exist_ok=True)

        results = {}
        combined_text = ""
        combined_json = None

        # DOCX 파일 처리
        if docx_file:
            logger.info(f"Processing DOCX file: {docx_file.filename}")
            docx_contents = await docx_file.read()

            # 타임스탬프가 추가된 파일명 (wzruleid 포함)
            base_name = os.path.splitext(docx_file.filename)[0]
            timestamped_docx_name = f"{wzruleid}_{base_name}_{timestamp}.docx"
            docx_save_path = os.path.join(docx_folder, timestamped_docx_name)

            # DOCX 파일 저장
            with open(docx_save_path, 'wb') as f:
                f.write(docx_contents)
            logger.info(f"DOCX file saved: {docx_save_path}")

            # docx2json.py의 독립 함수를 사용한 DOCX 파싱
            docx_parsed = False
            try:
                logger.info("[PARSE-REVISION] Attempting to use docx2json module...")
                # docx2json 모듈 임포트
                import sys
                applib_path_import = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'applib'))
                if applib_path_import not in sys.path:
                    sys.path.insert(0, applib_path_import)
                    logger.info(f"[PARSE-REVISION] Added path to sys.path: {applib_path_import}")

                from docx2json import process_docx_file
                logger.info("[PARSE-REVISION] Successfully imported process_docx_file")

                logger.info("[PARSE-REVISION] Using docx2json.process_docx_file() for DOCX parsing")

                # process_docx_file 함수 호출
                logger.info(f"[PARSE-REVISION] Calling process_docx_file with file: {docx_file.filename}, wzruleid: {wzruleid}")
                parse_result = process_docx_file(
                    file_content=docx_contents,
                    filename=docx_file.filename,
                    wzruleid=wzruleid
                )
                logger.info(f"[PARSE-REVISION] process_docx_file returned, success: {parse_result.get('success')}")

                if parse_result.get('success'):
                    # DOCX JSON 데이터 생성
                    docx_json_data = {
                        '문서정보': parse_result.get('document_info', {}),
                        '조문내용': parse_result.get('sections', [])
                    }

                    # DOCX JSON 파일 저장 (wzruleid 포함)
                    timestamped_docx_json_name = f"{wzruleid}_{base_name}_{timestamp}.json"
                    docx_json_save_path = os.path.join(docx_json_folder, timestamped_docx_json_name)

                    import json
                    with open(docx_json_save_path, 'w', encoding='utf-8') as f:
                        json.dump(docx_json_data, f, ensure_ascii=False, indent=2)
                    logger.info(f"DOCX JSON saved: {docx_json_save_path}")

                    # 파싱 결과를 기존 형식에 맞춰 변환
                    results['docx'] = {
                        'text': '\n'.join([s.get('내용', '') for s in parse_result.get('sections', [])]),
                        'json': docx_json_data,
                        'json_path': docx_json_save_path,
                        'filename': docx_file.filename
                    }
                    docx_parsed = True
                    logger.info(f"DOCX successfully parsed using docx2json library: {docx_file.filename}")
                else:
                    logger.warning(f"docx2json parsing failed: {parse_result.get('error')}")

            except ImportError as e:
                logger.warning(f"Failed to import docx2json: {e}")
            except Exception as e:
                logger.warning(f"docx2json parsing failed: {e}")

            # docx2json 파싱 실패 시 기존 방법 사용
            if not docx_parsed and DOCX_AVAILABLE:
                logger.info("Falling back to direct DOCX parsing")
                # applib의 DOCX 파싱 유틸리티 사용
                tmp_docx_path = None
                try:
                    # 임시 파일로 저장
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_docx:
                        tmp_docx.write(docx_contents)
                        tmp_docx_path = tmp_docx.name

                    # DOCX 파싱 시도
                    docx_parsed = False
                    doc = None
                    sections = []
                    metadata = {}
                    doc_title = docx_file.filename.replace('.docx', '')

                    # 방법 1: python-docx 라이브러리 사용
                    try:
                        import docx as python_docx
                        doc = python_docx.Document(tmp_docx_path)

                        # 메타데이터 추출 - 런타임 함수 사용
                        metadata = runtime_extract_metadata(doc)
                        logger.info(f"DOCX metadata extracted: {metadata}")

                        # 순차적 번호 추출 및 섹션 형식으로 변환 - 런타임 함수 사용
                        extract_results = runtime_extract_numbers_from_docx(tmp_docx_path)
                        sections = runtime_convert_to_sections_format(extract_results)

                        # 문서 제목 추출
                        if doc.paragraphs and doc.paragraphs[0].text.strip():
                            doc_title = doc.paragraphs[0].text.strip()

                        docx_parsed = True
                        logger.info("DOCX parsed using python-docx library")

                    except ImportError:
                        logger.warning("python-docx library not available, trying alternative method")

                        # 방법 2: unzip을 사용한 XML 파싱
                        try:
                            import subprocess
                            import re

                            # DOCX에서 document.xml 추출
                            result = subprocess.run(
                                ["unzip", "-p", tmp_docx_path, "word/document.xml"],
                                capture_output=True,
                                text=False
                            )

                            if result.returncode == 0:
                                xml_content = result.stdout.decode('utf-8', errors='ignore')
                                # <w:t> 태그에서 텍스트 추출
                                texts = re.findall(r'<w:t[^>]*>([^<]+)</w:t>', xml_content)

                                # 텍스트를 문단으로 그룹화
                                current_para = []
                                for i, text in enumerate(texts):
                                    current_para.append(text)
                                    # 간단한 문단 구분 (문장 종료 기호 확인)
                                    if any(text.endswith(p) for p in ['.', '!', '?', '\n']) or i == len(texts) - 1:
                                        para_text = ''.join(current_para).strip()
                                        if para_text:
                                            sections.append({
                                                'seq': len(sections) + 1,
                                                '내용': para_text,
                                                '레벨': 1,
                                                '번호': '',
                                                '관련이미지': []
                                            })
                                        current_para = []

                                docx_parsed = True
                                logger.info("DOCX parsed using unzip/XML method")
                        except Exception as e:
                            logger.error(f"Failed to parse DOCX using unzip: {e}")

                    # 파싱 실패 시 기본값 설정
                    if not docx_parsed:
                        sections = [{
                            'seq': 1,
                            '내용': f"DOCX 파일을 파싱할 수 없습니다. 파일명: {docx_file.filename}",
                            '레벨': 1,
                            '번호': '',
                            '관련이미지': []
                        }]
                        logger.warning("Using fallback DOCX content")

                    # 관련기준 처리
                    related_standards = []
                    if "관련기준" in metadata and metadata["관련기준"]:
                        related_standards = [item.strip() for item in metadata["관련기준"].split("\n") if item.strip()]

                    # 문서 정보 구조화
                    document_info = {
                        "규정명": doc_title,
                        "내규종류": "규정",
                        "규정표기명": doc_title,
                        "제정일": metadata.get("제정일", "").strip(),
                        "최종개정일": metadata.get("최종개정일", "").strip(),
                        "최종검토일": metadata.get("최종검토일", "").strip(),
                        "담당부서": metadata.get("담당부서", "").strip(),
                        "관련기준": related_standards,
                        "조문갯수": len(sections),
                        "이미지개수": 0
                    }

                    # 최종 JSON 구조 생성
                    document_structure = {
                        "문서정보": document_info,
                        "조문내용": sections
                    }

                    # 텍스트 내용 추출 (미리보기용)
                    docx_text = []
                    for section in sections:
                        if section.get('번호'):
                            docx_text.append(f"{section['번호']} {section.get('내용', '')}")
                        else:
                            docx_text.append(section.get('내용', ''))

                    docx_parsed_text = "\n".join(docx_text)

                    results['docx'] = {
                        'text': docx_parsed_text,
                        'json': document_structure,
                        'filename': docx_file.filename
                    }

                    logger.info(f"DOCX parsing completed: {docx_file.filename}")

                except Exception as e:
                    logger.error(f"Error parsing DOCX: {e}")
                    results['docx'] = {
                        'text': f"DOCX 파일 파싱 오류: {str(e)}",
                        'json': None,
                        'filename': docx_file.filename,
                        'error': True
                    }
                finally:
                    # 임시 파일 삭제
                    if tmp_docx_path and os.path.exists(tmp_docx_path):
                        try:
                            os.unlink(tmp_docx_path)
                        except:
                            pass
            elif not docx_parsed:  # docx2json 파싱이 실패했을 때만 실행
                # DOCX 파싱 라이브러리가 없는 경우 - Flask 앱 호출 시도
                logger.info("Attempting to use Flask app for DOCX parsing")
                tmp_docx_path = None
                try:
                    # 임시 파일로 저장
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_docx:
                        tmp_docx.write(docx_contents)
                        tmp_docx_path = tmp_docx.name

                    # unzip을 사용한 간단한 텍스트 추출 시도
                    import subprocess
                    import json
                    import re

                    # DOCX는 ZIP 형식이므로 unzip으로 텍스트 추출
                    result = subprocess.run(
                        ["unzip", "-p", tmp_docx_path, "word/document.xml"],
                        capture_output=True,
                        text=False,
                        timeout=10
                    )

                    if result.returncode == 0:
                        # XML에서 텍스트 추출
                        xml_content = result.stdout.decode('utf-8', errors='ignore')
                        # <w:t> 태그에서 텍스트 추출
                        texts = re.findall(r'<w:t[^>]*>([^<]+)</w:t>', xml_content)

                        # 텍스트를 섹션으로 구성
                        sections = []
                        current_text = []

                        for text in texts:
                            current_text.append(text)
                            # 문장 끝이나 단락 구분
                            if any(text.strip().endswith(p) for p in ['.', '!', '?']) or len(current_text) > 10:
                                combined = ''.join(current_text).strip()
                                if combined:
                                    sections.append({
                                        'seq': len(sections) + 1,
                                        '내용': combined,
                                        '레벨': 1,
                                        '번호': '',
                                        '관련이미지': []
                                    })
                                current_text = []

                        # 남은 텍스트 처리
                        if current_text:
                            combined = ''.join(current_text).strip()
                            if combined:
                                sections.append({
                                    'seq': len(sections) + 1,
                                    '내용': combined,
                                    '레벨': 1,
                                    '번호': '',
                                    '관련이미지': []
                                })

                        results['docx'] = {
                            'text': "\n".join([s['내용'] for s in sections]),
                            'json': {'문서정보': {}, '조문내용': sections},
                            'filename': docx_file.filename
                        }
                        logger.info("DOCX parsed successfully using unzip/XML extraction")
                    else:
                        # unzip 실패 시 간단한 대체 텍스트
                        results['docx'] = {
                            'text': f"[DOCX 파일: {docx_file.filename}]",
                            'json': {
                                '문서정보': {},
                                '조문내용': [{
                                    'seq': 1,
                                    '내용': f"DOCX 파일 내용을 추출할 수 없습니다: {docx_file.filename}",
                                    '레벨': 1,
                                    '번호': '',
                                    '관련이미지': []
                                }]
                            },
                            'filename': docx_file.filename
                        }
                        logger.warning(f"Failed to extract DOCX content using unzip: {result.stderr}")

                except Exception as e:
                    logger.error(f"Failed to parse DOCX using Flask app: {e}")
                    results['docx'] = {
                        'text': f"DOCX 파싱 실패: {str(e)}",
                        'json': None,
                        'filename': docx_file.filename,
                        'error': True
                    }
                finally:
                    if tmp_docx_path and os.path.exists(tmp_docx_path):
                        try:
                            os.unlink(tmp_docx_path)
                        except:
                            pass

        # PDF 파일 처리
        if pdf_file:
            logger.info(f"Processing PDF file: {pdf_file.filename}")
            pdf_contents = await pdf_file.read()

            # 타임스탬프가 추가된 파일명 (wzruleid 포함)
            base_name = os.path.splitext(pdf_file.filename)[0]
            timestamped_pdf_name = f"{wzruleid}_{base_name}_{timestamp}.pdf"
            timestamped_txt_name = f"{wzruleid}_{base_name}_{timestamp}.txt"
            timestamped_json_name = f"{wzruleid}_{base_name}_{timestamp}.json"

            pdf_save_path = os.path.join(pdf_folder, timestamped_pdf_name)
            txt_save_path = os.path.join(pdf_txt_folder, timestamped_txt_name)
            json_save_path = os.path.join(txt_json_folder, timestamped_json_name)

            # PDF 파일 저장
            with open(pdf_save_path, 'wb') as f:
                f.write(pdf_contents)
            logger.info(f"PDF file saved: {pdf_save_path}")

            # PDF 파싱 - 기존 pdf2txt.py 사용
            try:
                # PDF를 텍스트로 변환
                extract_text_from_pdf(pdf_save_path, txt_save_path)
                logger.info(f"Text file saved: {txt_save_path}")

                # 텍스트 파일 읽기
                with open(txt_save_path, 'r', encoding='utf-8') as f:
                    parsed_text = f.read()

                # JSON 변환 시도 (txt2json.py 사용)
                try:
                    parser = MentalHealthRegulationParser()
                    pdf_parsed_json = parser.parse_txt_to_json(parsed_text)
                    logger.info(f"PDF 파일 파싱 성공: {pdf_file.filename}")

                    # JSON 파일 저장
                    if pdf_parsed_json:
                        import json
                        with open(json_save_path, 'w', encoding='utf-8') as f:
                            json.dump(pdf_parsed_json, f, ensure_ascii=False, indent=2)
                        logger.info(f"JSON file saved: {json_save_path}")
                except Exception as e:
                    logger.warning(f"JSON 구조 파싱 실패, 텍스트만 반환: {e}")
                    pdf_parsed_json = None

                results['pdf'] = {
                    'text': parsed_text,
                    'json': pdf_parsed_json,
                    'filename': pdf_file.filename
                }

            except Exception as e:
                logger.error(f"Error parsing PDF: {e}")
                raise HTTPException(status_code=400, detail=f"PDF 파일 파싱 실패: {str(e)}")


        # 결과 통합
        if results:
            # PDF를 우선으로, 없으면 DOCX 사용
            if 'pdf' in results and not results['pdf'].get('error'):
                combined_text = results['pdf']['text']
                combined_json = results['pdf'].get('json')
            elif 'docx' in results and not results['docx'].get('error'):
                combined_text = results['docx']['text']
                combined_json = results['docx'].get('json')
            else:
                # 둘 다 오류인 경우
                error_msgs = []
                if 'pdf' in results and results['pdf'].get('error'):
                    error_msgs.append(f"PDF: {results['pdf']['text']}")
                if 'docx' in results and results['docx'].get('error'):
                    error_msgs.append(f"DOCX: {results['docx']['text']}")
                raise HTTPException(status_code=400, detail="\n".join(error_msgs))

        if not combined_text:
            raise HTTPException(status_code=400, detail="파일에서 텍스트를 추출할 수 없습니다.")

        # DB에 개정 이력 저장 (선택사항)
        if rule_id and reason:
            db_manager = get_db_connection()
            with db_manager.get_connection() as conn:
                with conn.cursor() as cur:
                    # 개정 이력 테이블이 있다면 저장
                    # cur.execute("""
                    #     INSERT INTO wz_rule_revision_history
                    #     (rule_id, reason, revision_date, created_by)
                    #     VALUES (%s, %s, %s, %s)
                    # """, (rule_id, reason, revision_date, user['username']))
                    pass

        response_data = {
            "success": True,
            "message": "파일이 성공적으로 파싱되었습니다.",
            "parsed_content": {
                "text": combined_text[:10000],  # 최대 10000자까지
                "total_length": len(combined_text),
                "files_processed": [],
                # 개별 파싱 결과 추가
                "pdf_result": None,
                "docx_result": None
            }
        }

        # 처리된 파일 정보 추가
        if 'pdf' in results:
            response_data["parsed_content"]["files_processed"].append({
                "type": "PDF",
                "filename": results['pdf']['filename'],
                "success": not results['pdf'].get('error', False)
            })
            # PDF 전체 결과 저장
            response_data["parsed_content"]["pdf_result"] = {
                "text": results['pdf'].get('text', ''),
                "json": results['pdf'].get('json'),
                "error": results['pdf'].get('error', False)
            }

        if 'docx' in results:
            response_data["parsed_content"]["files_processed"].append({
                "type": "DOCX",
                "filename": results['docx']['filename'],
                "success": not results['docx'].get('error', False)
            })
            # DOCX 전체 결과 저장
            response_data["parsed_content"]["docx_result"] = {
                "text": results['docx'].get('text', ''),
                "json": results['docx'].get('json'),
                "error": results['docx'].get('error', False)
            }

        # JSON 파싱 결과가 있으면 추가
        if combined_json:
            response_data["parsed_content"]["structured_data"] = combined_json

        logger.info(f"[PARSE-REVISION] Successfully processed files, returning response")
        logger.info(f"[PARSE-REVISION] Response keys: {list(response_data.keys())}")
        return response_data

    except HTTPException as he:
        logger.error(f"[PARSE-REVISION] HTTPException: {he.detail}")
        raise
    except Exception as e:
        logger.error(f"[PARSE-REVISION] Unexpected error processing revision file: {e}")
        logger.error(f"[PARSE-REVISION] Error type: {type(e).__name__}")
        import traceback
        logger.error(f"[PARSE-REVISION] Traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/upload-parse-pdf")
async def upload_and_parse_pdf(
    pdf_file: UploadFile = File(...),
    rule_id: int = Form(...),
    user: Dict[str, Any] = Depends(get_current_user)
):
    """PDF 파일 업로드 및 파싱 (개별)"""
    import time
    request_start = time.time()
    logger.info(f"[TIMING] upload_and_parse_pdf started for rule_id: {rule_id}")

    try:
        from datetime import datetime
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        # wzruleid 조회
        db_config = {
            'host': settings.DB_HOST,
            'port': settings.DB_PORT,
            'database': settings.DB_NAME,
            'user': settings.DB_USER,
            'password': settings.DB_PASSWORD,
        }
        db_manager = DatabaseConnectionManager(**db_config)

        with db_manager.get_connection() as conn:
            with conn.cursor() as cur:
                cur.execute("SELECT wzruleid FROM wz_rule WHERE wzruleseq = %s", (rule_id,))
                result = cur.fetchone()
                wzruleid = result[0] if result and result[0] else 0

        logger.info(f"Retrieved wzruleid={wzruleid} for rule_id={rule_id}")

        # 저장 폴더 생성 (통합 경로 체계)
        upload_pdf_folder = os.path.join(applib_path, "pdf")
        pdf_txt_json_folder = os.path.join(applib_path, "pdf_txt")
        os.makedirs(upload_pdf_folder, exist_ok=True)
        os.makedirs(pdf_txt_json_folder, exist_ok=True)

        # PDF 파일 저장 (wzruleid 추가)
        base_name = os.path.splitext(pdf_file.filename)[0]
        pdf_filename = f"{wzruleid}_{base_name}_{timestamp}.pdf"
        pdf_save_path = os.path.join(upload_pdf_folder, pdf_filename)

        save_start = time.time()
        pdf_contents = await pdf_file.read()
        with open(pdf_save_path, 'wb') as f:
            f.write(pdf_contents)
        save_elapsed = time.time() - save_start
        logger.info(f"[TIMING] PDF saved in {save_elapsed:.2f}s: {pdf_save_path}")

        # PDF → TXT 변환 (wzruleid 추가)
        txt_filename = f"{wzruleid}_{base_name}_{timestamp}.txt"
        txt_save_path = os.path.join(pdf_txt_json_folder, txt_filename)

        txt_start = time.time()
        extract_text_from_pdf(pdf_save_path, txt_save_path)
        txt_elapsed = time.time() - txt_start
        logger.info(f"[TIMING] PDF→TXT conversion completed in {txt_elapsed:.2f}s")

        if txt_elapsed > 10:
            logger.warning(f"[BOTTLENECK] PDF→TXT took {txt_elapsed:.2f}s (>10s threshold)")

        # TXT → JSON 변환
        with open(txt_save_path, 'r', encoding='utf-8') as f:
            txt_content = f.read()

        json_parse_start = time.time()
        parser = MentalHealthRegulationParser()
        json_data = parser.parse_txt_to_json(txt_content)
        json_parse_elapsed = time.time() - json_parse_start
        logger.info(f"[TIMING] TXT→JSON parsing completed in {json_parse_elapsed:.2f}s")

        if json_parse_elapsed > 10:
            logger.warning(f"[BOTTLENECK] TXT→JSON parsing took {json_parse_elapsed:.2f}s (>10s threshold)")

        # JSON 저장 (TXT→JSON 결과를 txt_json 폴더에 저장, wzruleid 추가)
        txt_json_folder = os.path.join(applib_path, "txt_json")
        os.makedirs(txt_json_folder, exist_ok=True)

        json_filename = f"{wzruleid}_{base_name}_{timestamp}.json"
        json_save_path = os.path.join(txt_json_folder, json_filename)

        import json
        with open(json_save_path, 'w', encoding='utf-8') as f:
            json.dump(json_data, f, ensure_ascii=False, indent=2)
        logger.info(f"JSON saved: {json_save_path}")

        # wzFilePdf 컬럼 업데이트 (업로드 즉시)
        try:
            pdf_relative_path = f"applib/pdf/{pdf_filename}"
            with db_manager.get_connection() as conn:
                with conn.cursor() as cur:
                    cur.execute(
                        "UPDATE wz_rule SET wzFilePdf = %s WHERE wzruleseq = %s",
                        (pdf_relative_path, rule_id)
                    )
                    conn.commit()
            logger.info(f"✅ wzFilePdf updated immediately after upload: {pdf_relative_path}")
        except Exception as update_error:
            logger.error(f"Error updating wzFilePdf: {update_error}")

        total_elapsed = time.time() - request_start
        logger.info(f"[TIMING] upload_and_parse_pdf completed in {total_elapsed:.2f}s")

        if total_elapsed > 30:
            logger.warning(f"[BOTTLENECK] Total PDF processing took {total_elapsed:.2f}s (>30s threshold)")

        return {
            "success": True,
            "message": "PDF 파일 처리 완료",
            "files": {
                "pdf": pdf_filename,
                "txt": txt_filename,
                "json": json_filename
            },
            "json_path": json_save_path,  # 전체 경로 추가
            "filepath": json_save_path,    # 호환성을 위해 중복
            "path": json_save_path,         # 호환성을 위해 중복
            "parsed_data": json_data
        }

    except Exception as e:
        logger.error(f"Error processing PDF: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/upload-parse-docx")
async def upload_and_parse_docx(
    docx_file: UploadFile = File(...),
    rule_id: int = Form(...),
    user: Dict[str, Any] = Depends(get_current_user)
):
    """DOCX 파일 업로드 및 파싱 (개별)"""
    try:
        from datetime import datetime
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        # wzruleid 조회
        db_config = {
            'host': settings.DB_HOST,
            'port': settings.DB_PORT,
            'database': settings.DB_NAME,
            'user': settings.DB_USER,
            'password': settings.DB_PASSWORD,
        }
        db_manager = DatabaseConnectionManager(**db_config)

        with db_manager.get_connection() as conn:
            with conn.cursor() as cur:
                cur.execute("SELECT wzruleid FROM wz_rule WHERE wzruleseq = %s", (rule_id,))
                result = cur.fetchone()
                wzruleid = result[0] if result and result[0] else 0

        logger.info(f"Retrieved wzruleid={wzruleid} for rule_id={rule_id}")

        # 저장 폴더 생성 (통합 경로 체계)
        upload_docx_folder = os.path.join(applib_path, "docx")
        docx_json_folder = os.path.join(applib_path, "docx_json")
        os.makedirs(upload_docx_folder, exist_ok=True)
        os.makedirs(docx_json_folder, exist_ok=True)

        # DOCX 파일 저장 (wzruleid 추가)
        base_name = os.path.splitext(docx_file.filename)[0]
        docx_filename = f"{wzruleid}_{base_name}_{timestamp}.docx"
        docx_save_path = os.path.join(upload_docx_folder, docx_filename)

        docx_contents = await docx_file.read()
        with open(docx_save_path, 'wb') as f:
            f.write(docx_contents)
        logger.info(f"DOCX saved: {docx_save_path}")

        # DOCX → JSON 변환
        from docx2json import process_docx_file

        parse_result = process_docx_file(
            file_content=docx_contents,
            filename=docx_file.filename,
            wzruleid=wzruleid
        )

        json_data = None
        if parse_result.get('success'):
            json_data = {
                '문서정보': parse_result.get('document_info', {}),
                '조문내용': parse_result.get('sections', [])
            }

            # images와 preview_data도 복사 (있으면)
            if parse_result.get('images'):
                json_data['images'] = parse_result.get('images', [])
            if parse_result.get('preview_data'):
                json_data['preview_data'] = parse_result.get('preview_data')

            # JSON 저장 (wzruleid 추가)
            json_filename = f"{wzruleid}_{base_name}_{timestamp}.json"
            json_save_path = os.path.join(docx_json_folder, json_filename)

            import json
            with open(json_save_path, 'w', encoding='utf-8') as f:
                json.dump(json_data, f, ensure_ascii=False, indent=2)
            logger.info(f"DOCX JSON saved: {json_save_path}")

        # wzFileDocx 컬럼 업데이트 (업로드 즉시)
        try:
            docx_relative_path = f"applib/docx/{docx_filename}"
            with db_manager.get_connection() as conn:
                with conn.cursor() as cur:
                    cur.execute(
                        "UPDATE wz_rule SET wzFileDocx = %s WHERE wzruleseq = %s",
                        (docx_relative_path, rule_id)
                    )
                    conn.commit()
            logger.info(f"✅ wzFileDocx updated immediately after upload: {docx_relative_path}")
        except Exception as update_error:
            logger.error(f"Error updating wzFileDocx: {update_error}")

        return {
            "success": True,
            "message": "DOCX 파일 처리 완료",
            "files": {
                "docx": docx_filename,
                "json": json_filename if json_data else None
            },
            "json_path": json_save_path if json_data else None,  # 전체 경로 추가
            "filepath": json_save_path if json_data else None,    # 호환성을 위해 중복
            "path": json_save_path if json_data else None,         # 호환성을 위해 중복
            "parsed_data": json_data
        }

    except Exception as e:
        logger.error(f"Error processing DOCX: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/save-merged")
async def save_merged_data(
    data: Dict[str, Any],
    user: Dict[str, Any] = Depends(get_current_user)
):
    """병합된 데이터를 JSON 파일로 저장"""
    try:
        from datetime import datetime
        import json

        # 타임스탬프 생성
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        # wzruleid 조회
        rule_id = data.get('rule_id', 0)
        wzruleid = 0
        if rule_id:
            db_config = {
                'host': settings.DB_HOST,
                'port': settings.DB_PORT,
                'database': settings.DB_NAME,
                'user': settings.DB_USER,
                'password': settings.DB_PASSWORD,
            }
            db_manager = DatabaseConnectionManager(**db_config)

            with db_manager.get_connection() as conn:
                with conn.cursor() as cur:
                    cur.execute("SELECT wzruleid FROM wz_rule WHERE wzruleseq = %s", (rule_id,))
                    result = cur.fetchone()
                    wzruleid = result[0] if result and result[0] else 0

        logger.info(f"Retrieved wzruleid={wzruleid} for rule_id={rule_id}")

        # 병합 폴더 생성
        merged_folder = os.path.join(applib_path, "merge_json")
        os.makedirs(merged_folder, exist_ok=True)

        # 파일명 생성 - 문서명 기반으로 변경 (wzruleid 포함)
        merged_data = data.get('merged_data', {})
        if merged_data:
            document_name = extract_document_name_from_merged_data(merged_data)
        else:
            # 폴백: 기존 방식 사용
            pdf_filename = data.get('pdf_filename', 'unknown.pdf')
            base_name = os.path.splitext(pdf_filename)[0]
            document_name = base_name.replace(' ', '_')

        merged_filename = f"{wzruleid}_{document_name}_{timestamp}.json"
        merged_path = os.path.join(merge_json_folder, merged_filename)

        # 병합된 데이터 사용 (merge_json.py의 JSONMerger가 생성한 완전한 구조)
        merged_json = data.get('merged_data', {})

        # 규정ID만 추가 (문서정보가 있는 경우)
        if isinstance(merged_json, dict) and "문서정보" in merged_json:
            merged_json["문서정보"]["규정ID"] = data.get('rule_id', 0)
        elif not merged_json:
            # 병합 데이터가 없는 경우 기본 구조 생성
            merged_json = {
                "문서정보": {
                    "규정명": "",
                    "PDF파일": pdf_filename,
                    "DOCX파일": data.get('docx_filename', 'unknown.docx'),
                    "병합일시": timestamp,
                    "규정ID": data.get('rule_id', 0)
                },
                "병합내용": {},
                "처리정보": {
                    "처리자": user.get('username', 'unknown'),
                    "처리시간": data.get('timestamp', datetime.now().isoformat())
                }
            }

        # JSON 파일로 저장
        with open(merged_path, 'w', encoding='utf-8') as f:
            json.dump(merged_json, f, ensure_ascii=False, indent=2)

        logger.info(f"Merged data saved to: {merged_path}")

        # 상대경로로 변환
        from api.file_utils import get_relative_path
        merged_relative_path = get_relative_path(merged_path)

        return {
            "success": True,
            "file_path": merged_path,  # 절대경로 (호환성)
            "relative_path": merged_relative_path,  # 상대경로 (DB 저장용)
            "filename": merged_filename
        }

    except Exception as e:
        logger.error(f"Error saving merged data: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/merge-documents")
async def merge_documents(
    merge_data: Dict[str, Any],
    user: Dict[str, Any] = Depends(get_current_user)
):
    """PDF와 DOCX 파싱 결과를 병합 (개선된 버전)"""
    try:
        from datetime import datetime
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        # 병합 폴더 생성
        merged_folder = os.path.join(applib_path, "merge_json")
        os.makedirs(merged_folder, exist_ok=True)

        pdf_result = merge_data.get('pdf_result', {})
        docx_result = merge_data.get('docx_result', {})
        pdf_filename = merge_data.get('pdf_filename', 'unknown.pdf')

        # 문서명 기반 파일명 생성 (병합 후 결정)
        # 임시로 기본 파일명 사용 (병합 후 재설정됨)
        base_name = os.path.splitext(pdf_filename)[0]
        temp_merged_filename = f"merged_{base_name}_{timestamp}.json"
        temp_merged_path = os.path.join(merge_json_folder, temp_merged_filename)

        # merge_json.py 사용하여 병합
        pdf_json_data = pdf_result.get('json', {})
        docx_json_data = docx_result.get('json', {})

        # 새로운 병합 로직 사용
        try:
            # merge_json 모듈 동적 임포트
            merge_json_spec = importlib.util.spec_from_file_location(
                "merge_json",
                os.path.join(applib_path, "merge_json.py")
            )
            merge_json_module = importlib.util.module_from_spec(merge_json_spec)
            merge_json_spec.loader.exec_module(merge_json_module)

            # 직접 JSON 데이터 전달
            merger = merge_json_module.JSONMerger(pdf_json_data=pdf_json_data, docx_json_data=docx_json_data)

            # 병합 실행
            merged_data = merger.merge_regulation()
        except Exception as e:
            logger.warning(f"merge_json 사용 실패: {e}, 기본 병합 로직 사용")
            merged_data = None

        if not merged_data:
            # 폴백: 기존 단순 병합 로직
            merged_data = {
                "문서정보": {},
                "조문내용": []
            }

            # DOCX 문서정보 사용 (더 완전한 정보)
            if docx_json_data and docx_json_data.get('문서정보'):
                merged_data['문서정보'] = docx_json_data['문서정보']
            elif pdf_json_data and pdf_json_data.get('문서정보'):
                merged_data['문서정보'] = pdf_json_data['문서정보']

            # PDF와 DOCX 조문내용 병합
            pdf_articles = pdf_json_data.get('조문내용', [])
            docx_articles = docx_json_data.get('조문내용', [])

            # 병합 로직: PDF의 번호체계(제목) + DOCX의 상세 내용
            merged_articles = []

            logger.info(f"병합 시작 - PDF 조문: {len(pdf_articles)}개, DOCX 조문: {len(docx_articles)}개")

            # PDF를 기준으로 병합 (PDF의 번호체계 우선)
            for i, pdf_item in enumerate(pdf_articles):
                merged_item = {
                    # PDF의 제목/번호를 우선 사용
                    '제목': pdf_item.get('제목', ''),
                    '번호': pdf_item.get('번호', ''),
                    '내용': pdf_item.get('내용', ''),
                    '부가정보': pdf_item.get('부가정보', {})
                }

                # 같은 순서의 DOCX 항목이 있으면 내용을 DOCX에서 가져오기
                if i < len(docx_articles):
                    docx_item = docx_articles[i]

                    # DOCX의 내용을 우선 사용 (더 완전한 텍스트 정보)
                    if docx_item.get('내용'):
                        merged_item['내용'] = docx_item['내용']

                    # 제목이 PDF에 없고 DOCX에 있으면 사용
                    if not merged_item['제목'] and docx_item.get('제목'):
                        merged_item['제목'] = docx_item['제목']

                    # 이미지 정보는 DOCX에서만 가져오기
                    if docx_item.get('관련이미지'):
                        merged_item['관련이미지'] = docx_item['관련이미지']

                merged_articles.append(merged_item)
                logger.debug(f"병합된 조문 {i+1}: 제목='{merged_item.get('제목', '')}', 내용 길이={len(merged_item.get('내용', ''))}")

            # DOCX에만 있는 추가 항목들 (PDF보다 많은 경우)
            if len(docx_articles) > len(pdf_articles):
                for j, docx_item in enumerate(docx_articles[len(pdf_articles):], len(pdf_articles)):
                    merged_articles.append(docx_item)
                    logger.debug(f"DOCX 추가 조문 {j+1}: {docx_item.get('제목', '')}")

            logger.info(f"병합 완료 - 총 {len(merged_articles)}개 조문")

            merged_data['조문내용'] = merged_articles

        # 병합 후 문서명 기반으로 올바른 파일명 생성
        document_name = extract_document_name_from_merged_data(merged_data)
        merged_filename = f"merged_{document_name}_{timestamp}.json"
        merged_path = os.path.join(merge_json_folder, merged_filename)

        # 병합 결과 저장
        import json
        with open(merged_path, 'w', encoding='utf-8') as f:
            json.dump(merged_data, f, ensure_ascii=False, indent=2)

        logger.info(f"Merged file saved with document name: {merged_path}")

        return {
            "success": True,
            "message": "문서가 성공적으로 병합되었습니다.",
            "merged_path": merged_path,
            "merged_content": merged_data
        }

    except Exception as e:
        logger.error(f"Error merging documents: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/get/{rule_id}")
async def get_rule_detail(
    rule_id: int,
    user: Dict[str, Any] = Depends(get_current_user)
):
    """규정 상세 정보 조회"""
    try:
        db_config = {
            'host': settings.DB_HOST,
            'port': settings.DB_PORT,
            'database': settings.DB_NAME,
            'user': settings.DB_USER,
            'password': settings.DB_PASSWORD,
        }

        db_manager = DatabaseConnectionManager(**db_config)

        with db_manager.get_connection() as conn:
            with conn.cursor() as cur:
                cur.execute("""
                    SELECT wzRuleSeq, wzLevel, wzRuleId, wzName, wzEditType,
                           wzPubNo, wzEstabDate, wzLastRevDate, wzMgrDptNm,
                           wzMgrDptOrgCd, wzRelDptNm, wzRelDptOrgCd, wzCateSeq,
                           wzExecDate, wzLKndName, wzCloseDate, wzFileDocx, wzFilePdf,
                           wzNewFlag, content_text
                    FROM wz_rule
                    WHERE wzruleseq = %s
                """, (rule_id,))

                columns = [desc[0].lower() for desc in cur.description]
                row = cur.fetchone()

                if not row:
                    raise HTTPException(status_code=404, detail="규정을 찾을 수 없습니다.")

                rule_data = dict(zip(columns, row))

                # 날짜 포맷팅
                for date_field in ['wzestabdate', 'wzlastrevdate', 'wzexecdate', 'wzclosedate']:
                    if rule_data.get(date_field):
                        if isinstance(rule_data[date_field], datetime):
                            rule_data[date_field] = rule_data[date_field].strftime('%Y-%m-%d')
                        elif isinstance(rule_data[date_field], str):
                            rule_data[date_field] = rule_data[date_field][:10] if len(rule_data[date_field]) >= 10 else rule_data[date_field]

                return {
                    "success": True,
                    "data": rule_data
                }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error fetching rule detail: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.put("/update-basic/{rule_id}")
async def update_rule_basic_info(
    rule_id: int,
    request: Request,
    background_tasks: BackgroundTasks,
    user: Dict[str, Any] = Depends(get_current_user)
):
    """규정 기본정보 업데이트"""
    try:
        # Request body 파싱
        data = await request.json()
        logger.info(f"Updating basic info for rule {rule_id} with data: {data}")

        db_config = {
            'host': settings.DB_HOST,
            'port': settings.DB_PORT,
            'database': settings.DB_NAME,
            'user': settings.DB_USER,
            'password': settings.DB_PASSWORD,
        }

        db_manager = DatabaseConnectionManager(**db_config)

        # 업데이트 쿼리 생성
        update_fields = []
        params = []

        # 업데이트 가능한 필드들
        field_mapping = {
            'wzname': 'wzName',
            'wzpubno': 'wzPubNo',
            'wzmgrdptnm': 'wzMgrDptNm',
            'wzmgrdptorgcd': 'wzMgrDptOrgCd',    # 담당부서 코드 추가
            'wzreldptnm': 'wzRelDptNm',          # 유관부서명 추가
            'wzreldptorgcd': 'wzRelDptOrgCd',    # 유관부서 코드 추가
            'wzestabdate': 'wzEstabDate',
            'wzlastrevdate': 'wzLastRevDate',    # 개정일 추가
            'wzexecdate': 'wzExecDate',
            'wzlkndname': 'wzLKndName',
            'wznewflag': 'wzNewFlag'
        }

        for db_field, json_field in field_mapping.items():
            if json_field in data or db_field in data:
                value = data.get(json_field) or data.get(db_field)
                if value is not None:
                    update_fields.append(f"{db_field} = %s")
                    params.append(value)

        if not update_fields:
            return {"success": False, "message": "업데이트할 필드가 없습니다."}

        # ModifiedBy 추가
        update_fields.append("wzModifiedBy = %s")
        params.append(user.get('username', 'admin'))

        # WHERE 조건 추가
        params.append(rule_id)

        query = f"""
            UPDATE wz_rule
            SET {', '.join(update_fields)}
            WHERE wzruleseq = %s
        """

        row_count = db_manager.execute_query(query, params, commit=True)

        if not row_count or row_count == 0:
            raise HTTPException(status_code=404, detail="규정을 찾을 수 없습니다.")

        logger.info(f"Successfully updated basic info for rule {rule_id}")

        # 백그라운드 병합 작업 예약 (사용자 화면 업데이트)
        background_tasks.add_task(run_json_merge_and_summary)
        logger.info(f"✅ 기본정보 저장 완료, 백그라운드 병합 작업 예약됨")

        return {
            "success": True,
            "message": "기본정보가 업데이트되었습니다. (백그라운드에서 병합 중)",
            "updated_fields": len(update_fields) - 1  # ModifiedBy 제외
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error updating basic info: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/create-revision/{rule_id}")
async def create_revision(
    rule_id: int,
    request: Request,
    user: Dict[str, Any] = Depends(get_current_user)
):
    """규정 개정판 생성 - 기존 규정 복사하여 새 개정본 생성"""
    try:
        # Request body 파싱
        revision_data = await request.json()
        logger.info(f"Creating revision for rule {rule_id} with data: {revision_data}")

        # DB 연결
        db_config = {
            'host': settings.DB_HOST,
            'port': settings.DB_PORT,
            'database': settings.DB_NAME,
            'user': settings.DB_USER,
            'password': settings.DB_PASSWORD,
        }

        db_manager = DatabaseConnectionManager(**db_config)

        with db_manager.get_connection() as conn:
            with conn.cursor() as cur:
                # 1. 기존 규정 조회 (딕셔너리 커서 사용) - wzFileJson 추가
                cur.execute("""
                    SELECT wzRuleSeq, wzLevel, wzRuleId, wzName, wzEditType,
                           wzPubNo, wzEstabDate, wzLastRevDate, wzMgrDptNm,
                           wzMgrDptOrgCd, wzRelDptNm, wzRelDptOrgCd, wzCateSeq,
                           wzExecDate, wzLKndName, wzCloseDate, wzFileDocx, wzFilePdf,
                           content_text, wzNewFlag, wzFileJson
                    FROM wz_rule
                    WHERE wzruleseq = %s
                """, (rule_id,))

                columns = [desc[0].lower() for desc in cur.description]
                row = cur.fetchone()

                if not row:
                    raise HTTPException(status_code=404, detail="원본 규정을 찾을 수 없습니다.")

                # 딕셔너리로 변환
                original = dict(zip(columns, row))
                logger.info(f"Found original rule: {original.get('wzname')}")

                # wzruleid 처리: 원본에 없으면(0이면) 새로 생성하고 원본도 업데이트
                original_wzruleid = original.get('wzruleid')
                if not original_wzruleid or original_wzruleid == 0:
                    wzruleid = generate_unique_wzruleid(cur)
                    logger.info(f"Original wzruleid was 0, generated new wzruleid={wzruleid}")
                    # 원본도 업데이트하여 같은 계보로 묶음
                    cur.execute("""
                        UPDATE wz_rule SET wzruleid = %s WHERE wzruleseq = %s
                    """, (wzruleid, rule_id))
                    logger.info(f"Updated original rule {rule_id} with wzruleid={wzruleid}")
                else:
                    wzruleid = original_wzruleid
                    logger.info(f"Using original wzruleid={wzruleid}")

                # 2. 기존 파일들을 _old 폴더로 이동 (개정일 기준으로 파일 찾기)
                from api.file_utils import move_file_to_old

                wzruleseq = rule_id  # 현재 규정의 시퀀스
                # 원본 규정의 개정일 (이 날짜의 파일을 _old로 이동)
                original_revision_date = original.get('wzlastrevdate', '')
                logger.info(f"Original revision date: {original_revision_date}")

                # JSON 파일 이동: JSON은 이동하지 않음 (사용자 요청)
                new_json_old_path = move_file_to_old(wzruleseq, wzruleid, 'json', original_revision_date)
                logger.info(f"JSON moved: wzruleseq={wzruleseq} -> {new_json_old_path}")

                # PDF 파일 이동: 개정일과 매칭되는 파일을 _old로 이동
                new_pdf_old_path = move_file_to_old(wzruleseq, wzruleid, 'pdf', original_revision_date)
                logger.info(f"PDF moved: wzruleseq={wzruleseq}, date={original_revision_date} -> {new_pdf_old_path}")

                # DOCX 파일 이동: 개정일과 매칭되는 파일을 _old로 이동
                new_docx_old_path = move_file_to_old(wzruleseq, wzruleid, 'docx', original_revision_date)
                logger.info(f"DOCX moved: wzruleseq={wzruleseq}, date={original_revision_date} -> {new_docx_old_path}")

                # 3. 기존 규정을 연혁으로 변경 (이동된 파일 경로로 업데이트)
                # DB에 NOT NULL 제약조건이 있으므로 None 대신 빈 문자열 사용
                cur.execute("""
                    UPDATE wz_rule
                    SET wzNewFlag = '연혁',
                        wzCloseDate = %s,
                        wzModifiedBy = %s,
                        wzFileJson = %s,
                        wzFilePdf = %s,
                        wzFileDocx = %s
                    WHERE wzruleseq = %s
                """, (
                    revision_data.get('revision_date', datetime.now().strftime('%Y-%m-%d')),
                    user.get('username'),
                    new_json_old_path or '',  # applib/merge_json_old/{wzruleid}_{wzruleseq}.json
                    new_pdf_old_path or '',   # applib/pdf_old/{wzruleid}_{wzruleseq}.pdf
                    new_docx_old_path or '',  # applib/docx_old/{wzruleid}_{wzruleseq}.docx
                    rule_id
                ))

                # 4. 새 개정판 생성 (기존 데이터 복사)
                insert_query = """
                    INSERT INTO wz_rule (
                        wzLevel, wzRuleId, wzName, wzEditType,
                        wzPubNo, wzEstabDate, wzLastRevDate, wzMgrDptNm,
                        wzMgrDptOrgCd, wzRelDptNm, wzRelDptOrgCd, wzCateSeq,
                        wzExecDate, wzLKndName, wzFileDocx, wzFilePdf, wzFileJson,
                        content_text, wzNewFlag, wzCreatedBy, wzModifiedBy
                    ) VALUES (
                        %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s
                    ) RETURNING wzRuleSeq
                """

                params = (
                    original.get('wzlevel', 1),
                    wzruleid,  # 위에서 확인/생성한 wzruleid 사용
                    original.get('wzname'),
                    '개정',  # wzEditType을 '개정'으로 변경
                    original.get('wzpubno'),
                    original.get('wzestabdate'),
                    revision_data.get('revision_date', datetime.now().strftime('%Y-%m-%d')),  # 개정일
                    original.get('wzmgrdptnm'),
                    original.get('wzmgrdptorgcd'),
                    original.get('wzreldptnm'),
                    original.get('wzreldptorgcd'),
                    original.get('wzcateseq'),
                    revision_data.get('execution_date', datetime.now().strftime('%Y-%m-%d')),  # 시행일
                    original.get('wzlkndname'),
                    original.get('wzfiledocx', ''),
                    original.get('wzfilepdf', ''),
                    '',  # wzFileJson - 초기에는 빈 값, 나중에 저장 시 업데이트됨
                    original.get('content_text', ''),
                    '현행',  # wzNewFlag - 초기값은 '현행' (작성중은 제약조건 위반)
                    user.get('username'),
                    user.get('username')
                )

                cur.execute(insert_query, params)
                new_id = cur.fetchone()[0]

                # 커밋
                conn.commit()

                logger.info(f"Successfully created revision with new ID: {new_id}")

                return {
                    "success": True,
                    "message": "개정판이 생성되었습니다.",
                    "rule_id": new_id,  # JavaScript에서 rule_id로 받기 위해
                    "new_rule_id": new_id,
                    "original_rule_id": rule_id
                }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error creating revision: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/upload-comparison-table/{rule_id}")
async def upload_comparison_table(
    rule_id: int,
    comparison_file: UploadFile = File(...),
    user: Dict[str, Any] = Depends(get_current_user)
):
    """
    신구대비표 파일 업로드 (PDF, DOCX, HWP, HWPX 등 다양한 형식 지원)

    Args:
        rule_id: 규정 ID (wzRuleSeq)
        comparison_file: 신구대비표 파일
        user: 현재 사용자

    Returns:
        업로드 결과 및 파일 경로
    """
    try:
        logger.info(f"Uploading comparison table for rule {rule_id}")

        # 허용된 파일 확장자
        allowed_extensions = ['.pdf', '.docx', '.doc', '.hwp', '.hwpx', '.xlsx', '.xls']
        file_ext = Path(comparison_file.filename).suffix.lower()

        if file_ext not in allowed_extensions:
            raise HTTPException(
                status_code=400,
                detail=f"허용된 파일 형식: {', '.join(allowed_extensions)}"
            )

        # DB에서 규정 정보 조회
        db_config = {
            'host': settings.DB_HOST,
            'port': settings.DB_PORT,
            'database': settings.DB_NAME,
            'user': settings.DB_USER,
            'password': settings.DB_PASSWORD,
        }

        db_manager = DatabaseConnectionManager(**db_config)

        with db_manager.get_connection() as conn:
            with conn.cursor() as cur:
                cur.execute("""
                    SELECT wzRuleSeq, wzRuleId, wzPubNo, wzLastRevDate, wzNewFlag
                    FROM wz_rule
                    WHERE wzruleseq = %s
                """, (rule_id,))

                rule = cur.fetchone()
                if not rule:
                    raise HTTPException(status_code=404, detail="규정을 찾을 수 없습니다.")

                wzruleseq, wzruleid, wzpubno, wzlastrevdate, wznewflag = rule

                # 파일명 생성: wzRuleId_wzRuleSeq_개정일.{원본확장자}
                revision_date_str = wzlastrevdate.replace('.', '').replace('-', '') if wzlastrevdate else 'nodate'
                filename = f"{wzruleid}_{wzruleseq}_{revision_date_str}{file_ext}"

                # 저장 경로
                save_dir = Path(f'{settings.WWW_STATIC_PDF_DIR}/comparisonTable')
                save_dir.mkdir(parents=True, exist_ok=True)
                save_path = save_dir / filename

                # 기존 파일이 있으면 백업
                if save_path.exists():
                    backup_dir = Path(f'{settings.WWW_STATIC_PDF_DIR}/comparisonTable_backup')
                    backup_dir.mkdir(parents=True, exist_ok=True)
                    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                    base_name = filename.rsplit('.', 1)[0]  # 확장자 제외
                    backup_path = backup_dir / f"{base_name}_{timestamp}{file_ext}"
                    shutil.copy2(save_path, backup_path)
                    logger.info(f"Backed up existing file to: {backup_path}")

                # 파일 저장
                content = await comparison_file.read()
                with open(save_path, 'wb') as f:
                    f.write(content)

                logger.info(f"Saved comparison table to: {save_path}")

                # DB에 경로 저장 (상대 경로)
                relative_path = f"comparisonTable/{filename}"
                cur.execute("""
                    UPDATE wz_rule
                    SET wzFileComparison = %s,
                        wzModifiedBy = %s
                    WHERE wzruleseq = %s
                """, (relative_path, user.get('username'), rule_id))

                conn.commit()
                logger.info(f"Updated wzFileComparison in database: {relative_path}")

                return {
                    "success": True,
                    "message": "신구대비표가 업로드되었습니다.",
                    "file_path": relative_path,
                    "filename": filename,
                    "rule_id": rule_id,
                    "wzruleid": wzruleid
                }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error uploading comparison table: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@public_router.get("/comparison-table/{rule_id}")
async def get_comparison_table(
    rule_id: int
):
    """
    신구대비표 파일 정보 조회 (공개 API - 사용자 화면에서 접근)

    Args:
        rule_id: 규정 ID (wzRuleSeq)

    Returns:
        신구대비표 파일 정보
    """
    try:
        db_config = {
            'host': settings.DB_HOST,
            'port': settings.DB_PORT,
            'database': settings.DB_NAME,
            'user': settings.DB_USER,
            'password': settings.DB_PASSWORD,
        }

        db_manager = DatabaseConnectionManager(**db_config)

        with db_manager.get_connection() as conn:
            with conn.cursor() as cur:
                cur.execute("""
                    SELECT wzFileComparison, wzRuleId, wzPubNo, wzLastRevDate
                    FROM wz_rule
                    WHERE wzruleseq = %s
                """, (rule_id,))

                rule = cur.fetchone()
                if not rule:
                    raise HTTPException(status_code=404, detail="규정을 찾을 수 없습니다.")

                wzfilecomparison, wzruleid, wzpubno, wzlastrevdate = rule

                if not wzfilecomparison:
                    return {
                        "success": False,
                        "message": "신구대비표가 등록되지 않았습니다.",
                        "has_file": False,
                        "file_exists": False,
                        "wzRuleSeq": rule_id
                    }

                # 파일 존재 확인
                file_path = Path(settings.WWW_STATIC_PDF_DIR) / wzfilecomparison
                file_exists = file_path.exists()

                return {
                    "success": True,
                    "has_file": True,
                    "file_exists": file_exists,
                    "file_path": wzfilecomparison,
                    "full_path": str(file_path) if file_exists else None,
                    "wzRuleSeq": rule_id,
                    "wzRuleId": wzruleid,
                    "wzPubNo": wzpubno,
                    "revision_date": wzlastrevdate
                }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting comparison table info: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@public_router.get("/comparison-table/{rule_id}/download")
async def download_comparison_table(rule_id: int, request: Request):
    """
    신구대비표 파일 다운로드

    Args:
        rule_id: 규정 ID (wzRuleSeq)
        request: FastAPI Request 객체

    Returns:
        FileResponse
    """
    from fastapi.responses import FileResponse

    try:
        db_config = {
            'host': settings.DB_HOST,
            'port': settings.DB_PORT,
            'database': settings.DB_NAME,
            'user': settings.DB_USER,
            'password': settings.DB_PASSWORD,
        }

        db_manager = DatabaseConnectionManager(**db_config)

        with db_manager.get_connection() as conn:
            with conn.cursor() as cur:
                cur.execute("""
                    SELECT wzFileComparison, wzName, wzPubNo
                    FROM wz_rule
                    WHERE wzruleseq = %s
                """, (rule_id,))

                rule = cur.fetchone()
                if not rule:
                    raise HTTPException(status_code=404, detail="규정을 찾을 수 없습니다.")

                wzfilecomparison, wzname, wzpubno = rule

                if not wzfilecomparison:
                    raise HTTPException(status_code=404, detail="신구대비표가 등록되지 않았습니다.")

                # 파일 경로
                file_path = Path(settings.WWW_STATIC_PDF_DIR) / wzfilecomparison
                if not file_path.exists():
                    raise HTTPException(status_code=404, detail="파일이 존재하지 않습니다.")

                # 확장자 및 media_type 결정
                file_ext = file_path.suffix.lower()
                media_types = {
                    '.pdf': 'application/pdf',
                    '.hwpx': 'application/haansofthwpx',
                    '.hwp': 'application/x-hwp',
                    '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    '.doc': 'application/msword',
                    '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                    '.xls': 'application/vnd.ms-excel',
                }
                media_type = media_types.get(file_ext, 'application/octet-stream')

                # 다운로드 파일명
                download_filename = f"{wzname}_신구대비표{file_ext}" if wzname else file_path.name

                # 다운로드 이력 기록
                try:
                    ip_address = request.client.host if request.client else None
                    user_agent = request.headers.get('user-agent', '')
                    referer = request.headers.get('referer', '')
                    session_id = request.cookies.get('session_id', '')

                    query_download_logs.insert_download_log(
                        rule_id=rule_id,
                        rule_name=wzname,
                        rule_pubno=wzpubno,
                        file_type='comparison',
                        file_name=download_filename,
                        ip_address=ip_address,
                        user_agent=user_agent,
                        referer=referer,
                        session_id=session_id,
                        user_id=None
                    )
                except Exception as log_error:
                    logger.warning(f"Download log failed (ignored): {log_error}")

                return FileResponse(
                    path=str(file_path),
                    filename=download_filename,
                    media_type=media_type
                )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error downloading comparison table: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@public_router.get("/rule-file/{rule_id}/download/{file_type}")
async def download_rule_file(rule_id: int, file_type: str, request: Request):
    """
    규정 파일(PDF, DOCX) 다운로드

    Args:
        rule_id: 규정 ID (wzRuleSeq)
        file_type: 파일 종류 (pdf, docx)
        request: FastAPI Request 객체

    Returns:
        FileResponse
    """
    from fastapi.responses import FileResponse

    try:
        if file_type not in ['pdf', 'docx']:
            raise HTTPException(status_code=400, detail="잘못된 파일 타입입니다. (pdf, docx)")

        db_config = {
            'host': settings.DB_HOST,
            'port': settings.DB_PORT,
            'database': settings.DB_NAME,
            'user': settings.DB_USER,
            'password': settings.DB_PASSWORD,
        }

        db_manager = DatabaseConnectionManager(**db_config)

        with db_manager.get_connection() as conn:
            with conn.cursor() as cur:
                cur.execute("""
                    SELECT wzFilePdf, wzFileDocx, wzName, wzPubNo, wzLastRevDate
                    FROM wz_rule
                    WHERE wzruleseq = %s
                """, (rule_id,))

                rule = cur.fetchone()
                if not rule:
                    raise HTTPException(status_code=404, detail="규정을 찾을 수 없습니다.")

                wzfilepdf, wzfiledocx, wzname, wzpubno, wzlastrevdate = rule

                # 파일 경로 결정
                if file_type == 'pdf':
                    file_path_str = wzfilepdf
                    media_type = "application/pdf"
                    ext = ".pdf"
                else:  # docx
                    file_path_str = wzfiledocx
                    media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    ext = ".docx"

                if not file_path_str:
                    raise HTTPException(status_code=404, detail=f"{file_type.upper()} 파일이 등록되지 않았습니다.")

                # 다양한 경로 형식 처리
                file_path = None
                possible_paths = []

                if file_path_str.startswith('/'):
                    # 절대 경로
                    possible_paths.append(Path(file_path_str))
                elif file_path_str.startswith('fastapi/'):
                    # fastapi/applib/... 형식
                    possible_paths.append(Path(settings.BASE_DIR) / file_path_str)
                elif file_path_str.startswith('applib/'):
                    # applib/... 형식
                    possible_paths.append(Path(settings.FASTAPI_DIR) / file_path_str)

                # 파일명만 있거나 경로가 없는 경우 - 여러 위치에서 검색
                file_name = Path(file_path_str).name
                if file_type == 'pdf':
                    possible_paths.extend([
                        Path(settings.PDF_DIR) / file_name,
                        Path(settings.WWW_STATIC_PDF_DIR) / file_name,
                        Path(f'{settings.WWW_STATIC_PDF_DIR}/print') / file_name,
                    ])
                else:  # docx
                    possible_paths.extend([
                        Path(settings.DOCX_DIR) / file_name,
                        Path(settings.MERGE_JSON_DIR) / file_name,
                    ])

                # 존재하는 파일 찾기
                for p in possible_paths:
                    if p.exists():
                        file_path = p
                        break

                if not file_path or not file_path.exists():
                    logger.error(f"File not found. Tried paths: {possible_paths}")
                    raise HTTPException(status_code=404, detail="파일이 존재하지 않습니다.")

                # 다운로드 파일명
                date_str = (wzlastrevdate or '').replace('.', '').replace('-', '')[:8]
                if wzname and date_str:
                    download_filename = f"{wzname}_{date_str}{ext}"
                elif wzname:
                    download_filename = f"{wzname}{ext}"
                else:
                    download_filename = file_path.name

                # 다운로드 이력 기록
                try:
                    ip_address = request.client.host if request.client else None
                    user_agent = request.headers.get('user-agent', '')
                    referer = request.headers.get('referer', '')
                    session_id = request.cookies.get('session_id', '')

                    query_download_logs.insert_download_log(
                        rule_id=rule_id,
                        rule_name=wzname,
                        rule_pubno=wzpubno,
                        file_type=file_type,
                        file_name=download_filename,
                        ip_address=ip_address,
                        user_agent=user_agent,
                        referer=referer,
                        session_id=session_id,
                        user_id=None
                    )
                except Exception as log_error:
                    logger.warning(f"Download log failed (ignored): {log_error}")

                return FileResponse(
                    path=str(file_path),
                    filename=download_filename,
                    media_type=media_type
                )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error downloading rule file: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


# ============ 수정 이력파일 API ============

@router.post("/upload-history-file/{rule_id}")
async def upload_history_file(
    rule_id: int,
    history_file: UploadFile = File(...),
    user: Dict[str, Any] = Depends(get_current_user)
):
    """
    수정 이력파일 업로드 (여러 파일 지원, 모든 파일 형식 허용)

    Args:
        rule_id: 규정 ID (wzRuleSeq)
        history_file: 수정 이력파일 (모든 파일 형식)
        user: 현재 사용자

    Returns:
        업로드 결과 및 파일 경로
    """
    try:
        logger.info(f"Uploading history file for rule {rule_id}: {history_file.filename}")

        # DB에서 규정 정보 조회
        db_config = {
            'host': settings.DB_HOST,
            'port': settings.DB_PORT,
            'database': settings.DB_NAME,
            'user': settings.DB_USER,
            'password': settings.DB_PASSWORD,
        }

        db_manager = DatabaseConnectionManager(**db_config)

        with db_manager.get_connection() as conn:
            with conn.cursor() as cur:
                cur.execute("""
                    SELECT wzRuleSeq, wzRuleId, wzPubNo, wzLastRevDate, wzNewFlag, wzFileHistory
                    FROM wz_rule
                    WHERE wzruleseq = %s
                """, (rule_id,))

                rule = cur.fetchone()
                if not rule:
                    raise HTTPException(status_code=404, detail="규정을 찾을 수 없습니다.")

                wzruleseq, wzruleid, wzpubno, wzlastrevdate, wznewflag, existing_history = rule

                # 기존 파일 목록 파싱 (JSON 배열 또는 단일 파일명)
                existing_files = []
                if existing_history:
                    try:
                        existing_files = json.loads(existing_history)
                        if not isinstance(existing_files, list):
                            existing_files = [existing_history]
                    except json.JSONDecodeError:
                        # 기존 단일 파일명인 경우
                        existing_files = [existing_history]

                # 파일명 생성: wzRuleId_wzRuleSeq_개정일_history_순번.확장자
                revision_date_str = wzlastrevdate.replace('.', '').replace('-', '') if wzlastrevdate else 'nodate'
                timestamp = datetime.now().strftime('%H%M%S')
                # 원본 파일 확장자 유지
                original_ext = Path(history_file.filename).suffix.lower() if history_file.filename else '.pdf'
                filename = f"{wzruleid}_{wzruleseq}_{revision_date_str}_history_{timestamp}{original_ext}"

                # 저장 경로
                save_dir = Path(f'{settings.WWW_STATIC_DIR}/history')
                save_dir.mkdir(parents=True, exist_ok=True)
                save_path = save_dir / filename

                # 파일 저장
                content = await history_file.read()
                with open(save_path, 'wb') as f:
                    f.write(content)

                logger.info(f"Saved history file to: {save_path}")

                # 새 파일을 목록에 추가
                existing_files.append(filename)

                # DB에 JSON 배열로 저장
                files_json = json.dumps(existing_files, ensure_ascii=False)
                cur.execute("""
                    UPDATE wz_rule
                    SET wzFileHistory = %s,
                        wzModifiedBy = %s
                    WHERE wzruleseq = %s
                """, (files_json, user.get('username'), rule_id))

                conn.commit()
                logger.info(f"Updated wzFileHistory in database: {files_json}")

                return {
                    "success": True,
                    "message": "수정 이력파일이 업로드되었습니다.",
                    "file_path": filename,
                    "filename": filename,
                    "rule_id": rule_id,
                    "wzruleid": wzruleid,
                    "total_files": len(existing_files),
                    "files": existing_files
                }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error uploading history file: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@public_router.get("/history-file/{rule_id}")
async def get_history_file_info(
    rule_id: int
):
    """
    수정 이력파일 정보 조회 (공개 API) - 여러 파일 지원

    Args:
        rule_id: 규정 ID (wzRuleSeq)

    Returns:
        수정 이력파일 목록 정보
    """
    try:
        db_config = {
            'host': settings.DB_HOST,
            'port': settings.DB_PORT,
            'database': settings.DB_NAME,
            'user': settings.DB_USER,
            'password': settings.DB_PASSWORD,
        }

        db_manager = DatabaseConnectionManager(**db_config)

        with db_manager.get_connection() as conn:
            with conn.cursor() as cur:
                cur.execute("""
                    SELECT wzFileHistory, wzRuleId, wzPubNo, wzLastRevDate
                    FROM wz_rule
                    WHERE wzruleseq = %s
                """, (rule_id,))

                rule = cur.fetchone()
                if not rule:
                    raise HTTPException(status_code=404, detail="규정을 찾을 수 없습니다.")

                wzfilehistory, wzruleid, wzpubno, wzlastrevdate = rule

                if not wzfilehistory:
                    return {
                        "success": True,
                        "has_file": False,
                        "files": [],
                        "total_files": 0,
                        "wzruleid": wzruleid,
                        "wzpubno": wzpubno,
                        "wzlastrevdate": wzlastrevdate
                    }

                # JSON 배열 파싱 또는 단일 파일명 처리
                try:
                    files = json.loads(wzfilehistory)
                    if not isinstance(files, list):
                        files = [wzfilehistory]
                except json.JSONDecodeError:
                    files = [wzfilehistory]

                # 실제 존재하는 파일만 필터링
                save_dir = Path(f'{settings.WWW_STATIC_DIR}/history')
                valid_files = []
                for f in files:
                    file_path = save_dir / f
                    if file_path.exists():
                        file_size = file_path.stat().st_size
                        valid_files.append({
                            "filename": f,
                            "download_url": f"/static/history/{f}",
                            "size": file_size,
                            "size_str": f"{file_size / 1024:.1f} KB" if file_size < 1024*1024 else f"{file_size / 1024 / 1024:.1f} MB"
                        })

                return {
                    "success": True,
                    "has_file": len(valid_files) > 0,
                    "files": valid_files,
                    "total_files": len(valid_files),
                    "wzruleid": wzruleid,
                    "wzpubno": wzpubno,
                    "wzlastrevdate": wzlastrevdate
                }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting history file info: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@public_router.get("/history-file/{rule_id}/download")
async def download_history_file(
    rule_id: int,
    request: Request,
    filename: str = None
):
    """
    수정 이력파일 다운로드

    Args:
        rule_id: 규정 ID (wzRuleSeq)
        request: FastAPI Request 객체
        filename: 다운로드할 파일명 (없으면 첫 번째 파일)

    Returns:
        파일 다운로드 응답
    """
    from fastapi.responses import FileResponse

    try:
        db_config = {
            'host': settings.DB_HOST,
            'port': settings.DB_PORT,
            'database': settings.DB_NAME,
            'user': settings.DB_USER,
            'password': settings.DB_PASSWORD,
        }

        db_manager = DatabaseConnectionManager(**db_config)

        with db_manager.get_connection() as conn:
            with conn.cursor() as cur:
                cur.execute("""
                    SELECT wzFileHistory, wzName
                    FROM wz_rule
                    WHERE wzruleseq = %s
                """, (rule_id,))

                rule = cur.fetchone()
                if not rule:
                    raise HTTPException(status_code=404, detail="규정을 찾을 수 없습니다.")

                wzfilehistory, wzname = rule

                if not wzfilehistory:
                    raise HTTPException(status_code=404, detail="수정 이력파일이 등록되지 않았습니다.")

                # JSON 배열 파싱 또는 단일 파일명 처리
                try:
                    files = json.loads(wzfilehistory)
                    if not isinstance(files, list):
                        files = [wzfilehistory]
                except json.JSONDecodeError:
                    files = [wzfilehistory]

                # 파일명이 지정되지 않으면 첫 번째 파일
                target_filename = filename if filename else (files[0] if files else None)
                if not target_filename:
                    raise HTTPException(status_code=404, detail="다운로드할 파일이 없습니다.")

                # 보안: 지정된 파일이 목록에 있는지 확인
                if target_filename not in files:
                    raise HTTPException(status_code=404, detail="파일을 찾을 수 없습니다.")

                file_path = Path(f'{settings.WWW_STATIC_DIR}/history') / target_filename
                if not file_path.exists():
                    raise HTTPException(status_code=404, detail="파일이 존재하지 않습니다.")

                # 원본 파일 확장자 추출
                file_ext = Path(target_filename).suffix.lower()  # .pdf, .hwpx, .docx 등

                # 확장자별 media_type 결정
                media_types = {
                    '.pdf': 'application/pdf',
                    '.hwpx': 'application/haansofthwpx',
                    '.hwp': 'application/x-hwp',
                    '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    '.doc': 'application/msword',
                    '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                    '.xls': 'application/vnd.ms-excel',
                    '.pptx': 'application/vnd.openxmlformats-officedocument.presentationml.presentation',
                    '.ppt': 'application/vnd.ms-powerpoint',
                }
                media_type = media_types.get(file_ext, 'application/octet-stream')

                # 다운로드 파일명 생성 (원본 확장자 유지)
                if wzname:
                    history_num = len(files) if len(files) > 1 else ''
                    download_filename = f"{wzname}_수정이력_{history_num}{file_ext}" if history_num else f"{wzname}_수정이력{file_ext}"
                else:
                    download_filename = target_filename

                # 다운로드 이력 기록
                try:
                    ip_address = request.client.host if request.client else None
                    user_agent = request.headers.get('user-agent', '')
                    referer = request.headers.get('referer', '')
                    session_id = request.cookies.get('session_id', '')

                    query_download_logs.insert_download_log(
                        rule_id=rule_id,
                        rule_name=wzname,
                        rule_pubno=None,
                        file_type='history',
                        file_name=download_filename,
                        ip_address=ip_address,
                        user_agent=user_agent,
                        referer=referer,
                        session_id=session_id,
                        user_id=None
                    )
                except Exception as log_error:
                    logger.warning(f"Download log failed (ignored): {log_error}")

                return FileResponse(
                    path=str(file_path),
                    filename=download_filename,
                    media_type=media_type
                )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error downloading history file: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/history-file/{rule_id}")
async def delete_history_file(
    rule_id: int,
    filename: str = None,
    user: Dict[str, Any] = Depends(get_current_user)
):
    """
    수정 이력파일 삭제 (특정 파일 또는 전체)

    Args:
        rule_id: 규정 ID (wzRuleSeq)
        filename: 삭제할 파일명 (없으면 전체 삭제)
        user: 현재 사용자

    Returns:
        삭제 결과
    """
    try:
        db_config = {
            'host': settings.DB_HOST,
            'port': settings.DB_PORT,
            'database': settings.DB_NAME,
            'user': settings.DB_USER,
            'password': settings.DB_PASSWORD,
        }

        db_manager = DatabaseConnectionManager(**db_config)

        with db_manager.get_connection() as conn:
            with conn.cursor() as cur:
                cur.execute("""
                    SELECT wzFileHistory FROM wz_rule WHERE wzruleseq = %s
                """, (rule_id,))

                rule = cur.fetchone()
                if not rule or not rule[0]:
                    raise HTTPException(status_code=404, detail="수정 이력파일이 없습니다.")

                wzfilehistory = rule[0]

                # JSON 배열 파싱 또는 단일 파일명 처리
                try:
                    files = json.loads(wzfilehistory)
                    if not isinstance(files, list):
                        files = [wzfilehistory]
                except json.JSONDecodeError:
                    files = [wzfilehistory]

                save_dir = Path(f'{settings.WWW_STATIC_DIR}/history')
                backup_dir = Path(f'{settings.WWW_STATIC_DIR}/history/backup')
                backup_dir.mkdir(parents=True, exist_ok=True)

                if filename:
                    # 특정 파일만 삭제
                    if filename not in files:
                        raise HTTPException(status_code=404, detail="파일을 찾을 수 없습니다.")

                    file_path = save_dir / filename
                    if file_path.exists():
                        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                        backup_path = backup_dir / f"{filename.replace('.pdf', '')}_{timestamp}.pdf"
                        shutil.move(str(file_path), str(backup_path))
                        logger.info(f"Moved history file to backup: {backup_path}")

                    # 목록에서 제거
                    files.remove(filename)

                    # DB 업데이트
                    if files:
                        files_json = json.dumps(files, ensure_ascii=False)
                        cur.execute("""
                            UPDATE wz_rule
                            SET wzFileHistory = %s,
                                wzModifiedBy = %s
                            WHERE wzruleseq = %s
                        """, (files_json, user.get('username'), rule_id))
                    else:
                        cur.execute("""
                            UPDATE wz_rule
                            SET wzFileHistory = NULL,
                                wzModifiedBy = %s
                            WHERE wzruleseq = %s
                        """, (user.get('username'), rule_id))

                    conn.commit()

                    return {
                        "success": True,
                        "message": f"파일 '{filename}'이 삭제되었습니다.",
                        "remaining_files": files,
                        "total_files": len(files)
                    }
                else:
                    # 전체 삭제
                    for f in files:
                        file_path = save_dir / f
                        if file_path.exists():
                            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                            backup_path = backup_dir / f"{f.replace('.pdf', '')}_{timestamp}.pdf"
                            shutil.move(str(file_path), str(backup_path))
                            logger.info(f"Moved history file to backup: {backup_path}")

                    # DB 업데이트
                    cur.execute("""
                        UPDATE wz_rule
                        SET wzFileHistory = NULL,
                            wzModifiedBy = %s
                        WHERE wzruleseq = %s
                    """, (user.get('username'), rule_id))

                    conn.commit()

                    return {
                        "success": True,
                        "message": f"수정 이력파일 {len(files)}개가 모두 삭제되었습니다.",
                        "remaining_files": [],
                        "total_files": 0
                    }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting history file: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/save-edited")
async def save_edited_content(
    request: Dict[str, Any],
    user: Dict[str, Any] = Depends(get_current_user)
):
    """편집된 내용을 applib/edited 폴더에 저장"""
    try:
        # 요청 데이터 파싱
        rule_id = request.get('rule_id')
        rule_title = request.get('title', 'untitled')
        content_type = request.get('content_type', 'article')  # article or json
        content = request.get('content', '')

        # applib/edited 폴더 경로
        edited_folder = os.path.join(
            os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
            'applib', 'edited'
        )
        os.makedirs(edited_folder, exist_ok=True)

        # 타임스탬프 생성
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        # 파일명 생성 (특수문자 제거)
        safe_title = "".join(c for c in rule_title if c.isalnum() or c in (' ', '-', '_')).rstrip()
        safe_title = safe_title.replace(' ', '_')

        # 파일 저장
        if content_type == 'article':
            # 조문 내용 저장 (텍스트 형식)
            filename = f"{safe_title}_{rule_id}_{timestamp}.txt"
            filepath = os.path.join(edited_folder, filename)
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(content)
        else:
            # JSON 형식으로 저장
            filename = f"{safe_title}_{rule_id}_{timestamp}.json"
            filepath = os.path.join(edited_folder, filename)

            # content가 문자열인 경우 JSON 파싱 시도
            if isinstance(content, str):
                try:
                    content = json.loads(content)
                except:
                    pass

            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(content, f, ensure_ascii=False, indent=2)

        logger.info(f"Edited content saved: {filepath}")

        return {
            "success": True,
            "message": "편집 내용이 저장되었습니다.",
            "filename": filename,
            "path": filepath
        }

    except Exception as e:
        logger.error(f"Error saving edited content: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/merge-json")
async def merge_json_files(
    pdf_json_path: str = Form(...),
    docx_json_path: str = Form(...),
    rule_id: int = Form(...),
    user: Dict[str, Any] = Depends(get_current_user)
) -> Dict[str, Any]:
    """
    PDF와 DOCX JSON 파일을 병합

    Args:
        pdf_json_path: PDF JSON 파일 경로
        docx_json_path: DOCX JSON 파일 경로
        rule_id: 규정 ID
        user: 현재 사용자

    Returns:
        병합된 JSON 데이터
    """
    import time
    from .timescaledb_manager_v2 import DatabaseConnectionManager
    request_start = time.time()

    try:
        logger.info(f"=== Starting JSON merge ===")
        logger.info(f"[TIMING] merge_json_files started")
        logger.info(f"PDF path received: {pdf_json_path}")
        logger.info(f"DOCX path received: {docx_json_path}")
        logger.info(f"Rule ID: {rule_id}")
        logger.info(f"User: {user.get('username', 'unknown')}")

        # 파일 경로 확인
        if not pdf_json_path or not docx_json_path:
            logger.error(f"Missing file paths - PDF: {pdf_json_path}, DOCX: {docx_json_path}")
            raise HTTPException(status_code=400, detail="PDF 또는 DOCX JSON 경로가 제공되지 않았습니다")

        # 파일 존재 여부 확인
        if not os.path.exists(pdf_json_path):
            logger.error(f"PDF JSON file not found: {pdf_json_path}")
            raise HTTPException(status_code=404, detail=f"PDF JSON 파일을 찾을 수 없습니다: {pdf_json_path}")

        if not os.path.exists(docx_json_path):
            logger.error(f"DOCX JSON file not found: {docx_json_path}")
            raise HTTPException(status_code=404, detail=f"DOCX JSON 파일을 찾을 수 없습니다: {docx_json_path}")

        logger.info(f"Files exist - PDF: {os.path.getsize(pdf_json_path)} bytes, DOCX: {os.path.getsize(docx_json_path)} bytes")

        # merge_json.py import
        try:
            # applib 경로를 sys.path에 추가
            import sys
            sys.path.insert(0, settings.APPLIB_DIR)
            logger.info("Importing JSONMerger from merge_json")
            from merge_json import JSONMerger
            logger.info("JSONMerger imported successfully")
        except ImportError as ie:
            logger.error(f"Failed to import JSONMerger: {ie}")
            raise HTTPException(status_code=500, detail=f"JSONMerger 모듈을 로드할 수 없습니다: {ie}")

        # JSON 병합 수행
        merge_start = time.time()
        logger.info("Creating JSONMerger instance")
        merger = JSONMerger(
            pdf_json_path=pdf_json_path,
            docx_json_path=docx_json_path
        )

        # JSON 파일 로드
        load_start = time.time()
        logger.info("Loading JSON files")
        load_result = merger.load_json_files()
        load_elapsed = time.time() - load_start
        logger.info(f"[TIMING] JSON files loaded in {load_elapsed:.2f}s")
        if not load_result:
            logger.error("Failed to load JSON files")
            logger.error(f"PDF data loaded: {merger.pdf_data is not None}")
            logger.error(f"DOCX data loaded: {merger.docx_data is not None}")
            raise HTTPException(status_code=400, detail="JSON 파일 로드 실패")

        logger.info("JSON files loaded successfully")
        logger.info(f"PDF articles count: {len(merger.pdf_data.get('조문내용', []))}")
        logger.info(f"DOCX articles count: {len(merger.docx_data.get('조문내용', []))}")

        # 병합 수행
        merge_process_start = time.time()
        logger.info("Starting merge process")
        # merge_regulation 메서드 사용 (merge가 아님)
        merged_data = merger.merge_regulation()
        merge_process_elapsed = time.time() - merge_process_start
        logger.info(f"[TIMING] merge_regulation completed in {merge_process_elapsed:.2f}s")

        if merge_process_elapsed > 20:
            logger.warning(f"[BOTTLENECK] merge_regulation took {merge_process_elapsed:.2f}s (>20s threshold)")

        if not merged_data:
            logger.error("Merge returned empty data")
            raise HTTPException(status_code=500, detail="병합 결과가 비어있습니다")

        logger.info(f"Merge completed successfully")
        logger.info(f"Merged articles count: {len(merged_data.get('조문내용', []))}")

        # wzruleid 조회
        wzruleid = 0
        db_config = {
            'host': settings.DB_HOST,
            'port': settings.DB_PORT,
            'database': settings.DB_NAME,
            'user': settings.DB_USER,
            'password': settings.DB_PASSWORD,
        }
        db_manager = DatabaseConnectionManager(**db_config)
        with db_manager.get_connection() as conn:
            with conn.cursor() as cur:
                cur.execute("SELECT wzruleid FROM wz_rule WHERE wzruleseq = %s", (rule_id,))
                result = cur.fetchone()
                wzruleid = result[0] if result and result[0] else 0

        logger.info(f"Retrieved wzruleid={wzruleid} for rule_id={rule_id}")

        # 병합된 JSON 저장
        merged_dir = settings.MERGE_JSON_DIR

        logger.info(f"Creating merge directory: {merged_dir}")
        os.makedirs(merged_dir, exist_ok=True)

        # wzruleid 기반 파일명 생성
        merged_filename = f"{wzruleid}.json"
        merged_filepath = os.path.join(merged_dir, merged_filename)

        logger.info(f"Saving merged JSON to: {merged_filepath}")
        with open(merged_filepath, 'w', encoding='utf-8') as f:
            json.dump(merged_data, f, ensure_ascii=False, indent=2)

        logger.info(f"Merged JSON saved successfully: {merged_filepath}")

        # www/static/file에 파일 복사
        www_dir = settings.WWW_STATIC_FILE_DIR
        os.makedirs(www_dir, exist_ok=True)
        www_filepath = os.path.join(www_dir, merged_filename)

        try:
            shutil.copy2(merged_filepath, www_filepath)
            logger.info(f"Copied JSON to www/static/file: {www_filepath}")
        except Exception as copy_error:
            logger.error(f"Failed to copy to www/static/file: {copy_error}")

        # DB에 저장할 상대 경로 (www/static/file/{wzruleid}.json)
        db_json_path = f"www/static/file/{wzruleid}.json"

        # 텍스트 추출 (편집용)
        text_content = ""
        if '조문내용' in merged_data:
            for article in merged_data['조문내용']:
                if isinstance(article, dict):
                    number = article.get('번호', '')
                    content = article.get('내용', '')
                    if number:
                        text_content += f"{number} {content}\n"
                    else:
                        text_content += f"{content}\n"

        logger.info(f"Text content extracted: {len(text_content)} characters")

        # DB에 JSON 파일 경로 저장
        try:
            db_config = {
                'host': settings.DB_HOST,
                'port': settings.DB_PORT,
                'database': settings.DB_NAME,
                'user': settings.DB_USER,
                'password': settings.DB_PASSWORD,
            }

            db_manager = DatabaseConnectionManager(**db_config)

            # wzFileJson 컬럼 업데이트
            update_query = """
                UPDATE wz_rule
                SET wzFileJson = %s
                WHERE wzruleseq = %s
            """

            # RETURNING 제거하여 rowcount 반환받도록 수정
            rowcount = db_manager.execute_query(update_query, (db_json_path, rule_id), commit=True)

            if rowcount and rowcount > 0:
                logger.info(f"Updated wzFileJson for rule {rule_id}: {db_json_path}")
            else:
                logger.warning(f"No record updated for rule {rule_id}")

            # 색인 작업 트리거 (백그라운드)
            try:
                import threading
                from .indexing_service import index_single_regulation

                # 규정명 조회
                with db_manager.get_connection() as conn:
                    with conn.cursor() as cur:
                        cur.execute("SELECT wzname FROM wz_rule WHERE wzruleseq = %s", (rule_id,))
                        result = cur.fetchone()
                        rule_name = result[0] if result else f"Rule_{rule_id}"

                # 백그라운드에서 색인 실행
                threading.Thread(
                    target=index_single_regulation,
                    args=(db_manager, rule_id, rule_name, db_json_path),
                    daemon=True
                ).start()
                logger.info(f"Indexing triggered for rule {rule_id}")
            except Exception as index_error:
                logger.warning(f"Failed to trigger indexing: {index_error}")
                # 색인 실패해도 병합은 성공으로 처리

        except Exception as db_error:
            logger.error(f"Failed to update wzFileJson in DB: {db_error}")
            # DB 업데이트 실패해도 병합은 성공으로 처리

        total_elapsed = time.time() - request_start
        logger.info(f"[TIMING] merge_json_files total time: {total_elapsed:.2f}s")
        logger.info("=== JSON merge completed successfully ===")

        if total_elapsed > 60:
            logger.warning(f"[BOTTLENECK] Total merge request took {total_elapsed:.2f}s (>60s threshold)")

        return {
            "success": True,
            "message": "JSON 파일이 성공적으로 병합되었습니다.",
            "merged_data": merged_data,
            "text_content": text_content,
            "filepath": merged_filepath,
            "filename": merged_filename,
            "json_path": merged_filepath  # 클라이언트에서 사용하기 위해 추가
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error during JSON merge: {e}")
        logger.error(f"Error type: {type(e)}")
        logger.error(f"Error details:", exc_info=True)
        raise HTTPException(status_code=500, detail=f"병합 중 오류 발생: {str(e)}")


def sync_json_to_static_file(
    rule_id: int,
    wzpubno: str,
    new_merged_json_path: str,
    mode: str = "revision"
) -> bool:
    """
    개정 시 static/file 디렉토리 동기화

    Args:
        rule_id: 규정 ID
        wzpubno: 규정 공포번호 (파일명 검색용)
        new_merged_json_path: 새로 생성된 병합 JSON 경로
        mode: "revision" 또는 "new"

    Returns:
        bool: 성공 여부
    """
    try:
        # 디렉토리 경로 설정
        static_file_dir = Path(settings.WWW_STATIC_FILE_DIR)
        static_file_old_dir = Path(f"{settings.WWW_STATIC_DIR}/file_old")
        static_file_old_dir.mkdir(parents=True, exist_ok=True)

        new_json_path = Path(new_merged_json_path)

        if mode == "revision":
            # Step 1: 기존 JSON 파일 찾기 (wzpubno 기반)
            # 예: wzpubno가 "1.1.1"이면 "merged_1.1.1.*" 패턴의 파일 검색
            # 공백 포함하여 정확한 매칭 (예: "merged_1.1.1. 정확한...")
            pattern = f"merged_{wzpubno} *.json"
            old_json_files = list(static_file_dir.glob(pattern))

            # 패턴 매칭 실패 시 더 넓은 패턴 시도
            if not old_json_files:
                pattern = f"merged_{wzpubno}*.json"
                old_json_files = list(static_file_dir.glob(pattern))

            if old_json_files:
                # 가장 최근 파일 선택 (여러 개 있을 경우)
                old_json_file = sorted(old_json_files, key=lambda f: f.stat().st_mtime)[-1]

                # 타임스탬프 추가하여 file_old로 복사
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                old_filename = f"{old_json_file.stem}_backup_{timestamp}.json"
                old_dest_path = static_file_old_dir / old_filename

                shutil.copy2(old_json_file, old_dest_path)
                logger.info(f"✅ Moved old JSON to file_old: {old_json_file.name} -> {old_dest_path.name}")

                # 기존 파일 삭제 (static/file에서 제거)
                old_json_file.unlink()
                logger.info(f"✅ Removed old JSON from static/file: {old_json_file.name}")
            else:
                logger.warning(f"⚠️ No existing JSON found in static/file for wzpubno: {wzpubno}")

        elif mode == "edit":
            # 편집 모드: 백업/삭제 없이 덮어쓰기만 수행
            logger.info(f"📝 Edit mode: Will overwrite existing file in static/file (no backup)")
            pass

        # mode == "new": 신규 제정은 기존 파일 체크 불필요

        # Step 2: 새 병합 JSON을 static/file/로 복사
        new_filename = new_json_path.name
        dest_path = static_file_dir / new_filename

        shutil.copy2(new_json_path, dest_path)
        logger.info(f"✅ Copied new JSON to static/file: {new_filename}")

        # Step 3: 이미지 동기화 (wzruleid 기반)
        # wzruleid 추출 (파일명에서 또는 DB에서)
        try:
            # 파일명이 wzruleid.json 형식인 경우
            if new_filename.replace('.json', '').isdigit():
                wzruleid = new_filename.replace('.json', '')
            else:
                # DB에서 wzruleid 조회
                from .timescaledb_manager_v2 import DatabaseConnectionManager
                from settings import settings
                db_config = {
                    'host': settings.DB_HOST,
                    'database': settings.DB_NAME,
                    'user': settings.DB_USER,
                    'password': settings.DB_PASSWORD,
                    'port': settings.DB_PORT
                }
                db_manager = DatabaseConnectionManager(**db_config)
                with db_manager.get_connection() as conn:
                    with conn.cursor() as cursor:
                        cursor.execute("SELECT wzruleid FROM wz_rule WHERE wzruleseq = %s", (rule_id,))
                        result = cursor.fetchone()
                        wzruleid = str(result[0]) if result and result[0] else None

            if wzruleid:
                # 이미지 소스 및 타겟 경로
                source_image_dirs = [
                    Path(f"{settings.FASTAPI_DIR}/static/extracted_images/{wzruleid}"),
                    Path(f"{settings.APPLIB_DIR}/static/extracted_images/{wzruleid}")
                ]
                target_image_dir = Path(f"{settings.EXTRACTED_IMAGES_DIR}/{wzruleid}")

                # 타겟 디렉토리 생성
                target_image_dir.mkdir(parents=True, exist_ok=True)

                # 이미지 파일 동기화
                images_synced = 0
                for source_dir in source_image_dirs:
                    if source_dir.exists():
                        for img_file in source_dir.glob('*'):
                            if img_file.is_file() and not img_file.name.startswith('.'):
                                target_file = target_image_dir / img_file.name
                                shutil.copy2(img_file, target_file)
                                images_synced += 1
                        break  # 첫 번째 존재하는 소스에서만 복사

                if images_synced > 0:
                    logger.info(f"✅ Synced {images_synced} images to www/static/extracted_images/{wzruleid}")
                else:
                    logger.info(f"ℹ️ No images found for wzruleid {wzruleid}")
            else:
                logger.warning(f"⚠️ Could not determine wzruleid for image sync")

        except Exception as img_sync_error:
            logger.warning(f"⚠️ Image sync failed (continuing): {img_sync_error}")

        return True

    except Exception as e:
        logger.error(f"❌ Failed to sync JSON to static/file: {e}")
        logger.error(f"Error details:", exc_info=True)
        return False


def run_json_merge_and_summary(rule_id: int):
    """
    백그라운드에서 JSON_ALL.py와 create_summary.py 실행

    Args:
        rule_id: 규정 ID (로깅용)
    """
    try:
        logger.info(f"[Background] Starting JSON merge and summary for rule {rule_id}")

        import subprocess
        import time

        # 약간의 지연을 둬서 연속 편집 시 부하 감소
        time.sleep(2)

        applib_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'applib')

        # JSON_ALL.py 실행
        json_all_path = os.path.join(applib_path, "JSON_ALL.py")

        logger.info(f"[Background] Executing JSON_ALL.py")
        start_time = time.time()

        json_all_result = subprocess.run(
            ["python3", json_all_path],
            capture_output=True,
            text=True,
            cwd=applib_path,
            timeout=30  # 30초 타임아웃
        )

        json_all_time = time.time() - start_time

        if json_all_result.returncode == 0:
            logger.info(f"[Background] JSON_ALL.py completed in {json_all_time:.2f}s")
        else:
            logger.warning(f"[Background] JSON_ALL.py failed: {json_all_result.stderr}")
            # 실패해도 계속 진행

        # create_summary.py 실행
        summary_path = os.path.join(applib_path, "create_summary.py")

        logger.info(f"[Background] Executing create_summary.py")
        start_time = time.time()

        summary_result = subprocess.run(
            ["python3", summary_path],
            capture_output=True,
            text=True,
            cwd=applib_path,
            timeout=30  # 30초 타임아웃
        )

        summary_time = time.time() - start_time

        if summary_result.returncode == 0:
            logger.info(f"[Background] create_summary.py completed in {summary_time:.2f}s")
        else:
            logger.warning(f"[Background] create_summary.py failed: {summary_result.stderr}")

        logger.info(f"[Background] JSON merge and summary completed for rule {rule_id}")

    except subprocess.TimeoutExpired:
        logger.error(f"[Background] Timeout during JSON merge for rule {rule_id}")
    except Exception as e:
        logger.error(f"[Background] Error during JSON merge for rule {rule_id}: {e}")


def _parse_content_to_articles(content: str) -> list:
    """
    텍스트 내용을 조문 단위로 파싱
    """
    import re
    articles = []
    lines = content.strip().split('\n')

    for line in lines:
        if line.strip():
            # 제목 패턴 매칭
            match = re.match(r'^(제\d+조|제\d+조\s*\([^)]+\)|①|②|③|④|⑤|\d+\.)', line)
            if match:
                articles.append({
                    "번호": match.group(1),
                    "내용": line[match.end():].strip()
                })
            else:
                articles.append({
                    "내용": line.strip()
                })

    return articles


def update_articles_content(articles: list, edited_content: str) -> None:
    """
    기존 조문 구조(seq, 레벨 등)를 유지하면서 내용만 업데이트

    중복 번호 처리 개선: 순차적 매칭 방식 사용

    Args:
        articles: 기존 조문 리스트 (seq, 레벨 포함)
        edited_content: 편집된 텍스트 내용
    """
    import re

    # 편집된 내용을 줄 단위로 파싱
    content_lines = [line.strip() for line in edited_content.strip().split('\n') if line.strip()]

    # 순차적으로 (번호, 내용) 리스트 생성
    content_list = []

    for line in content_lines:
        # 다양한 번호 패턴 매칭
        patterns = [
            r'^(제\d+조)',           # 제1조
            r'^(\d+\.)',             # 1.
            r'^(\d+\))',             # 1)
            r'^(\(\d+\))',           # (1)
            r'^([①-⑳])',            # ①
            r'^([가-하]\.)',         # 가.
            r'^(\([가-하]\))',       # (가)
        ]

        matched = False
        for pattern in patterns:
            match = re.match(pattern, line)
            if match:
                number = match.group(1).strip()
                content_after_number = line[len(match.group(1)):].strip()
                content_list.append((number, content_after_number))
                matched = True
                break

        if not matched and line:
            # 번호가 없는 경우 전체 내용 저장
            content_list.append(('', line))

    # 순차적으로 매칭 (같은 번호가 여러 개 있어도 순서대로 매칭)
    content_idx = 0

    for article in articles:
        if content_idx >= len(content_list):
            break

        article_number = article.get("번호", "").strip()

        # 번호가 일치하면 업데이트하고 다음으로
        if article_number and article_number == content_list[content_idx][0]:
            article["내용"] = content_list[content_idx][1]
            logger.debug(f"Updated article seq={article.get('seq')} {article_number}")
            content_idx += 1
        elif not article_number and content_list[content_idx][0] == '':
            # 번호 없는 조문 매칭
            article["내용"] = content_list[content_idx][1]
            logger.debug(f"Updated article seq={article.get('seq')} (no number)")
            content_idx += 1


@router.post("/save-edited-content")
async def save_edited_content(
    background_tasks: BackgroundTasks,
    rule_id: int = Form(...),
    content: str = Form(...),
    mode: Optional[str] = Form(None),
    is_revision: bool = Form(False),
    revision_date: Optional[str] = Form(None),
    execution_date: Optional[str] = Form(None),
    merged_json_path: Optional[str] = Form(None),
    use_merged_json: Optional[str] = Form(None),
    merged_json_data: Optional[str] = Form(None),
    user: Dict[str, Any] = Depends(get_current_user)
) -> Dict[str, Any]:
    """
    편집된 내용 저장

    Args:
        rule_id: 규정 ID
        content: 편집된 내용
        mode: 모드 (new, edit, revision)
        is_revision: 개정 여부
        revision_date: 개정일 (개정인 경우)
        execution_date: 시행일자 (개정인 경우)
        user: 현재 사용자

    Returns:
        저장 결과
    """
    try:
        logger.info(f"Saving edited content for rule {rule_id}, mode: {mode}, is_revision: {is_revision}")

        # applib 경로 정의
        applib_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'applib')

        # DB 연결
        db_config = {
            'host': settings.DB_HOST,
            'port': settings.DB_PORT,
            'database': settings.DB_NAME,
            'user': settings.DB_USER,
            'password': settings.DB_PASSWORD,
        }

        db_manager = DatabaseConnectionManager(**db_config)

        # 모드에 따른 처리 구분
        if mode == 'revision' or is_revision:
            logger.info(f"Processing revision with dates - Revision: {revision_date}, Execution: {execution_date}")

            # 개정 처리 전 wzNewFlag가 '현행'인 다른 버전의 JSON 파일을 연혁으로 이동
            # (이전 버전을 보존하기 위함)
            try:
                # 같은 규정의 '현행' 버전 조회 (현재 rule_id와 다른 레코드)
                with db_manager.get_connection() as conn:
                    with conn.cursor() as cursor:
                        cursor.execute("""
                            SELECT wzRuleSeq, wzFileJson, wzpubno, wzname
                            FROM wz_rule
                            WHERE wzpubno = (SELECT wzpubno FROM wz_rule WHERE wzruleseq = %s)
                            AND wzNewFlag = '현행'
                            AND wzruleseq != %s
                        """, (rule_id, rule_id))
                        existing = cursor.fetchone()

                if existing and existing[1]:
                    old_rule_seq = existing[0]
                    old_json_path = existing[1]
                    rule_pubno = existing[2]
                    rule_name = existing[3]

                    # merge_json에 있는 파일인 경우만 이동
                    if 'merge_json/' in old_json_path and os.path.exists(old_json_path):
                        logger.info(f"Moving previous '현행' JSON to history: {old_json_path}")

                        # merge_json_old 폴더로 복사
                        old_json_dir = os.path.join(applib_path, 'merge_json_old')
                        os.makedirs(old_json_dir, exist_ok=True)

                        import shutil
                        old_filename = os.path.basename(old_json_path)
                        new_json_old_path = os.path.join(old_json_dir, old_filename)

                        try:
                            shutil.copy2(old_json_path, new_json_old_path)
                            logger.info(f"Copied JSON to history: {old_json_path} -> {new_json_old_path}")

                            # DB 업데이트: 이전 버전의 wzFileJson을 merge_json_old 경로로 변경
                            cursor.execute("""
                                UPDATE wz_rule
                                SET wzFileJson = %s
                                WHERE wzruleseq = %s
                            """, (new_json_old_path, old_rule_seq))
                            conn.commit()
                            logger.info(f"Updated wzFileJson for previous version (rule {old_rule_seq})")

                        except Exception as copy_error:
                            logger.warning(f"Failed to copy JSON to history: {copy_error}")

            except Exception as e:
                logger.warning(f"Error checking existing JSON: {e}")

            # 개정 처리 - wzNewFlag를 '현행'으로 변경하고 날짜 업데이트
            # merge_json 폴더에 저장할 준비
            from datetime import datetime
            import json

            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

            # 규정 정보 조회 (파일명용 - wzruleid 포함)
            with db_manager.get_connection() as conn:
                with conn.cursor() as cursor:
                    cursor.execute(
                        "SELECT wzpubno, wzname, wzruleid FROM wz_rule WHERE wzruleseq = %s",
                        (rule_id,)
                    )
                    rule_info = cursor.fetchone()

            rule_pubno = rule_info[0] if rule_info and rule_info[0] else f"rule_{rule_id}"
            rule_name = rule_info[1] if rule_info and rule_info[1] else "unknown"
            wzruleid = rule_info[2] if rule_info and len(rule_info) > 2 and rule_info[2] else 0
            safe_name = "".join(c for c in rule_name if c.isalnum() or c in (' ', '-')).replace(' ', '_')[:50]

            # merge_json에 저장
            merge_json_dir = os.path.join(applib_path, 'merge_json')
            os.makedirs(merge_json_dir, exist_ok=True)

            # 파일 업로드를 통해 생성된 병합 JSON이 있는 경우 기존 파일 재사용
            found_merged_json = None

            # 1순위: wzruleid.json 파일이 이미 존재하는지 확인
            wzruleid_json_path = os.path.join(merge_json_dir, f"{wzruleid}.json")
            if os.path.exists(wzruleid_json_path):
                logger.info(f"[Revision] Using existing wzruleid.json: {wzruleid_json_path}")
                found_merged_json = wzruleid_json_path
            # 2순위: 파라미터로 전달된 merged JSON 사용
            elif use_merged_json == 'true' and merged_json_path and os.path.exists(merged_json_path):
                logger.info(f"[Revision] Using existing merged JSON from parameter: {merged_json_path}")
                found_merged_json = merged_json_path
            else:
                # 3순위: merge_json 폴더에서 최근 파일 찾기 (이전 방식 호환)
                # (병합 후 바로 저장 시 발생하는 경우)
                import glob
                pattern = os.path.join(merge_json_dir, "merged_*.json")
                all_merged_files = glob.glob(pattern)

                # 최근 5분 이내에 생성된 파일 중 규정번호와 일치하는 파일 찾기
                from datetime import datetime, timedelta
                cutoff_time = datetime.now() - timedelta(minutes=5)

                recent_files = []
                for filepath in all_merged_files:
                    file_mtime = datetime.fromtimestamp(os.path.getmtime(filepath))
                    if file_mtime >= cutoff_time:
                        # 파일 내용 확인 (규정ID 또는 규정번호 일치 여부)
                        try:
                            with open(filepath, 'r', encoding='utf-8') as f:
                                file_json = json.load(f)
                                # 문서정보에 규정ID가 있거나, 규정번호가 일치하는 경우
                                if file_json.get("문서정보", {}).get("규정ID") == rule_id or \
                                   file_json.get("문서정보", {}).get("규정표기명", "").startswith(rule_pubno or ""):
                                    recent_files.append((filepath, file_mtime, file_json))
                        except:
                            pass

                # 가장 최근 파일 선택
                if recent_files:
                    recent_files.sort(key=lambda x: x[1], reverse=True)
                    found_merged_json = recent_files[0][0]
                    logger.info(f"[Revision] Auto-found recent merged JSON: {found_merged_json}")

            if found_merged_json and os.path.exists(found_merged_json):
                # 기존 병합 JSON 읽기
                with open(found_merged_json, 'r', encoding='utf-8') as f:
                    merged_json = json.load(f)

                # merged_로 시작하는 파일이면 wzruleid.json으로 rename
                if os.path.basename(found_merged_json).startswith("merged_"):
                    new_merged_filepath = os.path.join(merge_json_dir, f"{wzruleid}.json")
                    logger.info(f"[Revision] Renaming {found_merged_json} -> {new_merged_filepath}")
                    import shutil
                    shutil.move(found_merged_json, new_merged_filepath)
                    merged_filepath = new_merged_filepath
                else:
                    # 이미 wzruleid.json 형식인 경우 그대로 사용
                    merged_filepath = found_merged_json

                # 개정 정보만 업데이트 (나머지 문서정보는 모두 보존)
                if "문서정보" in merged_json:
                    merged_json["문서정보"]["규정ID"] = rule_id
                    merged_json["문서정보"]["개정일"] = revision_date
                    merged_json["문서정보"]["시행일"] = execution_date
                    merged_json["문서정보"]["최종개정일"] = revision_date
                    merged_json["문서정보"]["수정자"] = user.get('username', 'unknown')
                else:
                    merged_json["문서정보"] = {
                        "규정ID": rule_id,
                        "개정일": revision_date,
                        "시행일": execution_date,
                        "최종개정일": revision_date,
                        "수정자": user.get('username', 'unknown')
                    }

                # 편집된 텍스트로 조문내용 업데이트 (구조 유지)
                # 개정 모드에서는 파일 업로드로 생성된 병합 JSON의 내용을 그대로 사용
                # content 파라미터로 덮어쓰지 않음 (중복 번호 문제 방지)
                if content and content.strip():
                    logger.info(f"[Revision] Content parameter received (length: {len(content)}), but ignoring to preserve merged JSON structure")
                    logger.debug(f"[Revision] Content preview: {content[:200]}...")
                else:
                    logger.info("[Revision] No content parameter received, using merged JSON as-is")

                # 기존 파일을 덮어쓰기 (중복 생성 방지)
                with open(merged_filepath, 'w', encoding='utf-8') as f:
                    json.dump(merged_json, f, ensure_ascii=False, indent=2)
                logger.info(f"[Revision] Updated existing merged JSON: {merged_filepath}")

            else:
                # 병합 JSON이 없는 경우에만 새 파일 생성 (wzruleid 기반)
                merged_filename = f"{wzruleid}.json"
                merged_filepath = os.path.join(merge_json_dir, merged_filename)
                logger.info(f"[Revision] Creating new merged JSON: {merged_filepath}")

                if merged_json_data:
                    # 프론트엔드에서 전송한 JSON 데이터 사용
                    logger.info("Using merged JSON data from frontend")
                    merged_json = json.loads(merged_json_data)

                    # 개정 정보 업데이트
                    if "문서정보" not in merged_json:
                        merged_json["문서정보"] = {}
                    merged_json["문서정보"]["규정ID"] = rule_id  # 규정ID 누락 방지
                    merged_json["문서정보"]["개정일"] = revision_date
                    merged_json["문서정보"]["시행일"] = execution_date
                    merged_json["문서정보"]["수정자"] = user.get('username', 'unknown')
                else:
                    # 기존 방식: 텍스트 파싱 (간소화된 구조)
                    merged_json = {
                        "문서정보": {
                            "규정ID": rule_id,
                            "규정명": rule_name,
                            "개정일": revision_date,
                            "시행일": execution_date,
                            "수정자": user.get('username', 'unknown')
                        },
                        "조문내용": _parse_content_to_articles(content)
                    }

                with open(merged_filepath, 'w', encoding='utf-8') as f:
                    json.dump(merged_json, f, ensure_ascii=False, indent=2)
                logger.info(f"[Revision] Created new merged JSON: {merged_filepath}")

            # ✅ static/file 동기화 (개정 시)
            sync_success = sync_json_to_static_file(
                rule_id=rule_id,
                wzpubno=rule_pubno,
                new_merged_json_path=merged_filepath,
                mode="revision"
            )

            if sync_success:
                logger.info("✅ Successfully synced JSON to static/file and file_old")
            else:
                logger.warning("⚠️ Failed to sync JSON to static/file (continuing with DB update)")

            # 상대경로로 변환하여 DB에 저장
            from api.file_utils import get_relative_path
            merged_relative = get_relative_path(merged_filepath)

            query = """
                UPDATE wz_rule
                SET wzFileJson = %s,
                    wzlastrevdate = %s,
                    wzexecdate = %s,
                    wzNewFlag = %s,
                    wzmodifiedby = %s
                WHERE wzruleseq = %s
            """
            params = (merged_relative, revision_date, execution_date, '현행', user.get('username', 'admin'), rule_id)

        elif mode == 'new':
            # 제정 모드 - wzNewFlag를 '현행'으로 변경
            # merge_json 폴더에 저장
            from datetime import datetime
            import json

            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

            # 규정 정보 조회 (wzruleid 포함)
            with db_manager.get_connection() as conn:
                with conn.cursor() as cursor:
                    cursor.execute(
                        "SELECT wzpubno, wzname, wzruleid FROM wz_rule WHERE wzruleseq = %s",
                        (rule_id,)
                    )
                    rule_info = cursor.fetchone()

            rule_pubno = rule_info[0] if rule_info and rule_info[0] else f"rule_{rule_id}"
            rule_name = rule_info[1] if rule_info and rule_info[1] else "unknown"
            wzruleid = rule_info[2] if rule_info and len(rule_info) > 2 and rule_info[2] else 0

            merge_json_dir = os.path.join(applib_path, 'merge_json')
            os.makedirs(merge_json_dir, exist_ok=True)

            # wzruleid 기반 파일명 생성
            merged_filename = f"{wzruleid}.json"
            merged_filepath = os.path.join(merge_json_dir, merged_filename)

            # 파일 업로드를 통해 생성된 병합 JSON이 있는 경우
            if use_merged_json == 'true' and merged_json_path and os.path.exists(merged_json_path):
                logger.info(f"[New mode] Using merged JSON from upload: {merged_json_path}")
                # 기존 병합 JSON 읽기
                with open(merged_json_path, 'r', encoding='utf-8') as f:
                    merged_json = json.load(f)

                # 제정 정보 업데이트
                if "문서정보" not in merged_json:
                    merged_json["문서정보"] = {}
                merged_json["문서정보"]["규정ID"] = rule_id
                merged_json["문서정보"]["모드"] = "new"
                merged_json["문서정보"]["수정자"] = user.get('username', 'unknown')

                # ⚠️ content로 덮어쓰지 않음 (merged_json에 이미 레벨 정보가 포함되어 있음)
                if content and content.strip():
                    logger.info("[New mode] Content parameter received but ignoring to preserve merged JSON structure (including levels)")
                else:
                    logger.info("[New mode] No content parameter, using merged JSON as-is")

            elif merged_json_data:
                # 프론트엔드에서 전송한 JSON 데이터 사용 (레벨 정보 포함)
                logger.info("[New mode] Using merged JSON data from frontend (includes level info)")
                merged_json = json.loads(merged_json_data)

                # 제정 정보 업데이트
                if "문서정보" not in merged_json:
                    merged_json["문서정보"] = {}
                merged_json["문서정보"]["규정ID"] = rule_id
                merged_json["문서정보"]["모드"] = "new"
                merged_json["문서정보"]["수정자"] = user.get('username', 'unknown')

                # ⚠️ content 파라미터는 무시 (merged_json_data에 이미 레벨 정보가 포함되어 있음)
                logger.info("[New mode] Ignoring content parameter to preserve level information from merged_json_data")

            else:
                # 기존 방식: 텍스트 파싱
                merged_json = {
                    "문서정보": {
                        "규정ID": rule_id,
                        "규정명": rule_name,
                        "모드": "new",
                        "수정자": user.get('username', 'unknown')
                    },
                    "조문내용": _parse_content_to_articles(content)
                }

            with open(merged_filepath, 'w', encoding='utf-8') as f:
                json.dump(merged_json, f, ensure_ascii=False, indent=2)

            logger.info(f"New regulation saved to merge_json: {merged_filepath}")

            # ✅ static/file 동기화 (제정 시)
            sync_success = sync_json_to_static_file(
                rule_id=rule_id,
                wzpubno=rule_pubno,
                new_merged_json_path=merged_filepath,
                mode="new"
            )

            if sync_success:
                logger.info("✅ Successfully synced new regulation JSON to static/file")
            else:
                logger.warning("⚠️ Failed to sync new regulation JSON to static/file (continuing with DB update)")

            # DB에 저장할 상대 경로 (www/static/file/{wzruleid}.json)
            db_json_path = f"www/static/file/{wzruleid}.json"

            query = """
                UPDATE wz_rule
                SET wzFileJson = %s,
                    wzNewFlag = %s,
                    wzmodifiedby = %s
                WHERE wzruleseq = %s
            """
            params = (db_json_path, '현행', user.get('username', 'admin'), rule_id)

            # 색인 작업 트리거 (백그라운드)
            try:
                import threading
                from .indexing_service import index_single_regulation

                # 백그라운드에서 색인 실행
                threading.Thread(
                    target=index_single_regulation,
                    args=(db_manager, rule_id, rule_name, db_json_path),
                    daemon=True
                ).start()
                logger.info(f"Indexing triggered for rule {rule_id} (new mode)")
            except Exception as index_error:
                logger.warning(f"Failed to trigger indexing: {index_error}")

        else:
            # 일반 편집 - 내용만 업데이트
            # merge_json 폴더에 저장
            from datetime import datetime
            import json

            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

            # 규정 정보 조회
            with db_manager.get_connection() as conn:
                with conn.cursor() as cursor:
                    cursor.execute(
                        "SELECT wzpubno, wzname FROM wz_rule WHERE wzruleseq = %s",
                        (rule_id,)
                    )
                    rule_info = cursor.fetchone()

            rule_pubno = rule_info[0] if rule_info and rule_info[0] else f"rule_{rule_id}"
            rule_name = rule_info[1] if rule_info and rule_info[1] else "unknown"
            safe_name = "".join(c for c in rule_name if c.isalnum() or c in (' ', '-')).replace(' ', '_')[:50]

            merge_json_dir = os.path.join(applib_path, 'merge_json')
            os.makedirs(merge_json_dir, exist_ok=True)

            # 파일 업로드를 통해 생성된 병합 JSON이 있는 경우
            if use_merged_json == 'true' and merged_json_path and os.path.exists(merged_json_path):
                logger.info(f"[Edit mode] Using merged JSON from upload: {merged_json_path}")
                # 기존 병합 JSON 파일을 그대로 사용 (새 파일 생성하지 않음)
                merged_filepath = merged_json_path

                # 기존 병합 JSON 읽기
                with open(merged_json_path, 'r', encoding='utf-8') as f:
                    merged_json = json.load(f)

                # 편집 정보 업데이트
                if "문서정보" not in merged_json:
                    merged_json["문서정보"] = {}
                merged_json["문서정보"]["규정ID"] = rule_id
                merged_json["문서정보"]["모드"] = "edit"
                merged_json["문서정보"]["수정자"] = user.get('username', 'unknown')

                # ⚠️ content로 덮어쓰지 않음 (merged_json에 이미 레벨 정보가 포함되어 있음)
                if content and content.strip():
                    logger.info("[Edit mode] Content parameter received but ignoring to preserve merged JSON structure (including levels)")
                else:
                    logger.info("[Edit mode] No content parameter, using merged JSON as-is")

            elif merged_json_data:
                # 프론트엔드에서 전송한 JSON 데이터 사용 (새 파일 생성)
                logger.info("[Edit mode] Using merged JSON data from frontend")
                merged_json = json.loads(merged_json_data)

                # 편집 정보 업데이트
                if "문서정보" not in merged_json:
                    merged_json["문서정보"] = {}
                merged_json["문서정보"]["규정ID"] = rule_id
                merged_json["문서정보"]["모드"] = "edit"
                merged_json["문서정보"]["수정자"] = user.get('username', 'unknown')

                # 문서명 기반 파일명 생성
                document_name = extract_document_name_from_merged_data(merged_json)
                merged_filename = f"merged_{document_name}_{timestamp}.json"
                merged_filepath = os.path.join(merge_json_dir, merged_filename)

            else:
                # 기존 방식: 텍스트 파싱 (새 파일 생성)
                merged_json = {
                    "문서정보": {
                        "규정ID": rule_id,
                        "규정명": rule_name,
                        "모드": "edit",
                        "수정자": user.get('username', 'unknown')
                    },
                    "조문내용": _parse_content_to_articles(content)
                }

                # 문서명 기반 파일명 생성
                document_name = extract_document_name_from_merged_data(merged_json)
                merged_filename = f"merged_{document_name}_{timestamp}.json"
                merged_filepath = os.path.join(merge_json_dir, merged_filename)

            with open(merged_filepath, 'w', encoding='utf-8') as f:
                json.dump(merged_json, f, ensure_ascii=False, indent=2)

            logger.info(f"[FINAL] Edit saved to merge_json: {merged_filepath}")

            # ✅ static/file 동기화 (편집 시)
            sync_success = sync_json_to_static_file(
                rule_id=rule_id,
                wzpubno=rule_pubno,
                new_merged_json_path=merged_filepath,
                mode="edit"
            )

            if sync_success:
                logger.info("✅ Successfully synced edited JSON to static/file")
            else:
                logger.warning("⚠️ Failed to sync edited JSON to static/file (continuing with DB update)")

            # 상대경로로 변환하여 DB에 저장
            from api.file_utils import get_relative_path
            merged_relative = get_relative_path(merged_filepath)

            query = """
                UPDATE wz_rule
                SET wzFileJson = %s,
                    wzmodifiedby = %s
                WHERE wzruleseq = %s
            """
            params = (merged_relative, user.get('username', 'admin'), rule_id)

        # execute_query 메서드 사용, commit=True로 설정
        row_count = db_manager.execute_query(query, params, commit=True)

        if not row_count or row_count == 0:
            logger.error("Failed to update database - no rows affected")
            raise HTTPException(status_code=500, detail="데이터베이스 업데이트 실패 - 해당 규정을 찾을 수 없습니다")

        # 텍스트 백업 저장 (선택적)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        edited_dir = os.path.join(applib_path, 'edited')
        os.makedirs(edited_dir, exist_ok=True)

        txt_filename = f"edited_{rule_id}_{timestamp}.txt"
        txt_filepath = os.path.join(edited_dir, txt_filename)

        with open(txt_filepath, 'w', encoding='utf-8') as f:
            f.write(content)

        logger.info(f"Text backup saved to: {txt_filepath}")
        logger.info(f"Database updated successfully for rule {rule_id}")

        # 개정 모드인 경우 PDF를 static/pdf/print 폴더로 복사
        if mode == 'revision' or is_revision:
            try:
                # applib/pdf 폴더에서 최신 PDF 파일 찾기
                pdf_dir = os.path.join(applib_path, 'pdf')
                import glob
                # wzruleid를 포함하는 최신 PDF 파일 찾기
                pdf_pattern = os.path.join(pdf_dir, f"{wzruleid}_*.pdf")
                pdf_files = sorted(glob.glob(pdf_pattern), key=os.path.getmtime, reverse=True)

                if pdf_files:
                    source_pdf = pdf_files[0]  # 가장 최근 파일

                    # static/pdf/print 폴더로 복사
                    print_dir = f'{settings.WWW_STATIC_PDF_DIR}/print'
                    os.makedirs(print_dir, exist_ok=True)

                    # 파일명 형식: {규정번호}_{규정명}_{개정일자}개정.pdf
                    # 예: 1.1.1._정확한_환자_확인_202503개정.pdf
                    safe_pubno = rule_pubno.replace('.', '').strip('.')
                    safe_name_for_print = "".join(c if c.isalnum() or c in (' ', '-') else '_' for c in rule_name).replace(' ', '_')
                    revision_date_short = revision_date.replace('.', '').replace('-', '')[:6] if revision_date else timestamp[:6]
                    print_filename = f"{rule_pubno}_{safe_name_for_print}_{revision_date_short}개정.pdf"
                    dest_pdf = os.path.join(print_dir, print_filename)

                    import shutil
                    shutil.copy2(source_pdf, dest_pdf)
                    logger.info(f"✅ PDF copied to print folder: {dest_pdf}")

                    # wzFilePdf 컬럼 업데이트 (파일명만 저장)
                    try:
                        update_pdf_query = """
                            UPDATE wz_rule
                            SET wzFilePdf = %s
                            WHERE wzruleseq = %s
                        """
                        db_manager.execute_query(update_pdf_query, (print_filename, rule_id), commit=True)
                        logger.info(f"✅ wzFilePdf updated: {print_filename}")
                    except Exception as update_error:
                        logger.error(f"Error updating wzFilePdf: {update_error}")

                    # 나중에 JSON 파일에 PDF 파일명 저장할 수 있도록 변수에 저장
                    current_pdf_filename = print_filename
                else:
                    logger.warning(f"⚠️ No PDF file found for wzruleid {wzruleid} in {pdf_dir}")

            except Exception as pdf_copy_error:
                logger.error(f"Error copying PDF to print folder: {pdf_copy_error}")
                # PDF 복사 실패해도 계속 진행

            # DOCX 파일 경로 업데이트
            try:
                docx_dir = os.path.join(applib_path, 'docx')
                import glob
                docx_pattern = os.path.join(docx_dir, f"{wzruleid}_*.docx")
                docx_files = sorted(glob.glob(docx_pattern), key=os.path.getmtime, reverse=True)

                if docx_files:
                    latest_docx = docx_files[0]
                    docx_filename = os.path.basename(latest_docx)
                    docx_relative_path = f"applib/docx/{docx_filename}"

                    update_docx_query = """
                        UPDATE wz_rule
                        SET wzFileDocx = %s
                        WHERE wzruleseq = %s
                    """
                    db_manager.execute_query(update_docx_query, (docx_relative_path, rule_id), commit=True)
                    logger.info(f"✅ wzFileDocx updated: {docx_relative_path}")
                else:
                    logger.warning(f"⚠️ No DOCX file found for wzruleid {wzruleid} in {docx_dir}")

            except Exception as docx_update_error:
                logger.error(f"Error updating wzFileDocx: {docx_update_error}")

        # JSON_ALL.py와 create_summary.py 백그라운드 실행
        background_tasks.add_task(run_json_merge_and_summary)
        logger.info(f"✅ 내용 저장 완료 (mode: {mode}), 백그라운드 병합 작업 예약됨")

        # 저장된 JSON 파일 정보 가져오기
        json_filepath = None
        json_filename = None

        # DB에서 방금 업데이트한 wzFileJson 경로 가져오기
        with db_manager.get_connection() as conn:
            with conn.cursor() as cursor:
                cursor.execute(
                    "SELECT wzFileJson FROM wz_rule WHERE wzruleseq = %s",
                    (rule_id,)
                )
                result = cursor.fetchone()
                if result and result[0]:
                    json_filepath = result[0]
                    json_filename = os.path.basename(json_filepath)

        return {
            "success": True,
            "message": "내용이 저장되었습니다. (백그라운드에서 병합 중)",
            "files": {
                "txt_filename": txt_filename,
                "txt_filepath": txt_filepath,
                "json_filename": json_filename,
                "json_filepath": json_filepath
            },
            "is_revision": is_revision,
            "revision_date": revision_date if is_revision else None,
            "execution_date": execution_date if is_revision else None
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error saving edited content: {e}")
        logger.error(f"Error details:", exc_info=True)
        raise HTTPException(status_code=500, detail=f"저장 중 오류 발생: {str(e)}")


# ========== 비동기 처리 엔드포인트 ==========
from .service_rule_editor_async import (
    upload_and_parse_pdf_async,
    upload_and_parse_docx_async,
    get_job_status,
    AsyncDocumentProcessor
)

@router.post("/upload-parse-pdf-async")
async def upload_parse_pdf_async_endpoint(
    background_tasks: BackgroundTasks,
    pdf_file: UploadFile = File(...),
    rule_id: int = Form(...),
    user: Dict[str, Any] = Depends(get_current_user)
):
    """PDF 파일 비동기 업로드 및 파싱 (즉시 응답)"""
    try:
        # 파일 내용 읽기
        pdf_content = await pdf_file.read()

        # 비동기 처리 시작
        result = await upload_and_parse_pdf_async(
            pdf_content,
            pdf_file.filename,
            rule_id,
            background_tasks
        )

        return result
    except Exception as e:
        logger.error(f"Error in async PDF upload: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/upload-parse-docx-async")
async def upload_parse_docx_async_endpoint(
    background_tasks: BackgroundTasks,
    docx_file: UploadFile = File(...),
    rule_id: int = Form(...),
    user: Dict[str, Any] = Depends(get_current_user)
):
    """DOCX 파일 비동기 업로드 및 파싱 (즉시 응답)"""
    try:
        # 파일 내용 읽기
        docx_content = await docx_file.read()

        # 비동기 처리 시작
        result = await upload_and_parse_docx_async(
            docx_content,
            docx_file.filename,
            rule_id,
            background_tasks
        )

        return result
    except Exception as e:
        logger.error(f"Error in async DOCX upload: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/job-status/{job_id}")
async def get_job_status_endpoint(
    job_id: str,
    user: Dict[str, Any] = Depends(get_current_user)
):
    """작업 상태 조회"""
    try:
        return await get_job_status(job_id)
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting job status: {e}")
        raise HTTPException(status_code=500, detail=str(e))

def run_json_merge_and_summary():
    """
    백그라운드에서 JSON 병합 및 요약 파일 생성
    1. JSON_ALL.py 실행 (merge_json 폴더의 모든 JSON 병합)
    2. create_summary.py 실행 (요약 파일 생성)
    """
    try:
        logger.info("🔄 백그라운드 작업 시작: JSON 병합 및 요약 생성")

        # 1. JSON_ALL.py 실행
        logger.info("  Step 1: JSON_ALL.py 실행 중...")
        result1 = subprocess.run(
            [PYTHON_EXECUTABLE, JSON_ALL_SCRIPT],
            capture_output=True,
            text=True,
            timeout=60  # 60초 타임아웃
        )

        if result1.returncode == 0:
            logger.info("  ✅ JSON_ALL.py 완료")
            logger.debug(f"  출력: {result1.stdout[:200]}")
        else:
            logger.error(f"  ❌ JSON_ALL.py 실패: {result1.stderr}")
            return False

        # 2. wzRuleSeq 추가 (merged_severance.json에)
        logger.info("  Step 2: wzRuleSeq 추가 중...")
        add_wzruleseq_script = '/tmp/add_wzruleseq_to_both_jsons.py'
        result2 = subprocess.run(
            [PYTHON_EXECUTABLE, add_wzruleseq_script],
            capture_output=True,
            text=True,
            timeout=30
        )

        if result2.returncode == 0:
            logger.info("  ✅ wzRuleSeq 추가 완료")
        else:
            logger.warning(f"  ⚠️ wzRuleSeq 추가 실패 (계속 진행): {result2.stderr[:200]}")

        # 3. DB로부터 직접 summary 생성 (신규 방식)
        logger.info("  Step 3: DB로부터 summary 생성 중...")
        summary_output = f'{settings.WWW_STATIC_FILE_DIR}/summary_kbregulation.json'
        summary_success = generate_summary_from_db(summary_output)

        if summary_success:
            logger.info("  ✅ summary 생성 완료")
        else:
            logger.error("  ❌ summary 생성 실패")
            return False

        logger.info("✅ 백그라운드 작업 완료: 병합 파일 및 요약 파일 생성됨")
        return True

    except subprocess.TimeoutExpired:
        logger.error("❌ 백그라운드 작업 타임아웃")
        return False
    except Exception as e:
        logger.error(f"❌ 백그라운드 작업 실패: {e}")
        return False

@router.put("/update-json/{rule_id}")
async def update_existing_json(
    rule_id: int,
    request: Request,
    background_tasks: BackgroundTasks,
    user: Dict[str, Any] = Depends(get_current_user)
):
    """기존 JSON 파일 덮어쓰기 + 백그라운드 병합"""
    try:
        data = await request.json()
        updated_json = data.get('json_data')

        if not updated_json:
            raise HTTPException(status_code=400, detail="JSON 데이터 없음")

        db_config = {
            'host': settings.DB_HOST,
            'port': settings.DB_PORT,
            'database': settings.DB_NAME,
            'user': settings.DB_USER,
            'password': settings.DB_PASSWORD,
        }
        db_manager = DatabaseConnectionManager(**db_config)

        with db_manager.get_connection() as conn:
            with conn.cursor() as cur:
                cur.execute("SELECT wzFileJson FROM wz_rule WHERE wzruleseq = %s", (rule_id,))
                row = cur.fetchone()

                if not row or not row[0]:
                    raise HTTPException(status_code=404, detail="JSON 경로 없음")

                json_path_from_db = row[0]

                # 상대경로를 절대경로로 변환
                from api.file_utils import get_absolute_path
                json_path = get_absolute_path(json_path_from_db)

                if not os.path.exists(json_path):
                    raise HTTPException(status_code=404, detail=f"파일 없음: {json_path}")

                # ===== 추가: preview_data 동기화 =====
                if 'preview_data' in updated_json:
                    updated_json['preview_data']['조문내용'] = updated_json['조문내용']
                    updated_json['preview_data']['문서정보']['조문갯수'] = len(updated_json['조문내용'])

                # ===== 추가: 문서정보 조문갯수 동기화 =====
                if '문서정보' in updated_json:
                    updated_json['문서정보']['조문갯수'] = len(updated_json['조문내용'])

                # 덮어쓰기 (필수: ensure_ascii=False)
                with open(json_path, 'w', encoding='utf-8') as f:
                    json.dump(updated_json, f, ensure_ascii=False, indent=2)

                # ===== 추가: www/static/file 경로에도 복사 =====
                try:
                    www_dir = settings.WWW_STATIC_FILE_DIR

                    # 원본 파일명 그대로 사용
                    filename = os.path.basename(json_path)
                    www_path = f"{www_dir}/{filename}"

                    # 디렉토리 생성
                    os.makedirs(www_dir, exist_ok=True)

                    # 파일 복사
                    shutil.copy2(json_path, www_path)
                    logger.info(f"✅ www/static/file 덮어쓰기 완료: {filename}")

                except Exception as copy_error:
                    logger.warning(f"⚠️ www/static/file 복사 실패 (무시): {copy_error}")

                # content_text 동기화
                text_content = ""
                for article in updated_json.get('조문내용', []):
                    if article.get('번호'):
                        text_content += f"{article['번호']} {article['내용']}\n\n"

                cur.execute("""
                    UPDATE wz_rule
                    SET content_text = %s, wzModifiedBy = %s
                    WHERE wzruleseq = %s
                """, (text_content.strip(), user.get('username'), rule_id))

                conn.commit()

                # 백그라운드에서 JSON 병합 및 요약 생성
                background_tasks.add_task(run_json_merge_and_summary)
                logger.info(f"✅ JSON 파일 업데이트 완료 (rule_id={rule_id}), 백그라운드 병합 작업 예약됨")

                return {
                    "success": True,
                    "message": "JSON 업데이트 완료 (병합 작업은 백그라운드에서 진행 중)",
                    "articles_count": len(updated_json.get('조문내용', []))
                }
    except Exception as e:
        logger.error(f"Error updating JSON: {e}")
        raise HTTPException(status_code=500, detail=str(e))